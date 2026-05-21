import os
import io
import json
import re
import time
import uuid
import markdown
from flask import Flask, render_template, request, send_file
from dotenv import load_dotenv
import yfinance as yf
import anthropic
import google.generativeai as genai
from google.api_core import exceptions as google_exceptions
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm, mm
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, KeepTogether, ListFlowable, ListItem, Image as RLImage,
)
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.ticker import FuncFormatter

load_dotenv()

import sys as _sys
import logging as _logging
from pathlib import Path as _Path
import rag as _rag

# Ticker activo durante la sesión de la app. Se arranca sin ticker
# (`python app.py`) y el usuario selecciona la empresa al entrar en la
# home a través de un selector que lista las carpetas existentes bajo
# `Empresas/`. Compatibilidad: si se pasa `python app.py META`, se usa
# directamente y se salta el selector.
# `--recrawl` fuerza la re-indexación del corpus RAG aunque el caché esté
# fresco.
_argv = [a for a in _sys.argv[1:] if a.strip()]
RECRAWL_AL_ARRANCAR = "--recrawl" in _argv
_argv_sin_flags = [a for a in _argv if not a.startswith("--")]
TICKER_FIJO = _argv_sin_flags[0].strip().upper() if _argv_sin_flags else None

# Logging: append continuo a `Empresas/{TICKER}/auditoria/sesion.txt` para que
# cada compañía tenga su propio histórico aislado. Si arrancas sin TICKER_FIJO
# (modo libre), cae al fallback `auditoria/` en la raíz.
# Si quieres más detalle, define LOG_LEVEL=DEBUG en .env.
_ROOT_DIR_LOG = _Path(__file__).resolve().parent
if TICKER_FIJO:
    _AUDITORIA_DIR = _ROOT_DIR_LOG / "Empresas" / TICKER_FIJO / "auditoria"
else:
    _AUDITORIA_DIR = _ROOT_DIR_LOG / "auditoria"
_AUDITORIA_DIR.mkdir(parents=True, exist_ok=True)
_LOG_FILE = _AUDITORIA_DIR / "sesion.txt"
_LOG_LEVEL = os.getenv("LOG_LEVEL", "INFO").upper()

_log_formatter = _logging.Formatter(
    fmt="[%(asctime)s.%(msecs)03d] %(levelname)-5s %(name)-18s %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
_file_handler = _logging.FileHandler(_LOG_FILE, mode="a", encoding="utf-8")
_file_handler.setFormatter(_log_formatter)
_stream_handler = _logging.StreamHandler()
_stream_handler.setFormatter(_log_formatter)
_logging.basicConfig(
    level=getattr(_logging, _LOG_LEVEL, _logging.INFO),
    handlers=[_file_handler, _stream_handler],
    force=True,
)
# Silenciamos ruido de librerías de terceros (peticiones HTTP, yfinance, etc.).
for _ruidoso in ("urllib3", "yfinance", "google", "anthropic", "httpx", "httpcore"):
    _logging.getLogger(_ruidoso).setLevel(_logging.WARNING)

logger = _logging.getLogger("dfin")
logger.info("=" * 70)
logger.info(
    "Módulo cargado · TICKER=%s · LOG_LEVEL=%s · fichero=%s",
    TICKER_FIJO or "(ninguno)", _LOG_LEVEL, _LOG_FILE,
)


# --- Coste de llamadas a la IA --------------------------------------------
#
# Tarifa pública por millón de tokens (USD). Actualizar si los proveedores
# cambian precios. Si un modelo no está mapeado, se usa el default Sonnet.

_PRECIOS_TOKENS_USD_M = {
    "claude-opus-4-7":              {"input": 15.00, "output": 75.00},
    "claude-opus-4-6":              {"input": 15.00, "output": 75.00},
    "claude-sonnet-4-6":            {"input":  3.00, "output": 15.00},
    "claude-haiku-4-5-20251001":    {"input":  0.80, "output":  4.00},
    "gemini-2.5-flash":             {"input":  0.30, "output":  2.50},
    "gemini-2.5-flash-lite":        {"input":  0.10, "output":  0.40},
}
_PRECIO_DEFECTO = {"input": 3.00, "output": 15.00}
_COSTE_FILE = _AUDITORIA_DIR / "coste_tokens.txt"

# Acumulado en memoria para no leer/parsear el fichero en cada llamada.
_COSTE_ACUMULADO = {"usd": 0.0, "llamadas": 0}


def _coste_usd(modelo, in_toks, out_toks):
    tarifa = _PRECIOS_TOKENS_USD_M.get(modelo, _PRECIO_DEFECTO)
    return (in_toks * tarifa["input"] + out_toks * tarifa["output"]) / 1_000_000


def _registrar_coste(modelo, modulo, usage):
    """Calcula el coste de una llamada y lo añade a coste_tokens.txt (append)."""
    in_toks = int(usage.get("input_tokens", 0) or 0)
    out_toks = int(usage.get("output_tokens", 0) or 0)
    if in_toks == 0 and out_toks == 0:
        return  # nada que registrar
    coste = _coste_usd(modelo, in_toks, out_toks)
    _COSTE_ACUMULADO["usd"] += coste
    _COSTE_ACUMULADO["llamadas"] += 1
    from datetime import datetime as _dt
    ts = _dt.now().strftime("%Y-%m-%d %H:%M:%S")
    linea = (
        f"[{ts}] modulo={modulo} modelo={modelo} "
        f"input_tokens={in_toks} output_tokens={out_toks} "
        f"coste_usd={coste:.6f} acumulado_usd={_COSTE_ACUMULADO['usd']:.6f} "
        f"llamada_n={_COSTE_ACUMULADO['llamadas']}\n"
    )
    _AUDITORIA_DIR.mkdir(parents=True, exist_ok=True)
    with open(_COSTE_FILE, "a", encoding="utf-8") as f:
        f.write(linea)
    _logging.getLogger("dfin.coste").info(
        "Coste %s · %s · in=%d out=%d → $%.6f (acumulado $%.4f)",
        modulo, modelo, in_toks, out_toks, coste, _COSTE_ACUMULADO["usd"],
    )


def _cargar_acumulado_coste():
    """Al arrancar, recupera el acumulado leyendo la última línea de coste_tokens.txt."""
    if not _COSTE_FILE.exists():
        return
    try:
        contenido = _COSTE_FILE.read_text(encoding="utf-8").splitlines()
        for linea in reversed(contenido):
            if "acumulado_usd=" not in linea:
                continue
            partes = dict(p.split("=", 1) for p in linea.split() if "=" in p)
            _COSTE_ACUMULADO["usd"] = float(partes.get("acumulado_usd", 0.0))
            _COSTE_ACUMULADO["llamadas"] = int(partes.get("llamada_n", 0))
            return
    except Exception as e:
        _logging.getLogger("dfin.coste").warning("No se pudo cargar acumulado de coste: %s", e)


_cargar_acumulado_coste()

app = Flask(__name__)

# Branding mutable: arranca con el ticker (o el genérico) y se rellena con el
# nombre real de la compañía cuando se cargan datos de Yahoo por primera vez.
_BRANDING = {
    "empresa": TICKER_FIJO or "DFin AI",
    "tenant_sufijo": TICKER_FIJO or "DFin AI",
}


_ROOT_DIR = _Path(__file__).resolve().parent


# --- Tema visual personalizable por empresa --------------------------------
#
# Carga `Empresas/{TICKER}/tema.txt` (formato clave=valor por línea, líneas
# que empiezan por # se ignoran). Valores aceptados:
#
#   color_primario          color principal de botones, links activos, badges
#   color_primario_oscuro   variante oscura (hover, contornos)
#   color_secundario        color secundario (acentos, second-CTA)
#   color_acento            color de detalle (línea borde top en cards, etc.)
#   color_texto_topbar      color del texto del topbar (por defecto blanco)
#   fondo_topbar            fondo del topbar (por defecto var(--black))
#
# Todos los valores deben ser colores hex válidos (#RGB o #RRGGBB).
# Si el fichero no existe o un valor no es válido, se ignora y se mantiene
# el default.

_TEMA_CLAVES_VALIDAS = {
    "color_primario", "color_primario_oscuro", "color_secundario",
    "color_acento", "color_texto_topbar", "fondo_topbar",
}
_HEX_RE = re.compile(r"^#(?:[0-9a-fA-F]{3}|[0-9a-fA-F]{6})$")


def _color_texto_sobre(hex_color):
    """Devuelve '#FFFFFF' o '#000000' según el contraste óptimo con `hex_color`.

    Usa la luminancia relativa W3C (perceptiva), no el promedio RGB.
    Threshold L < 0.5 → texto blanco; en otro caso → negro.
    """
    if not hex_color or not _HEX_RE.match(hex_color):
        return "#FFFFFF"
    h = hex_color.lstrip("#")
    if len(h) == 3:
        h = "".join(c * 2 for c in h)
    try:
        r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
    except ValueError:
        return "#FFFFFF"
    def lin(c):
        c = c / 255
        return c / 12.92 if c <= 0.03928 else ((c + 0.055) / 1.055) ** 2.4
    L = 0.2126 * lin(r) + 0.7152 * lin(g) + 0.0722 * lin(b)
    return "#FFFFFF" if L < 0.5 else "#000000"


def _cargar_tema_empresa():
    """Lee Empresas/{TICKER}/tema.txt si existe y devuelve dict sanitizado.

    Añade automáticamente colores de texto con contraste correcto para que
    los botones primarios y badges sean legibles independientemente del
    color de fondo elegido.
    """
    if not TICKER_FIJO:
        return {}
    ruta = _ROOT_DIR / "Empresas" / TICKER_FIJO / "tema.txt"
    if not ruta.exists():
        return {}
    tema = {}
    for linea in ruta.read_text(encoding="utf-8").splitlines():
        s = linea.strip()
        if not s or s.startswith("#") or "=" not in s:
            continue
        clave, _, valor = s.partition("=")
        clave = clave.strip().lower()
        valor = valor.strip()
        if clave not in _TEMA_CLAVES_VALIDAS:
            continue
        if not _HEX_RE.match(valor):
            _logging.getLogger("dfin.tema").warning(
                "tema.txt: valor inválido para %s: %r (ignorado)", clave, valor,
            )
            continue
        tema[clave] = valor

    # Texto auto-calculado para mantener contraste sobre los fondos primarios.
    if tema.get("color_primario"):
        tema["color_texto_primario"] = _color_texto_sobre(tema["color_primario"])
    if tema.get("color_primario_oscuro"):
        tema["color_texto_primario_oscuro"] = _color_texto_sobre(tema["color_primario_oscuro"])

    if tema:
        _logging.getLogger("dfin.tema").info(
            "Tema cargado para %s: %s", TICKER_FIJO, ", ".join(f"{k}={v}" for k, v in tema.items()),
        )
    return tema


# --- Líneas maestras del equipo de IR --------------------------------------
#
# Fichero `Empresas/{TICKER}/lineas_maestras.txt` (texto libre) con las
# normas, criterios o comentarios de la dirección sobre cómo responder a
# los correos de inversores. Tiene PRIORIDAD ABSOLUTA sobre cualquier otra
# fuente: PDFs, web, Yahoo o búsqueda. La IA debe seguirlas siempre y
# alinear cualquier respuesta con su contenido.
# La app permite crearlo y editarlo desde la propia UI del módulo IR.

_LINEAS_MAESTRAS_MAX_BYTES = 100_000   # límite defensivo


def _ruta_lineas_maestras():
    if not TICKER_FIJO:
        return None
    return _ROOT_DIR / "Empresas" / TICKER_FIJO / "lineas_maestras.txt"


def _cargar_lineas_maestras():
    """Devuelve el contenido del fichero o '' si no existe / no hay ticker."""
    ruta = _ruta_lineas_maestras()
    if not ruta or not ruta.exists():
        return ""
    try:
        return ruta.read_text(encoding="utf-8").strip()
    except OSError as e:
        _logging.getLogger("dfin.ir.lineas").warning("No se pudo leer %s: %s", ruta, e)
        return ""


def _guardar_lineas_maestras(contenido):
    """Escribe (o crea) el fichero. Devuelve True si éxito."""
    ruta = _ruta_lineas_maestras()
    if not ruta:
        raise RuntimeError("No hay TICKER_FIJO definido.")
    if contenido and len(contenido.encode("utf-8")) > _LINEAS_MAESTRAS_MAX_BYTES:
        raise ValueError(
            f"El contenido supera el límite de {_LINEAS_MAESTRAS_MAX_BYTES // 1000} KB."
        )
    ruta.parent.mkdir(parents=True, exist_ok=True)
    ruta.write_text((contenido or "").strip() + "\n", encoding="utf-8")
    _logging.getLogger("dfin.ir.lineas").info(
        "Líneas maestras actualizadas para %s (%d caracteres)", TICKER_FIJO, len(contenido or ""),
    )
    return True


def _ruta_logo_empresa():
    if not TICKER_FIJO:
        return None
    ruta = _ROOT_DIR / "Empresas" / TICKER_FIJO / "logo" / "logo.png"
    return ruta if ruta.exists() else None


def _ruta_logo_deloitte():
    """Busca el logo de Deloitte en varias ubicaciones razonables.

    Devuelve la primera que exista o None. Probamos nombres comunes
    (`deloitte.png`, `deloitte_logo.png`, `logo_deloitte.png` y como
    último recurso `logo.png`) en `static/` y `static/img/` para no
    obligar a un nombre exacto. Los nombres más explícitos tienen
    prioridad sobre `logo.png` por si conviven con otro logo.
    """
    candidatos = [
        _ROOT_DIR / "static" / "deloitte_logo.png",
        _ROOT_DIR / "static" / "deloitte.png",
        _ROOT_DIR / "static" / "logo_deloitte.png",
        _ROOT_DIR / "static" / "img" / "deloitte_logo.png",
        _ROOT_DIR / "static" / "img" / "deloitte.png",
        _ROOT_DIR / "static" / "img" / "logo_deloitte.png",
        _ROOT_DIR / "static" / "logo.png",
        _ROOT_DIR / "static" / "img" / "logo.png",
    ]
    for ruta in candidatos:
        if ruta.exists():
            return ruta
    return None


@app.context_processor
def _inyectar_branding():
    ruta_emp = _ruta_logo_empresa()
    if ruta_emp:
        try:
            v = int(ruta_emp.stat().st_mtime)
        except OSError:
            v = 0
        empresa_url = f"/empresa-logo?v={v}"
    else:
        empresa_url = None
    ruta_del = _ruta_logo_deloitte()
    if ruta_del:
        try:
            rel = ruta_del.relative_to(_ROOT_DIR / "static")
            v = int(ruta_del.stat().st_mtime)
            deloitte_url = "/static/" + str(rel).replace("\\", "/") + f"?v={v}"
        except (ValueError, OSError):
            deloitte_url = "/deloitte-logo"
    else:
        deloitte_url = None
    return {
        "TICKER_FIJO": TICKER_FIJO,
        "EMPRESA_NOMBRE": _BRANDING["empresa"],
        "TENANT_NOMBRE": _BRANDING["tenant_sufijo"],
        "EMPRESA_LOGO_DISPONIBLE": ruta_emp is not None,
        "EMPRESA_LOGO_URL": empresa_url,
        "DELOITTE_LOGO_DISPONIBLE": ruta_del is not None,
        "DELOITTE_LOGO_URL": deloitte_url,
        "TEMA_EMPRESA": _cargar_tema_empresa(),
        "COSTE_ACUMULADO_USD": _COSTE_ACUMULADO.get("usd", 0.0),
        "COSTE_LLAMADAS": _COSTE_ACUMULADO.get("llamadas", 0),
    }


@app.route("/empresa-logo")
def empresa_logo():
    from flask import send_file, abort, make_response
    ruta = _ruta_logo_empresa()
    if not ruta:
        abort(404)
    # max_age=0 + Cache-Control restrictivo: el navegador hará revalidación.
    # El cache-buster ?v=mtime en la URL del template ya garantiza no
    # mezclar logos entre instancias.
    resp = make_response(send_file(str(ruta), mimetype="image/png", max_age=0))
    resp.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
    resp.headers["Pragma"] = "no-cache"
    resp.headers["Expires"] = "0"
    return resp

ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY")
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

# Credenciales de Gmail. El módulo Relación con Inversores lee los correos
# cuyo asunto empieza por `f"{TICKER_FIJO}-{IR_PREFIJO_ASUNTO}-"` (por defecto
# "META-Inversores-" si se arrancó con `python app.py META`) recibidos en las
# últimas dos horas. Si faltan credenciales, el módulo muestra una vista de
# error pidiendo configurarlas.
GMAIL_USER = os.getenv("GMAIL_USER")
GMAIL_APP_PASSWORD = os.getenv("GMAIL_APP_PASSWORD")
IR_PREFIJO_ASUNTO = os.getenv("IR_PREFIJO_ASUNTO", "Inversores")


def _prefijo_asunto_actual():
    base = (IR_PREFIJO_ASUNTO or "Inversores").strip()
    if TICKER_FIJO:
        return f"{TICKER_FIJO}-{base}-"
    return f"{base}-"

if GOOGLE_API_KEY:
    genai.configure(api_key=GOOGLE_API_KEY)

MODELOS_ANTHROPIC = {
    "claude-opus-4-6",
    "claude-sonnet-4-6",
    "claude-haiku-4-5-20251001",
}
MODELOS_GOOGLE = {
    "gemini-2.5-flash",
    "gemini-2.5-flash-lite",
}
MODELOS_DISPONIBLES = MODELOS_ANTHROPIC | MODELOS_GOOGLE
MODELO_POR_DEFECTO = "claude-sonnet-4-6"

# Regex compartidas para detectar tablas Markdown producidas por la IA
_RE_TABLA_SEP_MD = re.compile(r"^\s*\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)+\|?\s*$")
_RE_TABLA_FILA_MD = re.compile(r"^\s*\|.+\|\s*$")

# Cache en memoria para evitar re-llamar a las APIs al descargar Word
_download_cache = {}


def _normalizar_modelo(modelo):
    """Devuelve el modelo si es válido, o el modelo por defecto."""
    if modelo in MODELOS_DISPONIBLES:
        return modelo
    return MODELO_POR_DEFECTO


def _extraer_respuesta_claude(message):
    """Extrae texto + citas + uso de tokens de una respuesta Claude."""
    partes = []
    citas = []
    for block in getattr(message, "content", []) or []:
        btype = getattr(block, "type", None)
        if btype == "text":
            t = getattr(block, "text", "") or ""
            if t:
                partes.append(t)
            for cit in getattr(block, "citations", None) or []:
                url = getattr(cit, "url", None)
                title = getattr(cit, "title", None) or url
                if url:
                    entrada = {"url": url, "title": title}
                    if entrada not in citas:
                        citas.append(entrada)
    usage_obj = getattr(message, "usage", None)
    usage = {
        "input_tokens": int(getattr(usage_obj, "input_tokens", 0) or 0),
        "output_tokens": int(getattr(usage_obj, "output_tokens", 0) or 0),
    }
    return "".join(partes).strip(), citas, usage


def llamar_claude_con_reintentos(prompt, max_tokens, modelo, max_reintentos=4,
                                  permitir_busqueda=False):
    """Llama a la API de Anthropic con reintentos y backoff exponencial.

    Si `permitir_busqueda` es True, activa la herramienta `web_search` nativa
    (server-side managed) para que el modelo pueda consultar internet. Devuelve
    una tupla `(texto, citas)` donde `citas` es una lista de dicts con url/title.
    """
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY, max_retries=0)
    errores_reintentables = (
        anthropic.APIStatusError,
        anthropic.APIConnectionError,
        anthropic.RateLimitError,
    )
    status_reintentables = {429, 500, 502, 503, 504, 529}

    extra_kwargs = {}
    if permitir_busqueda:
        extra_kwargs["tools"] = [{
            "type": "web_search_20250305",
            "name": "web_search",
            "max_uses": 5,
        }]

    ultimo_error = None
    for intento in range(max_reintentos):
        try:
            message = client.messages.create(
                model=modelo,
                max_tokens=max_tokens,
                messages=[{"role": "user", "content": prompt}],
                **extra_kwargs,
            )
            return _extraer_respuesta_claude(message)  # (texto, citas, usage)
        except errores_reintentables as e:
            ultimo_error = e
            status = getattr(e, "status_code", None)
            if status is not None and status not in status_reintentables:
                raise
            if intento < max_reintentos - 1:
                espera = 2 ** intento
                time.sleep(espera)
                continue
            break

    if ultimo_error is not None:
        status = getattr(ultimo_error, "status_code", None)
        if status == 529:
            raise RuntimeError(
                "El servicio de IA está sobrecargado en este momento. "
                "Por favor, inténtalo de nuevo en unos segundos."
            )
        if status == 429:
            raise RuntimeError(
                "Se ha superado el límite de peticiones a la IA. "
                "Por favor, espera un momento antes de volver a intentarlo."
            )
        raise RuntimeError(
            f"Error al comunicar con el servicio de IA: {str(ultimo_error)}"
        )


def _extraer_texto_gemini(respuesta):
    """Extrae texto de una respuesta Gemini, o relanza un error legible.

    Maneja los casos en que `response.text` falla por bloqueo de seguridad,
    truncado por `MAX_TOKENS`, o candidatos vacíos.
    """
    candidatos = getattr(respuesta, "candidates", None) or []
    if not candidatos:
        prompt_feedback = getattr(respuesta, "prompt_feedback", None)
        raise RuntimeError(
            f"Gemini no devolvió respuesta. Detalle: {prompt_feedback}"
        )

    candidato = candidatos[0]
    partes = []
    contenido = getattr(candidato, "content", None)
    if contenido and getattr(contenido, "parts", None):
        for parte in contenido.parts:
            texto_parte = getattr(parte, "text", None)
            if texto_parte:
                partes.append(texto_parte)

    texto = "".join(partes).strip()
    if texto:
        return texto

    finish_reason = getattr(candidato, "finish_reason", None)
    if finish_reason and str(finish_reason).endswith("MAX_TOKENS"):
        raise RuntimeError(
            "Gemini agotó el presupuesto de tokens antes de producir texto "
            "(probablemente por el proceso de razonamiento interno). "
            "Inténtalo de nuevo o selecciona otro modelo."
        )
    raise RuntimeError(
        f"Gemini no generó texto (finish_reason={finish_reason})."
    )


def _extraer_citas_gemini(respuesta):
    """Extrae las fuentes usadas en grounding de Google Search (Gemini)."""
    citas = []
    candidatos = getattr(respuesta, "candidates", None) or []
    if not candidatos:
        return citas
    gm = getattr(candidatos[0], "grounding_metadata", None)
    if not gm:
        return citas
    for chunk in getattr(gm, "grounding_chunks", None) or []:
        web = getattr(chunk, "web", None)
        if not web:
            continue
        url = getattr(web, "uri", None)
        title = getattr(web, "title", None) or url
        if url:
            entrada = {"url": url, "title": title}
            if entrada not in citas:
                citas.append(entrada)
    return citas


def llamar_gemini_con_reintentos(prompt, max_tokens, modelo, max_reintentos=4,
                                  permitir_busqueda=False):
    """Llama a la API de Google Gemini con reintentos y backoff exponencial.

    Si `permitir_busqueda` es True, activa grounding con Google Search para
    que Gemini consulte internet. Devuelve `(texto, citas)`.
    """
    if not GOOGLE_API_KEY:
        raise RuntimeError(
            "No se ha configurado GOOGLE_API_KEY en el archivo .env. "
            "Añádela para usar los modelos de Google Gemini."
        )

    errores_reintentables = (
        google_exceptions.ResourceExhausted,
        google_exceptions.ServiceUnavailable,
        google_exceptions.InternalServerError,
        google_exceptions.DeadlineExceeded,
        google_exceptions.Aborted,
    )

    max_output_tokens = max(max_tokens * 4, 8192)

    # Gemini 2.5+ usa `google_search`. La variante antigua
    # `google_search_retrieval` (Gemini 1.5) ya no está soportada y devuelve
    # 400, por eso no la usamos como fallback. Si la búsqueda con
    # `google_search` falla, reintentamos SIN herramientas para que el
    # contenido siga generándose con el conocimiento del modelo + el prompt.
    tools_config = None
    if permitir_busqueda:
        tools_config = [{"google_search": {}}]

    def _ejecutar(tools):
        kwargs = dict(
            generation_config={
                "max_output_tokens": max_output_tokens,
                "temperature": 0.7,
            },
        )
        if tools is not None:
            kwargs["tools"] = tools
        cliente = genai.GenerativeModel(modelo)
        return cliente.generate_content(prompt, **kwargs)

    ultimo_error = None
    for intento in range(max_reintentos):
        try:
            try:
                respuesta = _ejecutar(tools_config)
            except Exception as primary_err:
                # Si el problema viene de la herramienta de búsqueda, hacemos
                # un retry sin herramientas (sigue siendo útil para el
                # usuario, sólo perdemos la posibilidad de citar noticias).
                msg = str(primary_err).lower()
                es_error_de_tool = (
                    "google_search" in msg
                    or "tool" in msg
                    or "grounding" in msg
                )
                if permitir_busqueda and es_error_de_tool:
                    respuesta = _ejecutar(None)
                else:
                    raise
            texto = _extraer_texto_gemini(respuesta)
            citas = _extraer_citas_gemini(respuesta) if permitir_busqueda else []
            meta = getattr(respuesta, "usage_metadata", None)
            usage = {
                "input_tokens": int(getattr(meta, "prompt_token_count", 0) or 0),
                "output_tokens": int(getattr(meta, "candidates_token_count", 0) or 0),
            }
            return texto, citas, usage
        except errores_reintentables as e:
            ultimo_error = e
            if intento < max_reintentos - 1:
                espera = 2 ** intento
                time.sleep(espera)
                continue
            break

    if isinstance(ultimo_error, google_exceptions.ResourceExhausted):
        raise RuntimeError(
            "Se ha superado el límite de peticiones a la IA. "
            "Por favor, espera un momento antes de volver a intentarlo."
        )
    if isinstance(ultimo_error, google_exceptions.ServiceUnavailable):
        raise RuntimeError(
            "El servicio de IA está sobrecargado en este momento. "
            "Por favor, inténtalo de nuevo en unos segundos."
        )
    raise RuntimeError(
        f"Error al comunicar con el servicio de IA: {str(ultimo_error)}"
    )


def llamar_ia_con_reintentos(prompt, max_tokens, modelo, permitir_busqueda=False, modulo=None):
    """Despacha la llamada al proveedor de IA y devuelve (texto, citas).

    `modulo` (str opcional) etiqueta la llamada en el registro de coste
    (`auditoria/coste_tokens.txt`). Si se omite, el coste se registra pero
    sin discriminar el origen.
    """
    modelo = _normalizar_modelo(modelo)
    if modelo in MODELOS_ANTHROPIC:
        texto, citas, usage = llamar_claude_con_reintentos(
            prompt, max_tokens, modelo, permitir_busqueda=permitir_busqueda,
        )
    else:
        texto, citas, usage = llamar_gemini_con_reintentos(
            prompt, max_tokens, modelo, permitir_busqueda=permitir_busqueda,
        )
    try:
        _registrar_coste(modelo, modulo or "(sin-modulo)", usage)
    except Exception as _e:
        _logging.getLogger("dfin.coste").warning("No se pudo registrar coste: %s", _e)
    return texto, citas


def _append_fuentes_consultadas(texto_md, citas):
    """No-op: la sección 'Fuentes consultadas' se ha retirado del informe."""
    return texto_md


def _limpiar_respuesta_ia(texto):
    """Saneamiento post-hoc de la respuesta del modelo.

    Elimina bloques de código y convierte cualquier tabla Markdown que el
    modelo haya emitido (pese a las instrucciones) en una lista con
    etiquetas claras, preservando los datos sin los caracteres `|` y `---`.
    """
    if not texto:
        return texto

    # Si el modelo introdujo una sección final de fuentes / enlaces pese a
    # las instrucciones, la eliminamos antes de seguir saneando.
    texto = re.sub(
        r"\n#{1,6}\s*(Fuentes consultadas|Fuentes|Referencias|Enlaces consultados|Bibliograf[íi]a)[^\n]*\n.*\Z",
        "",
        texto,
        flags=re.IGNORECASE | re.DOTALL,
    )

    texto = re.sub(r"```[^\n`]*\n?", "", texto)
    texto = texto.replace("```", "")

    lineas = texto.split("\n")
    resultado = []
    encabezados_tabla = None

    for linea in lineas:
        stripped = linea.strip()

        if _RE_TABLA_SEP_MD.match(stripped):
            # Fila separadora: descarta y marca que lo siguiente es cuerpo
            continue

        if _RE_TABLA_FILA_MD.match(stripped) and stripped.count("|") >= 2:
            celdas = [c.strip() for c in stripped.strip("|").split("|")]
            celdas = [c for c in celdas if c]
            if not celdas:
                continue
            if encabezados_tabla is None:
                encabezados_tabla = celdas
                continue
            if len(celdas) == len(encabezados_tabla):
                partes = [
                    f"**{h}:** {v}" for h, v in zip(encabezados_tabla, celdas)
                ]
                resultado.append("- " + " — ".join(partes))
            else:
                resultado.append("- " + " · ".join(celdas))
            continue

        # Línea fuera de tabla: reinicia el tracking de encabezados
        encabezados_tabla = None
        resultado.append(linea)

    # Colapsa múltiples líneas en blanco a una sola
    limpio = "\n".join(resultado)
    limpio = re.sub(r"\n{3,}", "\n\n", limpio).strip()
    return limpio


def obtener_datos_financieros(ticker_symbol):
    """Obtiene datos financieros de mercado para un ticker dado."""
    ticker = yf.Ticker(ticker_symbol)

    # Inicializar sesión con history antes de obtener fundamentales
    ticker.history(period="1d")

    income_stmt = ticker.income_stmt
    if income_stmt is None or income_stmt.empty:
        return None

    nombre = ticker.info.get("shortName", ticker_symbol.upper())
    moneda = (
        ticker.info.get("financialCurrency")
        or ticker.info.get("currency")
        or ""
    )
    moneda = str(moneda).upper().strip() if moneda else ""

    datos = {
        "ticker": ticker_symbol.upper(),
        "nombre": nombre,
        "moneda": moneda,
        "periodos": [],
    }

    columnas = income_stmt.columns[:4]  # Últimos 4 periodos disponibles

    for col in columnas:
        periodo = str(col.date()) if hasattr(col, "date") else str(col)

        def safe_get(df, keys):
            for key in keys:
                if key in df.index:
                    val = df.loc[key, col]
                    if val is not None and str(val) != "nan":
                        return float(val)
            return 0.0

        ingresos = safe_get(income_stmt, ["Total Revenue", "Revenue"])
        coste_ventas = safe_get(
            income_stmt,
            ["Cost Of Revenue", "Cost of Revenue", "Cost Of Goods Sold"],
        )
        margen_bruto = ingresos - coste_ventas
        gastos_operativos = safe_get(
            income_stmt,
            ["Operating Expense", "Total Operating Expenses", "Selling General And Administration"],
        )
        ebitda = safe_get(income_stmt, ["EBITDA", "Normalized EBITDA"])
        beneficio_neto = safe_get(
            income_stmt,
            ["Net Income", "Net Income Common Stockholders"],
        )

        # PyG ampliada
        ebit = safe_get(
            income_stmt,
            ["EBIT", "Operating Income", "Total Operating Income As Reported"],
        )
        da = safe_get(
            income_stmt,
            [
                "Reconciled Depreciation",
                "Depreciation And Amortization",
                "Depreciation Amortization Depletion Income Statement",
            ],
        )
        sga = safe_get(
            income_stmt,
            ["Selling General And Administration", "Selling General Administrative"],
        )
        rd = safe_get(income_stmt, ["Research And Development"])
        ingresos_financieros = safe_get(income_stmt, ["Interest Income"])
        gastos_financieros = safe_get(
            income_stmt, ["Interest Expense", "Interest Expense Non Operating"],
        )
        resultado_financiero_neto = safe_get(
            income_stmt,
            ["Net Non Operating Interest Income Expense"],
        )
        if not resultado_financiero_neto:
            resultado_financiero_neto = ingresos_financieros - gastos_financieros
        resultado_antes_impuestos = safe_get(
            income_stmt, ["Pretax Income", "Income Before Tax"],
        )
        impuestos = safe_get(
            income_stmt, ["Tax Provision", "Income Tax Expense Benefit"],
        )
        bpa_basico = safe_get(income_stmt, ["Basic EPS"])
        bpa_diluido = safe_get(income_stmt, ["Diluted EPS"])

        pct = lambda parte: round((parte / ingresos) * 100, 2) if ingresos else 0.0
        tasa_impositiva = (
            round((impuestos / resultado_antes_impuestos) * 100, 2)
            if resultado_antes_impuestos else 0.0
        )

        datos["periodos"].append(
            {
                "periodo": periodo,
                "ingresos": ingresos,
                "coste_ventas": coste_ventas,
                "margen_bruto": margen_bruto,
                "gastos_operativos": gastos_operativos,
                "ebitda": ebitda,
                "beneficio_neto": beneficio_neto,
                "pct_margen_bruto": pct(margen_bruto),
                "pct_gastos_operativos": pct(gastos_operativos),
                "pct_ebitda": pct(ebitda),
                "pct_beneficio_neto": pct(beneficio_neto),
                # Ampliación PyG
                "ebit": ebit,
                "pct_ebit": pct(ebit),
                "da": da,
                "pct_da": pct(da),
                "sga": sga,
                "pct_sga": pct(sga),
                "rd": rd,
                "pct_rd": pct(rd),
                "ingresos_financieros": ingresos_financieros,
                "gastos_financieros": gastos_financieros,
                "resultado_financiero_neto": resultado_financiero_neto,
                "resultado_antes_impuestos": resultado_antes_impuestos,
                "pct_rai": pct(resultado_antes_impuestos),
                "impuestos": impuestos,
                "tasa_impositiva": tasa_impositiva,
                "bpa_basico": bpa_basico,
                "bpa_diluido": bpa_diluido,
            }
        )

    datos["estados"] = _serializar_estados_yahoo(ticker)
    return datos


# Mapeo de orden + traducción a español por estado, alineado con la presentación
# de Yahoo Finance (es.finance.yahoo.com). Cada lista es el orden de aparición
# canónico; las filas que yfinance devuelva fuera de este mapeo se añaden al
# final manteniendo su nombre original.
_ORDEN_INCOME_STATEMENT = [
    ("Total Revenue", "Ingresos totales"),
    ("Operating Revenue", "Ingresos operativos"),
    ("Cost Of Revenue", "Coste de los ingresos"),
    ("Gross Profit", "Beneficio bruto"),
    ("Operating Expense", "Gastos operativos"),
    ("Research And Development", "Investigación y desarrollo"),
    ("Selling General And Administration", "Gastos generales y de administración"),
    ("Selling And Marketing Expense", "Gastos de ventas y marketing"),
    ("General And Administrative Expense", "Gastos generales y administrativos"),
    ("Other Operating Expenses", "Otros gastos operativos"),
    ("Operating Income", "Beneficio operativo"),
    ("Total Operating Income As Reported", "Beneficio operativo (declarado)"),
    ("Net Non Operating Interest Income Expense", "Resultado financiero neto"),
    ("Interest Income", "Ingresos por intereses"),
    ("Interest Expense", "Gastos por intereses"),
    ("Interest Income Non Operating", "Ingresos por intereses no operativos"),
    ("Interest Expense Non Operating", "Gastos por intereses no operativos"),
    ("Other Income Expense", "Otros ingresos / gastos"),
    ("Other Non Operating Income Expenses", "Otros ingresos / gastos no operativos"),
    ("Special Income Charges", "Cargos / ingresos especiales"),
    ("Gain On Sale Of Security", "Beneficio por venta de inversiones"),
    ("Pretax Income", "Beneficio antes de impuestos"),
    ("Tax Provision", "Provisión para impuestos sobre la renta"),
    ("Tax Effect Of Unusual Items", "Efecto fiscal de partidas inusuales"),
    ("Tax Rate For Calcs", "Tasa impositiva (para cálculos)"),
    ("Net Income Common Stockholders", "Beneficio neto para los accionistas comunes"),
    ("Net Income", "Beneficio neto"),
    ("Net Income Continuous Operations", "Beneficio neto de operaciones continuadas"),
    ("Net Income Discontinuous Operations", "Beneficio neto de operaciones discontinuadas"),
    ("Net Income From Continuing And Discontinued Operation", "Beneficio neto de operaciones continuadas y discontinuadas"),
    ("Net Income From Continuing Operation Net Minority Interest", "Beneficio neto de operación continua (interés minoritario neto)"),
    ("Net Income Including Noncontrolling Interests", "Beneficio neto incluyendo intereses minoritarios"),
    ("Minority Interests", "Intereses minoritarios"),
    ("Diluted NI Available To Com Stockholders", "Beneficio neto diluido para accionistas comunes"),
    ("Basic EPS", "BPA básico"),
    ("Diluted EPS", "BPA diluido"),
    ("Basic Average Shares", "Acciones medias básicas"),
    ("Diluted Average Shares", "Acciones medias diluidas"),
    ("Total Expenses", "Total gastos"),
    ("Normalized Income", "Beneficio normalizado"),
    ("Reconciled Cost Of Revenue", "Coste de los ingresos reconciliado"),
    ("Reconciled Depreciation", "Amortización reconciliada"),
    ("Total Unusual Items Excluding Goodwill", "Partidas inusuales (excl. fondo de comercio)"),
    ("Total Unusual Items", "Total partidas inusuales"),
    ("EBIT", "EBIT"),
    ("EBITDA", "EBITDA"),
    ("Normalized EBITDA", "EBITDA normalizado"),
]

_ORDEN_BALANCE_SHEET = [
    ("Total Assets", "Activos totales"),
    ("Current Assets", "Activos corrientes"),
    ("Cash Cash Equivalents And Short Term Investments", "Efectivo, equivalentes e inversiones a corto plazo"),
    ("Cash And Cash Equivalents", "Efectivo y equivalentes de efectivo"),
    ("Cash Financial", "Efectivo (financiero)"),
    ("Cash Equivalents", "Equivalentes de efectivo"),
    ("Other Short Term Investments", "Otras inversiones a corto plazo"),
    ("Receivables", "Cuentas por cobrar"),
    ("Accounts Receivable", "Cuentas por cobrar comerciales"),
    ("Other Receivables", "Otras cuentas por cobrar"),
    ("Inventory", "Existencias"),
    ("Prepaid Assets", "Pagos anticipados"),
    ("Other Current Assets", "Otros activos corrientes"),
    ("Total Non Current Assets", "Activos no corrientes"),
    ("Net PPE", "Inmovilizado material neto"),
    ("Gross PPE", "Inmovilizado material bruto"),
    ("Accumulated Depreciation", "Amortización acumulada"),
    ("Goodwill And Other Intangible Assets", "Fondo de comercio y otros intangibles"),
    ("Goodwill", "Fondo de comercio"),
    ("Other Intangible Assets", "Otros activos intangibles"),
    ("Investments And Advances", "Inversiones y anticipos"),
    ("Long Term Equity Investment", "Inversiones a largo plazo en patrimonio"),
    ("Other Non Current Assets", "Otros activos no corrientes"),
    ("Total Liabilities Net Minority Interest", "Pasivos totales (neto de intereses minoritarios)"),
    ("Current Liabilities", "Pasivos corrientes"),
    ("Payables And Accrued Expenses", "Cuentas por pagar y gastos devengados"),
    ("Payables", "Cuentas por pagar"),
    ("Accounts Payable", "Cuentas por pagar comerciales"),
    ("Current Debt And Capital Lease Obligation", "Deuda a corto plazo y arrendamientos"),
    ("Current Debt", "Deuda a corto plazo"),
    ("Current Capital Lease Obligation", "Arrendamientos a corto plazo"),
    ("Current Deferred Revenue", "Ingresos diferidos a corto plazo"),
    ("Other Current Liabilities", "Otros pasivos corrientes"),
    ("Total Non Current Liabilities Net Minority Interest", "Pasivos no corrientes (neto de intereses minoritarios)"),
    ("Long Term Debt And Capital Lease Obligation", "Deuda a largo plazo y arrendamientos"),
    ("Long Term Debt", "Deuda a largo plazo"),
    ("Long Term Capital Lease Obligation", "Arrendamientos a largo plazo"),
    ("Non Current Deferred Revenue", "Ingresos diferidos a largo plazo"),
    ("Non Current Deferred Taxes Liabilities", "Pasivos por impuestos diferidos a largo plazo"),
    ("Tradeand Other Payables Non Current", "Otras cuentas por pagar no corrientes"),
    ("Other Non Current Liabilities", "Otros pasivos no corrientes"),
    ("Total Equity Gross Minority Interest", "Patrimonio neto total"),
    ("Stockholders Equity", "Fondos propios"),
    ("Capital Stock", "Capital social"),
    ("Common Stock", "Acciones ordinarias"),
    ("Additional Paid In Capital", "Prima de emisión"),
    ("Retained Earnings", "Reservas / Beneficios acumulados"),
    ("Treasury Stock", "Acciones propias en cartera"),
    ("Gains Losses Not Affecting Retained Earnings", "Ganancias / pérdidas no afectan reservas"),
    ("Minority Interest", "Intereses minoritarios"),
    ("Total Capitalization", "Capitalización total"),
    ("Common Stock Equity", "Patrimonio en acciones ordinarias"),
    ("Net Tangible Assets", "Activos netos tangibles"),
    ("Working Capital", "Capital circulante"),
    ("Invested Capital", "Capital invertido"),
    ("Tangible Book Value", "Valor en libros tangible"),
    ("Total Debt", "Deuda total"),
    ("Net Debt", "Deuda neta"),
    ("Share Issued", "Acciones emitidas"),
    ("Ordinary Shares Number", "Número de acciones ordinarias"),
]

_ORDEN_CASHFLOW = [
    ("Operating Cash Flow", "Flujo de caja operativo"),
    ("Cash Flow From Continuing Operating Activities", "Flujo de caja de actividades operativas continuadas"),
    ("Net Income From Continuing Operations", "Beneficio neto de operaciones continuadas"),
    ("Depreciation Amortization Depletion", "Amortizaciones (D&A)"),
    ("Depreciation And Amortization", "Amortizaciones"),
    ("Deferred Income Tax", "Impuesto sobre la renta diferido"),
    ("Stock Based Compensation", "Retribución basada en acciones"),
    ("Change In Working Capital", "Variación del capital circulante"),
    ("Change In Receivables", "Variación de cuentas por cobrar"),
    ("Change In Inventory", "Variación de existencias"),
    ("Change In Payables And Accrued Expense", "Variación de cuentas por pagar y devengos"),
    ("Change In Other Current Assets", "Variación de otros activos corrientes"),
    ("Change In Other Current Liabilities", "Variación de otros pasivos corrientes"),
    ("Change In Other Working Capital", "Variación de otro capital circulante"),
    ("Other Non Cash Items", "Otras partidas no monetarias"),
    ("Investing Cash Flow", "Flujo de caja de inversión"),
    ("Cash Flow From Continuing Investing Activities", "Flujo de caja de actividades de inversión continuadas"),
    ("Capital Expenditure", "Inversiones en inmovilizado (CAPEX)"),
    ("Net PPE Purchase And Sale", "Compra / venta neta de inmovilizado"),
    ("Purchase Of PPE", "Compras de inmovilizado"),
    ("Sale Of PPE", "Ventas de inmovilizado"),
    ("Net Investment Purchase And Sale", "Compra / venta neta de inversiones"),
    ("Purchase Of Investment", "Compras de inversiones"),
    ("Sale Of Investment", "Ventas de inversiones"),
    ("Net Business Purchase And Sale", "Compra / venta neta de negocios"),
    ("Purchase Of Business", "Compras de negocios"),
    ("Sale Of Business", "Ventas de negocios"),
    ("Net Other Investing Changes", "Otras variaciones netas de inversión"),
    ("Financing Cash Flow", "Flujo de caja de financiación"),
    ("Cash Flow From Continuing Financing Activities", "Flujo de caja de actividades de financiación continuadas"),
    ("Net Issuance Payments Of Debt", "Emisiones / amortizaciones netas de deuda"),
    ("Net Long Term Debt Issuance", "Emisión neta de deuda a largo plazo"),
    ("Long Term Debt Issuance", "Emisión de deuda a largo plazo"),
    ("Long Term Debt Payments", "Amortización de deuda a largo plazo"),
    ("Net Short Term Debt Issuance", "Emisión neta de deuda a corto plazo"),
    ("Net Common Stock Issuance", "Emisión neta de acciones ordinarias"),
    ("Common Stock Issuance", "Emisión de acciones ordinarias"),
    ("Common Stock Payments", "Recompra de acciones ordinarias"),
    ("Cash Dividends Paid", "Dividendos pagados en efectivo"),
    ("Common Stock Dividend Paid", "Dividendos ordinarios pagados"),
    ("Net Other Financing Charges", "Otros cargos netos de financiación"),
    ("End Cash Position", "Saldo final de efectivo"),
    ("Beginning Cash Position", "Saldo inicial de efectivo"),
    ("Changes In Cash", "Variación neta de efectivo"),
    ("Effect Of Exchange Rate Changes", "Efecto del tipo de cambio"),
    ("Income Tax Paid Supplemental Data", "Impuesto sobre la renta pagado (info. complementaria)"),
    ("Interest Paid Supplemental Data", "Intereses pagados (info. complementaria)"),
    ("Free Cash Flow", "Flujo de caja libre"),
    ("Repurchase Of Capital Stock", "Recompra de capital social"),
    ("Issuance Of Capital Stock", "Emisión de capital social"),
    ("Issuance Of Debt", "Emisión de deuda"),
    ("Repayment Of Debt", "Amortización de deuda"),
]


def _ordenar_y_traducir_filas(df_index_to_valores, orden):
    """Devuelve la lista de filas {concepto, valores} ordenadas según `orden`.

    `df_index_to_valores` es un dict concepto_en_inglés → lista de valores.
    `orden` es la lista de tuplas (clave_yfinance, traducción_es).
    Las filas no presentes en `orden` se añaden al final con su nombre original.
    """
    en_orden = set()
    filas = []
    for clave, traduccion in orden:
        if clave in df_index_to_valores:
            filas.append({"concepto": traduccion, "valores": df_index_to_valores[clave]})
            en_orden.add(clave)
    for clave, valores in df_index_to_valores.items():
        if clave not in en_orden:
            filas.append({"concepto": clave, "valores": valores})
    return filas


def _serializar_estados_yahoo(ticker):
    """Devuelve los 3 estados financieros, traducidos al español y reordenados.

    Estructura: {nombre_estado: {"periodos": [str], "filas": [{"concepto", "valores": [float|None]}]}}.
    """
    import math
    fuentes = {
        "income_statement": (getattr(ticker, "income_stmt", None), _ORDEN_INCOME_STATEMENT),
        "balance_sheet": (getattr(ticker, "balance_sheet", None), _ORDEN_BALANCE_SHEET),
        "cashflow": (getattr(ticker, "cashflow", None), _ORDEN_CASHFLOW),
    }
    salida = {}
    for clave, (df, orden) in fuentes.items():
        if df is None or df.empty:
            salida[clave] = {"periodos": [], "filas": []}
            continue
        cols = list(df.columns[:4])
        periodos = [str(c.date()) if hasattr(c, "date") else str(c) for c in cols]
        crudo = {}
        for concepto in df.index:
            valores = []
            for col in cols:
                v = df.loc[concepto, col]
                if v is None:
                    valores.append(None)
                    continue
                try:
                    vf = float(v)
                    valores.append(None if math.isnan(vf) else vf)
                except (TypeError, ValueError):
                    valores.append(None)
            if any(v is not None for v in valores):
                crudo[str(concepto)] = valores
        filas = _ordenar_y_traducir_filas(crudo, orden)
        salida[clave] = {"periodos": periodos, "filas": filas}
    return salida


def _texto_datos_para_prompt(datos):
    """Genera el bloque de datos financieros para inyectar en prompts."""
    moneda = datos.get("moneda") or ""
    etiqueta_moneda = f" (moneda de reporting: {moneda})" if moneda else ""
    texto = f"Empresa: {datos['nombre']} ({datos['ticker']}){etiqueta_moneda}\n\n"
    for p in datos["periodos"]:
        texto += (
            f"Periodo: {p['periodo']}\n"
            f"  Ingresos: {p['ingresos']:,.0f}\n"
            f"  Coste de ventas: {p['coste_ventas']:,.0f}\n"
            f"  Margen bruto: {p['margen_bruto']:,.0f} ({p['pct_margen_bruto']}%)\n"
            f"  Gastos operativos: {p['gastos_operativos']:,.0f} ({p['pct_gastos_operativos']}%)\n"
            f"    - SG&A: {p.get('sga', 0):,.0f} ({p.get('pct_sga', 0)}%)\n"
            f"    - I+D: {p.get('rd', 0):,.0f} ({p.get('pct_rd', 0)}%)\n"
            f"  EBITDA: {p['ebitda']:,.0f} ({p['pct_ebitda']}%)\n"
            f"  D&A (amortizaciones): {p.get('da', 0):,.0f} ({p.get('pct_da', 0)}%)\n"
            f"  EBIT (resultado de explotación): {p.get('ebit', 0):,.0f} ({p.get('pct_ebit', 0)}%)\n"
            f"  Resultado financiero neto: {p.get('resultado_financiero_neto', 0):,.0f}"
            f" (ingresos fin.: {p.get('ingresos_financieros', 0):,.0f}, gastos fin.: {p.get('gastos_financieros', 0):,.0f})\n"
            f"  Resultado antes de impuestos: {p.get('resultado_antes_impuestos', 0):,.0f} ({p.get('pct_rai', 0)}%)\n"
            f"  Impuesto sobre beneficios: {p.get('impuestos', 0):,.0f}"
            f" (tasa efectiva: {p.get('tasa_impositiva', 0)}%)\n"
            f"  Beneficio neto: {p['beneficio_neto']:,.0f} ({p['pct_beneficio_neto']}%)\n"
            f"  BPA básico: {p.get('bpa_basico', 0)} | BPA diluido: {p.get('bpa_diluido', 0)}\n"
        )

    estados = datos.get("estados") or {}
    nombres = {
        "income_statement": "Estado de Ingresos (income statement)",
        "balance_sheet": "Balance General (balance sheet)",
        "cashflow": "Flujo de Caja (cashflow)",
    }
    for clave, titulo in nombres.items():
        bloque = estados.get(clave) or {}
        filas = bloque.get("filas") or []
        periodos = bloque.get("periodos") or []
        if not filas or not periodos:
            continue
        texto += f"\n--- {titulo} (datos brutos de Yahoo Finance) ---\n"
        texto += "Concepto | " + " | ".join(periodos) + "\n"
        for fila in filas:
            valores_txt = []
            for v in fila["valores"]:
                if v is None:
                    valores_txt.append("n/d")
                elif abs(v) >= 1:
                    valores_txt.append(f"{v:,.0f}")
                else:
                    valores_txt.append(f"{v:.4f}")
            texto += f"{fila['concepto']} | " + " | ".join(valores_txt) + "\n"
    return texto


# Bloque compartido con el ALCANCE del documento. Fija la perspectiva
# (memoria explicativa de la PyG para accionistas) y veta contenido de tesis
# de inversión.
_ALCANCE_MEMORIA = """ALCANCE Y PÚBLICO DEL DOCUMENTO (fundamental, no te desvíes):
- Eres un analista financiero que redacta la NOTA EXPLICATIVA de la cuenta
  de Pérdidas y Ganancias (PyG) que forma parte de las cuentas anuales que
  la compañía publica para sus accionistas.
- El lector es un accionista que quiere entender cómo ha ido la compañía
  en los ejercicios presentados y qué cabe esperar cualitativamente en los
  próximos ejercicios (drivers de negocio, tendencias, riesgos operativos).
- ESTO NO ES UNA TESIS DE INVERSIÓN. Está PROHIBIDO incluir:
  · Recomendaciones de compra, venta o mantenimiento.
  · Objetivos de precio o rangos de valoración.
  · Múltiplos de mercado (P/E, EV/EBITDA…), DCF, comparables cotizados.
  · Análisis técnico, rating, price target ni opinión bursátil.
- SÍ debes incluir: evolución interanual de ingresos, márgenes y resultado;
  drivers cualitativos que explican la variación; perspectivas cualitativas
  sobre el comportamiento del negocio; riesgos y fortalezas operativas."""


# Reglas de formato compartidas entre generación inicial y chat.
_REGLAS_FORMATO_MEMORIA = """Reglas de formato OBLIGATORIAS (bajo ningún concepto las infrinjas):
- PROHIBIDO usar tablas Markdown: no escribas NUNCA los caracteres `|`, `---` o `:---:`.
- En lugar de tablas, usa listas con guion `-` y negritas para las etiquetas, por ejemplo:
  `- **FY2024:** ingresos de 60.922 M USD (+125,8% interanual).`
- NO uses bloques de código ni acentos graves triples.
- PROHIBIDO incluir gráficos, charts, dashboards, diagramas, esquemas visuales,
  ASCII art o cualquier representación visual de datos. Tampoco uses frases del
  tipo "ver gráfico", "como muestra el dashboard", "véase la figura", "según el
  panel". La nota es ÚNICAMENTE texto en prosa profesional: todo análisis
  cuantitativo debe expresarse con cifras dentro del párrafo o en listas con guion.
- Estructura el contenido con encabezados Markdown `## Sección` y `### Subsección`.
- Usa párrafos en prosa profesional; resalta términos clave con `**negrita**` (uno o dos por párrafo).
- Cifras siempre con separadores de miles y unidades (millones / %).
- Asegúrate de cerrar todas las secciones de forma natural. Si NO hay un enfoque
  estricto definido por el usuario, cierra con una sección "## Conclusión". Si
  hay un enfoque estricto, NO añadas conclusión genérica ni secciones extra:
  el cierre debe ser una o dos frases finales sobre el enfoque pedido y nada más."""


_TIPOS_INFORME = {
    "nota_memoria": (
        "NOTA DE MEMORIA contable: documento formal que acompaña a las cuentas "
        "anuales para los accionistas, lenguaje técnico-contable pero accesible, "
        "evita opiniones de inversión."
    ),
    "consejo": (
        "INFORME AL CONSEJO de Administración: tono ejecutivo, foco en decisiones "
        "y palancas de gestión, destaca riesgos y oportunidades; estructura tipo "
        "executive summary + análisis + recomendaciones implícitas (no inversión)."
    ),
    "interno": (
        "INFORME INTERNO de management: orientado al equipo financiero, tono "
        "operativo y directo, sin perífrasis; puede entrar en detalle de variaciones "
        "y métricas operativas."
    ),
    "accionista": (
        "EXPLICACIÓN AL ACCIONISTA minorista: lenguaje claro y didáctico, evita "
        "jerga contable salvo definiéndola, foco en ‘qué significa esto para mí’ "
        "sin convertirse en recomendación de inversión."
    ),
}

_EXTENSIONES = {
    "breve": ("aproximadamente 400-500 palabras", 1800),
    "estandar": ("aproximadamente 800-1000 palabras", 4000),
    "extensa": ("aproximadamente 1500-2000 palabras", 6500),
}


def _bloque_configuracion_informe(tipo_informe, enfoque, extension):
    partes = []
    if tipo_informe and tipo_informe in _TIPOS_INFORME:
        partes.append(f"TIPO DE INFORME REQUERIDO: {_TIPOS_INFORME[tipo_informe]}")
    if enfoque and enfoque.strip():
        partes.append(
            "ENFOQUE ESTRICTO DEL INFORME (CRÍTICO, no lo amplíes):\n"
            f"{enfoque.strip()}\n\n"
            "REGLAS DE ALCANCE — son OBLIGATORIAS y prevalecen sobre cualquier\n"
            "convención de estructura habitual de notas de memoria o informes:\n"
            "- El documento debe limitarse EXCLUSIVAMENTE al enfoque indicado arriba.\n"
            "- NO añadas secciones como 'Perspectivas cualitativas', 'Riesgos',\n"
            "  'Resumen ejecutivo', 'Evolución general del negocio', 'Contexto sectorial',\n"
            "  'Conclusión', 'Recomendaciones' u otras secciones genéricas, salvo que\n"
            "  el propio enfoque las pida.\n"
            "- NO traigas magnitudes financieras ajenas al enfoque (p. ej. si el enfoque\n"
            "  es BPA, no hables de margen bruto, deuda, caja ni de operaciones que no\n"
            "  afecten directamente al BPA), salvo cuando sean estrictamente necesarias\n"
            "  para explicar el enfoque, y siempre en una sola frase de contexto.\n"
            "- El cierre del documento es una o dos frases finales sobre el enfoque\n"
            "  pedido; NO abras nuevos temas en el cierre.\n"
            "- Si la información disponible sobre el enfoque es escasa, redacta un\n"
            "  documento más breve antes que rellenarlo con apartados fuera de alcance."
        )
    if extension and extension in _EXTENSIONES:
        descripcion, _ = _EXTENSIONES[extension]
        partes.append(
            f"EXTENSIÓN OBJETIVO: {descripcion}. Si el enfoque no da para esa "
            "extensión, prioriza la regla de alcance sobre la de extensión y "
            "entrega un documento más corto."
        )
    if not partes:
        return ""
    return "\n\n" + "\n\n".join(partes) + "\n"


def generar_nota_memoria(datos, modelo, instrucciones_usuario=None,
                          permitir_busqueda=False,
                          tipo_informe=None, enfoque=None, extension=None):
    """Genera el informe narrativo a partir de los estados financieros.

    - `instrucciones_usuario`: texto libre que el usuario añade para orientar
      el tono, las secciones o el foco del informe.
    - `permitir_busqueda`: si es True, la IA puede consultar noticias recientes
      del sector en internet y citarlas en una sección final.
    - `tipo_informe`: clave de `_TIPOS_INFORME` que ajusta el registro y la
      audiencia objetivo del documento.
    - `enfoque`: texto libre que describe en qué parte de los estados se debe
      concentrar el análisis (puede mezclar selecciones predefinidas y texto
      libre del usuario).
    - `extension`: clave de `_EXTENSIONES` (breve / estandar / extensa).
    """
    tabla_texto = _texto_datos_para_prompt(datos)

    bloque_config = _bloque_configuracion_informe(tipo_informe, enfoque, extension)

    bloque_instrucciones = ""
    if instrucciones_usuario and instrucciones_usuario.strip():
        bloque_instrucciones = (
            "\n\nIndicaciones específicas del usuario (tienen PRIORIDAD sobre las "
            "decisiones por defecto, siempre que no contradigan el alcance ni las "
            "reglas de formato):\n"
            f"{instrucciones_usuario.strip()}\n"
        )

    bloque_busqueda = ""
    if permitir_busqueda:
        bloque_busqueda = (
            "\nCONTEXTO EXTERNO: tienes disponible una herramienta de búsqueda web. "
            "Úsala con criterio para localizar noticias relevantes de los últimos 12 "
            "meses sobre la compañía o su sector que ayuden a explicar la evolución "
            "del negocio. Incorpóralas en el cuerpo del informe como **hechos** "
            "(no como opinión), y asegúrate de que el análisis financiero siga "
            "siendo el centro del documento.\n"
        )

    max_tokens = _EXTENSIONES.get(extension or "", (None, 4000))[1]

    prompt = f"""{_ALCANCE_MEMORIA}

{_REGLAS_FORMATO_MEMORIA}{bloque_config}{bloque_instrucciones}{bloque_busqueda}
Datos financieros sobre los que elaborar el informe:
{tabla_texto}"""

    log_inf = _logging.getLogger("dfin.informe")
    import time as _t
    _t0 = _t.time()
    log_inf.info(
        "Generando informe: ticker=%s tipo=%s extension=%s modelo=%s busqueda=%s",
        (datos or {}).get("ticker", "?"), tipo_informe, extension, modelo, permitir_busqueda,
    )
    try:
        texto, citas = llamar_ia_con_reintentos(
            prompt, max_tokens=max_tokens, modelo=modelo,
            permitir_busqueda=permitir_busqueda, modulo="informe.generar",
        )
    except Exception as e:
        log_inf.exception("Fallo generando informe: %s", e)
        raise
    memoria = _limpiar_respuesta_ia(texto)
    log_inf.info(
        "Informe generado: %d palabras, lat=%.2fs",
        len((memoria or "").split()), _t.time() - _t0,
    )
    return _append_fuentes_consultadas(memoria, citas)


_RE_BLOQUE_RESPUESTA = re.compile(r"\[RESPUESTA\]\s*\n(.+?)(?=\n\[MEMORIA\]|\Z)", re.DOTALL)
_RE_BLOQUE_MEMORIA = re.compile(r"\[MEMORIA\]\s*\n(.+)", re.DOTALL)


def _parsear_respuesta_chat(texto_crudo):
    """Separa la respuesta conversacional y la memoria actualizada."""
    if not texto_crudo:
        return "", None
    m_mem = _RE_BLOQUE_MEMORIA.search(texto_crudo)
    memoria_nueva = None
    if m_mem:
        memoria_nueva = _limpiar_respuesta_ia(m_mem.group(1).strip())
        texto_restante = texto_crudo[: m_mem.start()]
    else:
        texto_restante = texto_crudo

    m_resp = _RE_BLOQUE_RESPUESTA.search(texto_restante)
    if m_resp:
        respuesta = m_resp.group(1).strip()
    else:
        # Sin etiquetas: si no hay memoria nueva lo tratamos todo como chat
        respuesta = texto_restante.strip()
    # Truncamos la respuesta conversacional a algo razonable
    if len(respuesta) > 1200:
        respuesta = respuesta[:1200].rsplit(".", 1)[0] + "…"
    return respuesta, memoria_nueva


def refinar_memoria_con_chat(datos, memoria_actual, historial, mensaje_usuario,
                              instrucciones_iniciales, modelo,
                              permitir_busqueda=False):
    """Ejecuta un turno de chat para refinar la memoria.

    Devuelve `(respuesta_chat, memoria_nueva_o_None, citas)`.
    """
    tabla_texto = _texto_datos_para_prompt(datos)

    historial_txt = ""
    for turno in historial:
        rol = "Usuario" if turno.get("role") == "user" else "Asistente"
        historial_txt += f"\n{rol}: {turno.get('content', '').strip()}"

    instrucciones_bloque = ""
    if instrucciones_iniciales and instrucciones_iniciales.strip():
        instrucciones_bloque = (
            "\nInstrucciones iniciales del usuario (al generar la memoria por primera vez):\n"
            f"{instrucciones_iniciales.strip()}\n"
        )

    bloque_busqueda = ""
    if permitir_busqueda:
        bloque_busqueda = (
            "\nEn este turno el usuario te permite consultar noticias recientes en "
            "internet. Úsalo si aporta a la memoria (p. ej. noticias del sector de "
            "los últimos 12 meses relevantes para explicar la evolución).\n"
        )

    prompt = f"""{_ALCANCE_MEMORIA}

Colaboras con el usuario para refinar la NOTA EXPLICATIVA de la PyG ya
generada. El usuario te pide ajustes, ampliaciones o cambios de enfoque;
aplica esos cambios CONSERVANDO el alcance descrito arriba.

{_REGLAS_FORMATO_MEMORIA}
{bloque_busqueda}
Datos financieros de referencia:
{tabla_texto}
{instrucciones_bloque}
MEMORIA ACTUAL:
---
{memoria_actual.strip()}
---
{historial_txt}

Nueva petición del usuario:
{mensaje_usuario.strip()}

RESPONDE ESTRICTAMENTE con este formato exacto (sin texto adicional fuera de las etiquetas):

[RESPUESTA]
(1 o 2 frases en español explicando qué cambios has aplicado a la memoria.)

[MEMORIA]
(la MEMORIA COMPLETA ACTUALIZADA en markdown, respetando las reglas de formato.
Si el usuario no pide cambios en la memoria, repite la memoria actual tal cual.
NO añadas una sección de fuentes / enlaces al final: el informe debe terminar
en la conclusión narrativa, sin listado de URLs ni "Fuentes consultadas".)"""

    texto, citas = llamar_ia_con_reintentos(
        prompt, max_tokens=4500, modelo=modelo,
        permitir_busqueda=permitir_busqueda, modulo="informe.chat",
    )
    respuesta, memoria_nueva = _parsear_respuesta_chat(texto)
    if not respuesta:
        respuesta = "Memoria actualizada."
    return respuesta, memoria_nueva, citas


# Nombre del sistema externo simulado para la detección de incoherencias.
SISTEMA_NOTA_OFICIAL = "Workiva"


def _extraer_json_incoherencias(texto_crudo):
    """Localiza el primer array JSON en la respuesta y lo parsea.

    El LLM a veces envuelve el JSON en prosa o en un bloque ```json```;
    aceptamos cualquiera de esas variantes.
    """
    if not texto_crudo:
        return []
    texto = texto_crudo.strip()
    # Bloque ```json ... ```
    m = re.search(r"```(?:json)?\s*(\[.*?\])\s*```", texto, re.DOTALL)
    if m:
        candidato = m.group(1)
    else:
        # Primer array que aparezca en el texto
        m2 = re.search(r"\[\s*\{.*\}\s*\]", texto, re.DOTALL)
        candidato = m2.group(0) if m2 else texto
    try:
        data = json.loads(candidato)
    except (json.JSONDecodeError, TypeError):
        return []
    if not isinstance(data, list):
        return []
    incoherencias = []
    for item in data:
        if not isinstance(item, dict):
            continue
        incoherencias.append({
            "tipo": str(item.get("tipo", "")).strip() or "Discrepancia",
            "severidad": str(item.get("severidad", "")).strip().lower() or "media",
            "informe": str(item.get("informe", "")).strip() or "Informe corporativo",
            "extracto_ia": str(item.get("extracto_ia", "")).strip(),
            "extracto_oficial": str(item.get("extracto_oficial", "")).strip(),
            "explicacion": str(item.get("explicacion", "")).strip(),
        })
    return incoherencias


# Catálogo de informes narrativos simulados disponibles en el repositorio
# corporativo. El LLM debe atribuir cada incoherencia a uno de estos
# documentos para que la simulación sea coherente.
_INFORMES_WORKIVA = [
    "Informe de gestión",
    "Informe al consejo de administración",
    "Informe de evolución del negocio",
    "Informe trimestral de resultados",
    "Memoria abreviada de cuentas anuales",
    "Presentación a inversores",
    "Informe de auditoría interna",
    "Comentario de la dirección sobre los resultados (MD&A)",
]


def detectar_incoherencias_memoria(datos, memoria, modelo):
    """Simula la detección de incoherencias entre la nota IA y un conjunto
    de informes narrativos corporativos almacenados en Workiva.

    Los informes no existen: pedimos al LLM que invente entre 4 y 6
    discrepancias verosímiles, ancladas en los datos y el texto reales,
    y atribuya cada una a un informe distinto del catálogo simulado.
    """
    tabla_texto = _texto_datos_para_prompt(datos)
    memoria_trunc = (memoria or "").strip()
    if len(memoria_trunc) > 6000:
        memoria_trunc = memoria_trunc[:6000] + "\n[...]"

    catalogo_informes = "\n".join(f"  - {n}" for n in _INFORMES_WORKIVA)

    prompt = f"""Eres un revisor financiero senior. Estás comparando la NOTA EXPLICATIVA
de la cuenta de Pérdidas y Ganancias generada automáticamente por IA contra
un CONJUNTO DE INFORMES NARRATIVOS CORPORATIVOS almacenados en el repositorio
{SISTEMA_NOTA_OFICIAL}. Cada informe del repositorio aborda, total o
parcialmente, la misma información financiera que la nota IA, pero ha sido
redactado por equipos distintos (dirección financiera, auditoría, relaciones
con inversores, etc.), por lo que es habitual encontrar pequeñas
divergencias entre ellos.

CATÁLOGO DE INFORMES DEL REPOSITORIO (SIMULADO — escoge entre estos nombres
literales para el campo "informe", usa AL MENOS 3 informes distintos a lo
largo de las incoherencias que reportes):
{catalogo_informes}

Tu tarea: identificar entre 4 y 6 incoherencias VEROSÍMILES entre la nota
IA y los distintos informes del repositorio. Cada incoherencia debe estar
ANCLADA en los datos financieros y en el texto real de la versión IA (cita
extractos cortos exactos de la versión IA). La redacción del informe
corporativo es SIMULADA: invéntala con sentido financiero — cifras parecidas
pero ligeramente distintas, otro número de decimales, otro driver atribuido,
otro año destacado, otro signo de tendencia, otro reparto del margen, etc.
Asegúrate de que el extracto del informe es ESTILÍSTICAMENTE COHERENTE con
el tipo de informe del que procede (p. ej. el "Informe al consejo" usa tono
ejecutivo y agregado; la "Presentación a inversores" es más promocional;
el "Informe de auditoría interna" es más técnico y prudente).

Tipos de incoherencia que debes cubrir (al menos 3 tipos distintos entre las
4-6 que reportes):
- "Cifra": las dos versiones reportan un importe diferente para la misma
  métrica (p. ej. ingresos FY2024).
- "Decimales": misma magnitud con distinta precisión (1 vs 2 decimales, o
  redondeo distinto).
- "Driver": ambas atribuyen la variación a causas distintas (p. ej. la IA
  habla de tipo de cambio y el informe corporativo de volumen).
- "Periodo": una destaca otro ejercicio como pico o valle.
- "Signo/tendencia": una habla de mejora interanual y la otra de empeoramiento
  marginal.
- "Unidad/escala": miles vs millones, o moneda local vs reporting.

DEVUELVE EXCLUSIVAMENTE un array JSON válido (sin texto fuera del array, sin
bloques de código markdown). Cada elemento del array tiene exactamente estas
claves:
  - "tipo": una de las categorías de arriba (string corto).
  - "severidad": "alta" | "media" | "baja".
  - "informe": nombre LITERAL del informe del catálogo de arriba.
  - "extracto_ia": fragmento literal (máx 220 caracteres) de la versión IA.
  - "extracto_oficial": fragmento inventado (máx 220 caracteres) tal como
    aparecería en el informe corporativo indicado.
  - "explicacion": 1-2 frases explicando por qué difieren y qué implicación
    tiene para el lector.

Datos financieros (referencia común a todos los documentos):
{tabla_texto}

VERSIÓN IA (íntegra):
{memoria_trunc}
"""

    texto, _citas = llamar_ia_con_reintentos(
        prompt, max_tokens=2500, modelo=modelo, permitir_busqueda=False,
        modulo="informe.incoherencias",
    )
    return _extraer_json_incoherencias(texto)


def generar_analisis_comparativo(datos_empresas, modelo):
    """Genera un análisis comparativo de hasta 3 empresas con el modelo elegido."""
    texto = ""
    for datos in datos_empresas:
        texto += f"\n{'='*50}\nEmpresa: {datos['nombre']} ({datos['ticker']})\n"
        for p in datos["periodos"][:1]:  # Último periodo para comparar
            texto += f"""Periodo más reciente: {p['periodo']}
  Ingresos: {p['ingresos']:,.0f}
  Coste de ventas: {p['coste_ventas']:,.0f}
  Margen bruto: {p['margen_bruto']:,.0f} ({p['pct_margen_bruto']}%)
  Gastos operativos: {p['gastos_operativos']:,.0f} ({p['pct_gastos_operativos']}%)
  EBITDA: {p['ebitda']:,.0f} ({p['pct_ebitda']}%)
  Beneficio neto: {p['beneficio_neto']:,.0f} ({p['pct_beneficio_neto']}%)
"""

    prompt = f"""Eres un analista financiero experto. Genera un análisis comparativo detallado
en español de las siguientes empresas. Compara sus métricas financieras, identifica
cuál tiene mejor rendimiento en cada categoría, analiza las diferencias en márgenes
y rentabilidad, y ofrece una conclusión sobre cuál presenta mejor salud financiera.

Reglas de formato OBLIGATORIAS (bajo ningún concepto las infrinjas):
- PROHIBIDO usar tablas Markdown: no escribas NUNCA los caracteres `|`, `---` o `:---:`.
- En lugar de tablas, usa listas con guion `-` y negritas para las etiquetas, por ejemplo:
  `- **Margen bruto:** Empresa A 45,2% vs Empresa B 38,1%.`
- NO uses bloques de código ni acentos graves triples.
- Estructura el contenido con encabezados Markdown `## Sección` y `### Subsección`.
- Usa párrafos en prosa profesional; resalta términos clave con `**negrita**` (uno o dos por párrafo).
- Cifras siempre con separadores de miles y unidades (millones / %).
- Asegúrate de cerrar todas las secciones: termina con una sección final de "## Conclusión".

Datos financieros comparativos:
{texto}"""

    texto, _citas = llamar_ia_con_reintentos(
        prompt, max_tokens=5000, modelo=modelo, modulo="comparador",
    )
    return _limpiar_respuesta_ia(texto)


_RE_INLINE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")


def _añadir_runs_con_formato(paragraph, texto):
    """Añade runs al párrafo aplicando **negrita**, *cursiva* y `código`."""
    for parte in _RE_INLINE.split(texto):
        if not parte:
            continue
        if parte.startswith("**") and parte.endswith("**") and len(parte) > 4:
            run = paragraph.add_run(parte[2:-2])
            run.bold = True
        elif parte.startswith("*") and parte.endswith("*") and len(parte) > 2:
            run = paragraph.add_run(parte[1:-1])
            run.italic = True
        elif parte.startswith("`") and parte.endswith("`") and len(parte) > 2:
            run = paragraph.add_run(parte[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(parte)


# ---------------------------------------------------------------------------
# Segmentación de la nota y asignación de gráficos por sección.
# Los documentos Word/PDF (y la vista web) intercalan los gráficos después de
# la sección ## cuyos títulos contengan determinadas palabras clave.
# ---------------------------------------------------------------------------

# Paleta corporativa Deloitte (matplotlib)
_MPL_GREEN = "#86BC25"
_MPL_GREEN_DARK = "#43B02A"
_MPL_BLACK = "#000000"
_MPL_GRAY = "#53565A"
_MPL_ORANGE = "#ED8B00"

_CHART_KEYWORDS = {
    "ingresos": ("ingresos", "facturación", "revenue", "ventas", "volumen de negocio", "cifra de negocio"),
    "margenes": ("margen", "márgenes", "rentabilidad", "ebitda", "beneficio"),
    "estructura": ("coste", "costes", "estructura", "gasto"),
    "metricas": ("métrica", "ratios", "clave", "kpi", "resumen", "magnitudes", "indicadores"),
}

_RE_H2 = re.compile(r"^\s*##\s+(.*?)\s*$")


def _segmentar_memoria(texto):
    """Divide la nota de memoria en secciones por encabezados `## `.

    Devuelve una lista de dicts `{title, body}`. Si el texto contiene
    contenido previo al primer `##`, ese bloque se agrupa como preámbulo
    sin título.
    """
    if not texto:
        return []

    lineas = texto.replace("\r\n", "\n").split("\n")
    secciones = []
    preambulo = []
    actual_title = None
    actual_body = []

    def _cerrar():
        nonlocal actual_title, actual_body
        if actual_title is not None:
            secciones.append({
                "title": actual_title,
                "body": "\n".join(actual_body).strip(),
            })
        actual_title = None
        actual_body = []

    for linea in lineas:
        m = _RE_H2.match(linea)
        if m:
            _cerrar()
            actual_title = m.group(1).strip(" *")
            actual_body = []
        else:
            if actual_title is None:
                preambulo.append(linea)
            else:
                actual_body.append(linea)
    _cerrar()

    resultado = []
    preambulo_txt = "\n".join(preambulo).strip()
    if preambulo_txt:
        resultado.append({"title": None, "body": preambulo_txt})
    resultado.extend(secciones)
    return resultado


def _asignar_graficos_a_secciones(secciones, graficos_disponibles):
    """Asocia cada gráfico a la sección cuyo título encaja mejor por keywords.

    Se puntúa cada sección por número de coincidencias de palabras clave y
    se asigna el gráfico a la de mayor score. De este modo "Métricas clave"
    gana a "Resumen ejecutivo" para el gráfico de métricas porque acumula
    más palabras clave específicas.
    """
    asignaciones = {i: [] for i in range(len(secciones))}

    for key in graficos_disponibles:
        palabras = _CHART_KEYWORDS.get(key, ())
        mejor_idx = None
        mejor_score = 0
        for i, sec in enumerate(secciones):
            titulo = (sec.get("title") or "").lower()
            if not titulo:
                continue
            score = sum(1 for pal in palabras if pal in titulo)
            if score > mejor_score:
                mejor_score = score
                mejor_idx = i
        if mejor_idx is not None:
            asignaciones[mejor_idx].append(key)

    resultado = [
        {**sec, "charts": asignaciones[i]}
        for i, sec in enumerate(secciones)
    ]
    asignados_total = {ch for lst in asignaciones.values() for ch in lst}
    pendientes = [k for k in graficos_disponibles if k not in asignados_total]
    return resultado, pendientes


def segmentar_y_asignar(nota):
    """Wrapper pensado para usarse desde las vistas y los exportadores."""
    secciones = _segmentar_memoria(nota)
    orden = ["ingresos", "margenes", "estructura", "metricas"]
    return _asignar_graficos_a_secciones(secciones, orden)


# ---------------------------------------------------------------------------
# Renderizado de gráficos a PNG para Word/PDF (matplotlib, paleta Deloitte).
# ---------------------------------------------------------------------------

def _formato_eje_cifra(v, _pos=None):
    abs_v = abs(v)
    if abs_v >= 1e9:
        return f"{v/1e9:,.1f}B"
    if abs_v >= 1e6:
        return f"{v/1e6:,.1f}M"
    if abs_v >= 1e3:
        return f"{v/1e3:,.0f}K"
    return f"{v:,.0f}"


def _figura_a_bytes(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf


def _chart_png_ingresos(graficos):
    fig, ax = plt.subplots(figsize=(6.4, 3.6))
    ax.bar(graficos["labels"], graficos["ingresos"], color=_MPL_GREEN, edgecolor=_MPL_GREEN_DARK)
    moneda = graficos.get("moneda", "")
    titulo = "Evolución de ingresos" + (f" ({moneda})" if moneda else "")
    ax.set_title(titulo, fontsize=11, fontweight="bold", loc="left", color=_MPL_BLACK)
    ax.yaxis.set_major_formatter(FuncFormatter(_formato_eje_cifra))
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.grid(axis="y", color="#E4E4E2", linewidth=0.6)
    ax.set_axisbelow(True)
    return _figura_a_bytes(fig)


def _chart_png_margenes(graficos):
    fig, ax = plt.subplots(figsize=(6.4, 3.6))
    labels = graficos["labels"]
    ax.plot(labels, graficos["margenes"]["margen_bruto"], marker="o", color=_MPL_GREEN_DARK, label="Margen bruto (%)")
    ax.plot(labels, graficos["margenes"]["ebitda"], marker="o", color=_MPL_BLACK, label="EBITDA (%)")
    ax.plot(labels, graficos["margenes"]["neto"], marker="o", color=_MPL_ORANGE, label="Margen neto (%)")
    ax.set_title("Evolución de márgenes", fontsize=11, fontweight="bold", loc="left", color=_MPL_BLACK)
    ax.yaxis.set_major_formatter(FuncFormatter(lambda v, _=None: f"{v:.0f}%"))
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.grid(axis="y", color="#E4E4E2", linewidth=0.6)
    ax.set_axisbelow(True)
    ax.legend(loc="lower center", bbox_to_anchor=(0.5, -0.22), ncol=3, frameon=False, fontsize=9)
    return _figura_a_bytes(fig)


def _chart_png_estructura(graficos):
    fig, ax = plt.subplots(figsize=(6.4, 3.8))
    etiquetas = graficos["estructura_costes"]["labels"]
    valores = [max(v, 0) for v in graficos["estructura_costes"]["values"]]
    colores = [_MPL_BLACK, _MPL_GRAY, _MPL_ORANGE, _MPL_GREEN]
    if sum(valores) <= 0:
        valores = [1] * len(etiquetas)
    wedges, _texts, autotexts = ax.pie(
        valores, labels=None, colors=colores, startangle=90,
        autopct=lambda p: f"{p:.1f}%", pctdistance=0.78,
        wedgeprops=dict(width=0.45, edgecolor="white", linewidth=2),
    )
    for t in autotexts:
        t.set_color("white")
        t.set_fontsize(9)
    moneda = graficos.get("moneda", "")
    sufijo_moneda = f" — {moneda}" if moneda else ""
    ax.set_title(
        f"Estructura de costes — {graficos['estructura_costes']['periodo']}{sufijo_moneda}",
        fontsize=11, fontweight="bold", loc="left", color=_MPL_BLACK,
    )
    ax.legend(wedges, etiquetas, loc="center left", bbox_to_anchor=(1.0, 0.5), frameon=False, fontsize=8.5)
    return _figura_a_bytes(fig)


def _chart_png_metricas(graficos):
    fig, ax = plt.subplots(figsize=(6.4, 3.2))
    labels = graficos["metricas_clave"]["labels"]
    valores = graficos["metricas_clave"]["values"]
    colores = [_MPL_BLACK, _MPL_GREEN_DARK, _MPL_GREEN, _MPL_ORANGE][: len(labels)]
    bars = ax.barh(labels, valores, color=colores, edgecolor="white")
    ax.invert_yaxis()
    moneda = graficos.get("moneda", "")
    sufijo_moneda = f" — {moneda}" if moneda else ""
    ax.set_title(
        f"Métricas clave — {graficos['metricas_clave']['periodo']}{sufijo_moneda}",
        fontsize=11, fontweight="bold", loc="left", color=_MPL_BLACK,
    )
    ax.xaxis.set_major_formatter(FuncFormatter(_formato_eje_cifra))
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.grid(axis="x", color="#E4E4E2", linewidth=0.6)
    ax.set_axisbelow(True)
    for bar, v in zip(bars, valores):
        ax.text(bar.get_width(), bar.get_y() + bar.get_height()/2,
                "  " + _formato_eje_cifra(v), va="center", fontsize=9, color=_MPL_BLACK)
    return _figura_a_bytes(fig)


_CHART_RENDERERS = {
    "ingresos": _chart_png_ingresos,
    "margenes": _chart_png_margenes,
    "estructura": _chart_png_estructura,
    "metricas": _chart_png_metricas,
}

_CHART_TITULOS = {
    "ingresos": "Evolución de ingresos",
    "margenes": "Evolución de márgenes",
    "estructura": "Estructura de costes",
    "metricas": "Métricas clave del último ejercicio",
}


def añadir_markdown_a_docx(doc, texto):
    """Convierte texto Markdown ligero a párrafos Word con encabezados y listas.

    Soporta encabezados `#`/`##`/`###`, listas con `-`/`*`/numeradas, y
    formato inline `**negrita**`, `*cursiva*` y `` `código` ``. Las filas de
    tablas Markdown se transforman en líneas legibles separadas por ` · `.
    """
    if not texto:
        return

    lineas = texto.replace("\r\n", "\n").split("\n")
    for linea in lineas:
        stripped = linea.strip()
        if not stripped:
            continue
        # Filas separadoras de tabla Markdown: |---|---|
        if _RE_TABLA_SEP_MD.match(stripped):
            continue
        # Filas de tabla Markdown: convertir a "celda1 · celda2 · celda3"
        if stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 2:
            celdas = [c.strip() for c in stripped.strip("|").split("|")]
            celdas = [c for c in celdas if c]
            if celdas:
                p = doc.add_paragraph()
                _añadir_runs_con_formato(p, " · ".join(celdas))
            continue
        # Encabezados
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(" *"), level=3)
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(" *"), level=2)
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(" *"), level=1)
            continue
        # Listas con viñetas
        m = re.match(r"^[-*]\s+(.*)", stripped)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            _añadir_runs_con_formato(p, m.group(1))
            continue
        # Listas numeradas
        m = re.match(r"^\d+\.\s+(.*)", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            _añadir_runs_con_formato(p, m.group(1))
            continue
        # Párrafo normal
        p = doc.add_paragraph()
        _añadir_runs_con_formato(p, stripped)


def _añadir_tabla_datos_word(doc, datos, escala):
    """Añade la tabla de datos con anchuras fijas para que no se corte."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    headers = ["Métrica"] + [p["periodo"] for p in datos["periodos"]]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = False
    # Layout fijo: fuerza al word a respetar los anchos de columna.
    tblPr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)

    ancho_metrica = Inches(1.8)
    n_periodos = max(len(datos["periodos"]), 1)
    ancho_periodo = Inches((6.3 - 1.8) / n_periodos)

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.width = ancho_metrica if i == 0 else ancho_periodo
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9.5)

    metricas = [
        ("Ingresos", "ingresos", None),
        ("Coste de ventas", "coste_ventas", None),
        ("Margen bruto", "margen_bruto", "pct_margen_bruto"),
        ("Gastos operativos", "gastos_operativos", "pct_gastos_operativos"),
        ("EBITDA", "ebitda", "pct_ebitda"),
        ("Beneficio neto", "beneficio_neto", "pct_beneficio_neto"),
    ]
    for nombre_metrica, key, pct_key in metricas:
        row = table.add_row()
        row.cells[0].text = nombre_metrica
        row.cells[0].width = ancho_metrica
        for i, p in enumerate(datos["periodos"]):
            valor = formatear_importe_escalado(p[key], escala)
            if pct_key:
                valor += f" ({formatear_pct_es(p[pct_key])})"
            row.cells[i + 1].text = valor
            row.cells[i + 1].width = ancho_periodo
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)


def _añadir_seccion_markdown_docx(doc, texto):
    """Añade a Word un bloque de Markdown *sin* encabezado (el body de una sección)."""
    if not texto:
        return
    lineas = texto.replace("\r\n", "\n").split("\n")
    for linea in lineas:
        stripped = linea.strip()
        if not stripped:
            continue
        if _RE_TABLA_SEP_MD.match(stripped):
            continue
        if stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 2:
            celdas = [c.strip() for c in stripped.strip("|").split("|") if c.strip()]
            if celdas:
                p = doc.add_paragraph()
                _añadir_runs_con_formato(p, " · ".join(celdas))
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(" *"), level=3)
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(" *"), level=2)
            continue
        m = re.match(r"^[-*]\s+(.*)", stripped)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            _añadir_runs_con_formato(p, m.group(1))
            continue
        m = re.match(r"^\d+\.\s+(.*)", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            _añadir_runs_con_formato(p, m.group(1))
            continue
        p = doc.add_paragraph()
        _añadir_runs_con_formato(p, stripped)


def _insertar_grafico_docx(doc, chart_key, graficos):
    renderer = _CHART_RENDERERS.get(chart_key)
    if renderer is None:
        return
    buf = renderer(graficos)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(buf, width=Inches(5.8))


def crear_documento_word(datos, nota):
    """Crea un documento Word con los datos financieros, gráficos y memoria."""
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Márgenes algo más estrechos para tablas anchas
    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)

    titulo = doc.add_heading(f"Informe Financiero — {datos['nombre']}", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titulo.runs:
        run.font.color.rgb = RGBColor(0x43, 0xB0, 0x2A)

    escala = determinar_escala_tabla(datos)
    moneda = datos.get("moneda") or ""
    subtitulo = doc.add_paragraph()
    run = subtitulo.add_run(f"Ticker: {datos['ticker']}")
    run.font.bold = True
    if moneda:
        subtitulo.add_run(f"  ·  Moneda de reporting: {moneda}")
    doc.add_paragraph("")

    doc.add_heading("Datos Financieros", level=1)
    etiqueta = etiqueta_unidad(escala, moneda)
    if etiqueta:
        p = doc.add_paragraph()
        r = p.add_run(etiqueta)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x53, 0x56, 0x5A)
    _añadir_tabla_datos_word(doc, datos, escala)
    doc.add_paragraph("")

    doc.add_heading("Nota de Memoria Explicativa", level=1)

    graficos = preparar_datos_graficos(datos)
    secciones_asignadas, pendientes = segmentar_y_asignar(nota)

    for sec in secciones_asignadas:
        if sec.get("title"):
            doc.add_heading(sec["title"], level=2)
        _añadir_seccion_markdown_docx(doc, sec["body"])
        for ch in sec.get("charts", []):
            _insertar_grafico_docx(doc, ch, graficos)

    if pendientes:
        doc.add_heading("Resumen visual", level=2)
        for ch in pendientes:
            _insertar_grafico_docx(doc, ch, graficos)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def crear_documento_comparativo(datos_empresas, analisis):
    """Crea un documento Word con el análisis comparativo."""
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)

    titulo = doc.add_heading("Informe Comparativo Financiero", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titulo.runs:
        run.font.color.rgb = RGBColor(0x43, 0xB0, 0x2A)

    empresas_nombres = ", ".join([d["nombre"] for d in datos_empresas])
    doc.add_paragraph(f"Empresas comparadas: {empresas_nombres}")

    # Escala y moneda para el comparativo: se elige la escala a partir del
    # máximo entre todas las empresas para que la tabla sea comparable.
    datos_combinados = {
        "periodos": [d["periodos"][0] for d in datos_empresas if d.get("periodos")]
    }
    escala = determinar_escala_tabla(datos_combinados)
    monedas = sorted({d.get("moneda", "") for d in datos_empresas if d.get("moneda")})
    moneda_display = ", ".join(monedas) if monedas else ""
    etiqueta = etiqueta_unidad(escala, moneda_display)
    if etiqueta:
        p = doc.add_paragraph()
        r = p.add_run(etiqueta)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x53, 0x56, 0x5A)
        if len(monedas) > 1:
            warn = doc.add_paragraph()
            wr = warn.add_run(
                "Atención: las empresas comparadas reportan en monedas distintas; "
                "las cifras no se han convertido a una moneda única."
            )
            wr.italic = True
            wr.font.size = Pt(9)
            wr.font.color.rgb = RGBColor(0xED, 0x8B, 0x00)
    doc.add_paragraph("")

    doc.add_heading("Tabla Comparativa", level=1)

    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    headers = ["Métrica"] + [d["nombre"] for d in datos_empresas]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = False
    tblPr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)

    ancho_metrica = Inches(1.8)
    n_empresas = max(len(datos_empresas), 1)
    ancho_empresa = Inches((6.3 - 1.8) / n_empresas)

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.width = ancho_metrica if i == 0 else ancho_empresa
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9.5)

    metricas = [
        ("Ingresos", "ingresos", None),
        ("Coste de ventas", "coste_ventas", None),
        ("Margen bruto", "margen_bruto", "pct_margen_bruto"),
        ("Gastos operativos", "gastos_operativos", "pct_gastos_operativos"),
        ("EBITDA", "ebitda", "pct_ebitda"),
        ("Beneficio neto", "beneficio_neto", "pct_beneficio_neto"),
    ]

    for nombre_metrica, key, pct_key in metricas:
        row = table.add_row()
        row.cells[0].text = nombre_metrica
        row.cells[0].width = ancho_metrica
        for i, datos in enumerate(datos_empresas):
            p = datos["periodos"][0]  # Último periodo
            valor = formatear_importe_escalado(p[key], escala)
            if pct_key:
                valor += f" ({formatear_pct_es(p[pct_key])})"
            row.cells[i + 1].text = valor
            row.cells[i + 1].width = ancho_empresa
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)

    doc.add_paragraph("")
    doc.add_heading("Análisis Comparativo", level=1)
    añadir_markdown_a_docx(doc, analisis)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _md_inline_a_html(texto):
    """Convierte **negrita**, *cursiva* y `código` a tags reportlab/HTML seguros."""
    if not texto:
        return ""
    seguro = (
        texto.replace("&", "&amp;")
             .replace("<", "&lt;")
             .replace(">", "&gt;")
    )
    seguro = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", seguro)
    seguro = re.sub(r"(?<!\w)\*([^*]+)\*(?!\w)", r"<i>\1</i>", seguro)
    seguro = re.sub(r"`([^`]+)`", r"<font face='Courier'>\1</font>", seguro)
    return seguro


def _markdown_a_flowables_pdf(texto, estilos):
    """Convierte texto Markdown ligero a una lista de flowables para ReportLab."""
    flow = []
    if not texto:
        return flow

    lineas = texto.replace("\r\n", "\n").split("\n")
    bullets = []

    def _volcar_bullets():
        if not bullets:
            return
        items = [
            ListItem(Paragraph(_md_inline_a_html(b), estilos["Body"]), leftIndent=10)
            for b in bullets
        ]
        flow.append(ListFlowable(items, bulletType="bullet", start="•",
                                 leftIndent=14, bulletFontSize=9))
        flow.append(Spacer(1, 4))
        bullets.clear()

    for linea in lineas:
        stripped = linea.strip()
        if not stripped:
            _volcar_bullets()
            flow.append(Spacer(1, 4))
            continue
        if _RE_TABLA_SEP_MD.match(stripped):
            continue
        if stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 2:
            _volcar_bullets()
            celdas = [c.strip() for c in stripped.strip("|").split("|") if c.strip()]
            if celdas:
                flow.append(Paragraph(_md_inline_a_html(" · ".join(celdas)), estilos["Body"]))
            continue
        if stripped.startswith("### "):
            _volcar_bullets()
            flow.append(Paragraph(_md_inline_a_html(stripped[4:].strip(" *")), estilos["H3"]))
            continue
        if stripped.startswith("## "):
            _volcar_bullets()
            flow.append(Paragraph(_md_inline_a_html(stripped[3:].strip(" *")), estilos["H2"]))
            continue
        if stripped.startswith("# "):
            _volcar_bullets()
            flow.append(Paragraph(_md_inline_a_html(stripped[2:].strip(" *")), estilos["H1"]))
            continue
        m = re.match(r"^[-*]\s+(.*)", stripped)
        if m:
            bullets.append(m.group(1))
            continue
        m = re.match(r"^\d+\.\s+(.*)", stripped)
        if m:
            bullets.append(m.group(1))
            continue
        _volcar_bullets()
        flow.append(Paragraph(_md_inline_a_html(stripped), estilos["Body"]))

    _volcar_bullets()
    return flow


DELOITTE_GREEN = colors.HexColor("#86BC25")
DELOITTE_GREEN_DARK = colors.HexColor("#43B02A")
DELOITTE_BLACK = colors.HexColor("#000000")
DELOITTE_GRAY_11 = colors.HexColor("#53565A")
DELOITTE_GRAY_BORDER = colors.HexColor("#D0D0CE")
DELOITTE_GRAY_BG = colors.HexColor("#F5F5F4")


def _estilos_pdf():
    base = getSampleStyleSheet()
    return {
        "Title": ParagraphStyle(
            "DFTitle", parent=base["Title"],
            fontName="Helvetica-Bold", fontSize=20, leading=24,
            textColor=DELOITTE_BLACK, alignment=TA_LEFT, spaceAfter=6,
        ),
        "Eyebrow": ParagraphStyle(
            "DFEyebrow", fontName="Helvetica-Bold", fontSize=8, leading=11,
            textColor=DELOITTE_GREEN_DARK, spaceAfter=2,
        ),
        "Subtitle": ParagraphStyle(
            "DFSub", fontName="Helvetica", fontSize=10, leading=14,
            textColor=DELOITTE_GRAY_11, spaceAfter=14,
        ),
        "H1": ParagraphStyle(
            "DFH1", fontName="Helvetica-Bold", fontSize=14, leading=18,
            textColor=DELOITTE_BLACK, spaceBefore=12, spaceAfter=6,
        ),
        "H2": ParagraphStyle(
            "DFH2", fontName="Helvetica-Bold", fontSize=12, leading=16,
            textColor=DELOITTE_BLACK, spaceBefore=10, spaceAfter=5,
        ),
        "H3": ParagraphStyle(
            "DFH3", fontName="Helvetica-Bold", fontSize=10.5, leading=14,
            textColor=DELOITTE_GRAY_11, spaceBefore=8, spaceAfter=3,
        ),
        "Body": ParagraphStyle(
            "DFBody", fontName="Helvetica", fontSize=9.5, leading=14,
            textColor=DELOITTE_BLACK, alignment=TA_JUSTIFY, spaceAfter=4,
        ),
    }


def _dibujar_branding_pdf(canvas, doc):
    canvas.saveState()
    # Barra lateral verde Deloitte
    canvas.setFillColor(DELOITTE_GREEN)
    canvas.rect(0, 0, 6, A4[1], stroke=0, fill=1)
    # Cabecera
    canvas.setFillColor(DELOITTE_BLACK)
    canvas.setFont("Helvetica-Bold", 10)
    canvas.drawString(18 * mm, A4[1] - 12 * mm, "DFin AI")
    canvas.setFillColor(DELOITTE_GREEN_DARK)
    canvas.setFont("Helvetica-Bold", 10)
    canvas.drawString(18 * mm + 18, A4[1] - 12 * mm, "· Financial Intelligence")
    canvas.setStrokeColor(DELOITTE_GREEN)
    canvas.setLineWidth(1)
    canvas.line(18 * mm, A4[1] - 14 * mm, A4[0] - 18 * mm, A4[1] - 14 * mm)
    # Pie
    canvas.setFillColor(DELOITTE_GRAY_11)
    canvas.setFont("Helvetica", 8)
    canvas.drawString(18 * mm, 10 * mm, "DFin AI · Informe confidencial")
    canvas.drawRightString(A4[0] - 18 * mm, 10 * mm, f"Página {canvas.getPageNumber()}")
    canvas.restoreState()


def _construir_tabla_datos_pdf(datos, ancho_total, escala):
    """Construye la tabla de datos financieros para PDF ajustada al ancho dado."""
    headers = ["Métrica"] + [p["periodo"] for p in datos["periodos"]]
    filas_metricas = [
        ("Ingresos", "ingresos", None),
        ("Coste de ventas", "coste_ventas", None),
        ("Margen bruto", "margen_bruto", "pct_margen_bruto"),
        ("Gastos operativos", "gastos_operativos", "pct_gastos_operativos"),
        ("EBITDA", "ebitda", "pct_ebitda"),
        ("Beneficio neto", "beneficio_neto", "pct_beneficio_neto"),
    ]
    body_rows = []
    for nombre_metrica, key, pct_key in filas_metricas:
        fila = [nombre_metrica]
        for p in datos["periodos"]:
            valor = formatear_importe_escalado(p[key], escala)
            if pct_key:
                valor += f" ({formatear_pct_es(p[pct_key])})"
            fila.append(valor)
        body_rows.append(fila)

    # Anchuras fijas para evitar cortes en la tabla
    n_periodos = max(len(datos["periodos"]), 1)
    ancho_metrica = min(42 * mm, ancho_total * 0.28)
    ancho_periodo = (ancho_total - ancho_metrica) / n_periodos
    col_widths = [ancho_metrica] + [ancho_periodo] * n_periodos

    tabla = Table([headers] + body_rows, colWidths=col_widths, hAlign="LEFT", repeatRows=1)
    tabla.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), DELOITTE_BLACK),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, 0), 8.5),
        ("ALIGN", (1, 0), (-1, 0), "RIGHT"),
        ("LINEBELOW", (0, 0), (-1, 0), 1.5, DELOITTE_GREEN),
        ("BOTTOMPADDING", (0, 0), (-1, 0), 7),
        ("TOPPADDING", (0, 0), (-1, 0), 7),
        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 1), (-1, -1), 8.5),
        ("ALIGN", (1, 1), (-1, -1), "RIGHT"),
        ("LINEBELOW", (0, 1), (-1, -1), 0.25, DELOITTE_GRAY_BORDER),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, DELOITTE_GRAY_BG]),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("FONTNAME", (0, 1), (0, -1), "Helvetica-Bold"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    return tabla


def _flowable_chart_pdf(chart_key, graficos, ancho_disponible):
    renderer = _CHART_RENDERERS.get(chart_key)
    if renderer is None:
        return []
    buf = renderer(graficos)
    img = RLImage(buf)
    # Escala manteniendo aspecto para ocupar como máximo el 90% del ancho.
    max_w = ancho_disponible * 0.92
    w, h = img.drawWidth, img.drawHeight
    escala = min(max_w / w, 1)
    img.drawWidth = w * escala
    img.drawHeight = h * escala
    return [Spacer(1, 6), img, Spacer(1, 10)]


def crear_documento_pdf(datos, nota):
    """Crea un PDF de la nota de memoria con branding DFin AI / paleta Deloitte."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=18 * mm, rightMargin=18 * mm,
        topMargin=22 * mm, bottomMargin=18 * mm,
        title=f"Informe DFin AI — {datos['nombre']}",
        author="DFin AI",
    )
    ancho_disponible = A4[0] - 36 * mm
    estilos = _estilos_pdf()
    story = []

    escala = determinar_escala_tabla(datos)
    moneda = datos.get("moneda") or ""

    story.append(Paragraph("INFORME FINANCIERO · NOTA DE MEMORIA", estilos["Eyebrow"]))
    story.append(Paragraph(datos["nombre"], estilos["Title"]))
    sub = f"Ticker <b>{datos['ticker']}</b> &nbsp;·&nbsp; {len(datos['periodos'])} ejercicios analizados"
    if moneda:
        sub += f" &nbsp;·&nbsp; Moneda: <b>{moneda}</b>"
    story.append(Paragraph(sub, estilos["Subtitle"]))

    story.append(Paragraph("Datos financieros — serie histórica", estilos["H1"]))
    etiqueta = etiqueta_unidad(escala, moneda)
    if etiqueta:
        story.append(Paragraph(f"<i>{etiqueta}</i>", estilos["Body"]))
        story.append(Spacer(1, 4))
    story.append(_construir_tabla_datos_pdf(datos, ancho_disponible, escala))
    story.append(Spacer(1, 14))

    story.append(Paragraph("Nota de memoria explicativa", estilos["H1"]))

    graficos = preparar_datos_graficos(datos)
    secciones_asignadas, pendientes = segmentar_y_asignar(nota)

    for sec in secciones_asignadas:
        if sec.get("title"):
            story.append(Paragraph(_md_inline_a_html(sec["title"]), estilos["H2"]))
        story.extend(_markdown_a_flowables_pdf(sec["body"], estilos))
        for ch in sec.get("charts", []):
            story.extend(_flowable_chart_pdf(ch, graficos, ancho_disponible))

    if pendientes:
        story.append(Paragraph("Resumen visual", estilos["H2"]))
        for ch in pendientes:
            story.extend(_flowable_chart_pdf(ch, graficos, ancho_disponible))

    doc.build(story, onFirstPage=_dibujar_branding_pdf, onLaterPages=_dibujar_branding_pdf)
    buffer.seek(0)
    return buffer


def formatear_numero(valor):
    """Formatea números grandes en formato europeo (miles con `.`, decimales con `,`).

    Escala el número a la unidad más cómoda visualmente:
    - < 1.000        → entero o 2 decimales según haya parte decimal
    - < 1.000.000    → entero con separadores de miles
    - < mil millones → millones con 2 decimales y sufijo "M"
    - < 1 billón     → miles de millones con 2 decimales y sufijo "mil M"
    - resto          → billones con 2 decimales y sufijo "B" (1 B = 10^12)
    """
    try:
        v = float(valor)
    except (TypeError, ValueError):
        return ""
    abs_v = abs(v)
    if abs_v >= 1e12:
        return _fmt_es(v / 1e12, 2) + " B"
    if abs_v >= 1e9:
        return _fmt_es(v / 1e9, 2) + " mil M"
    if abs_v >= 1e6:
        return _fmt_es(v / 1e6, 2) + " M"
    if abs_v >= 1000:
        return _fmt_es(v, 0)
    if abs_v == 0 or float(v).is_integer():
        return _fmt_es(v, 0)
    return _fmt_es(v, 2)


def _fmt_es(valor, decimales):
    """Aplica formato español: punto como separador de miles, coma como decimal."""
    s = f"{valor:,.{decimales}f}"
    return s.replace(",", "\x00").replace(".", ",").replace("\x00", ".")


def determinar_escala_tabla(datos):
    """Escoge una escala común para toda la tabla según la cifra máxima.

    Devuelve un dict con `divisor`, `sufijo` (unidad legible) y `decimales`.
    Usa la misma escala para todas las filas/columnas para garantizar que
    los importes caben en la tabla y son comparables entre sí.
    """
    vals = []
    for p in datos.get("periodos", []):
        for k in ("ingresos", "coste_ventas", "margen_bruto",
                  "gastos_operativos", "ebitda", "beneficio_neto"):
            v = p.get(k, 0) or 0
            vals.append(abs(float(v)))
    max_v = max(vals) if vals else 0

    if max_v >= 1e12:
        return {"divisor": 1e9, "sufijo": "miles de millones", "decimales": 2}
    if max_v >= 1e9:
        return {"divisor": 1e6, "sufijo": "millones", "decimales": 0}
    if max_v >= 1e6:
        return {"divisor": 1e6, "sufijo": "millones", "decimales": 1}
    if max_v >= 1e3:
        return {"divisor": 1e3, "sufijo": "miles", "decimales": 0}
    return {"divisor": 1, "sufijo": "", "decimales": 0}


def escala_para_estado(bloque):
    """Elige una escala común para un estado financiero completo.

    `bloque` es la entrada de `datos["estados"][clave]` con `periodos` y `filas`.
    Devuelve la misma estructura que `determinar_escala_tabla`. Las filas que
    parecen ratios o porcentajes (`Tax Rate`, `EPS`, `Acciones medias`) se
    excluyen del cálculo para no falsear la escala con valores pequeños o
    enormes según el caso.
    """
    EXCLUIDOS = (
        "BPA", "EPS", "Acciones medias", "Tasa impositiva", "Number", "Acciones emitidas",
    )
    vals = []
    for fila in bloque.get("filas", []):
        nombre = fila.get("concepto", "")
        if any(t.lower() in nombre.lower() for t in EXCLUIDOS):
            continue
        for v in fila.get("valores", []):
            if v is None:
                continue
            try:
                vals.append(abs(float(v)))
            except (TypeError, ValueError):
                pass
    max_v = max(vals) if vals else 0
    if max_v >= 1e12:
        return {"divisor": 1e9, "sufijo": "miles de millones", "decimales": 2}
    if max_v >= 1e9:
        return {"divisor": 1e6, "sufijo": "millones", "decimales": 0}
    if max_v >= 1e6:
        return {"divisor": 1e6, "sufijo": "millones", "decimales": 1}
    if max_v >= 1e3:
        return {"divisor": 1e3, "sufijo": "miles", "decimales": 0}
    return {"divisor": 1, "sufijo": "", "decimales": 0}


def formatear_celda_estado(valor, escala, concepto=""):
    """Aplica la escala del estado y devuelve la cadena formateada en español.

    Las filas tipo BPA/EPS, tasas impositivas o número de acciones se muestran
    SIN aplicar la escala (no son magnitudes monetarias en millones).
    """
    if valor is None:
        return ""
    try:
        v = float(valor)
    except (TypeError, ValueError):
        return ""
    nombre = (concepto or "").lower()
    if any(t in nombre for t in ("bpa", "eps", "acciones medias", "tasa impositiva", "acciones emitidas", "ordinary shares number")):
        # Magnitudes per-share o conteos: sin escala, hasta 2 decimales.
        decimales = 2 if abs(v) < 1000 else 0
        return _fmt_es(v, decimales)
    divisor = escala.get("divisor", 1) or 1
    decimales = escala.get("decimales", 0)
    return _fmt_es(v / divisor, decimales)



def formatear_importe_escalado(valor, escala):
    """Formatea un importe aplicando una escala y formato español."""
    divisor = escala.get("divisor", 1) or 1
    decimales = escala.get("decimales", 0)
    return _fmt_es(valor / divisor, decimales)


def formatear_pct_es(pct):
    """Formatea un porcentaje con un decimal y formato español."""
    return _fmt_es(pct, 1) + "%"


def etiqueta_unidad(escala, moneda):
    """Etiqueta legible del tipo 'Cifras en millones de USD'."""
    sufijo = escala.get("sufijo", "")
    moneda = (moneda or "").strip()
    if sufijo and moneda:
        return f"Cifras en {sufijo} de {moneda}"
    if sufijo:
        return f"Cifras en {sufijo}"
    if moneda:
        return f"Cifras en {moneda}"
    return ""


# KPIs por defecto (siempre se muestran en el cuadro de "Datos relevantes").
_KPI_BASE = [
    {"label": "Ingresos", "fuente": ("pyg", "ingresos"), "tipo": "moneda"},
    {"label": "Margen bruto", "fuente": ("pyg", "margen_bruto"), "pct": "pct_margen_bruto", "tipo": "moneda"},
    {"label": "EBITDA", "fuente": ("pyg", "ebitda"), "pct": "pct_ebitda", "tipo": "moneda"},
    {"label": "EBIT", "fuente": ("pyg", "ebit"), "pct": "pct_ebit", "tipo": "moneda"},
    {"label": "Beneficio neto", "fuente": ("pyg", "beneficio_neto"), "pct": "pct_beneficio_neto", "tipo": "moneda"},
    {"label": "BPA básico", "fuente": ("pyg", "bpa_basico"), "tipo": "per_share"},
]

# KPIs adicionales por enfoque (texto literal del chip seleccionado).
_KPI_POR_ENFOQUE = {
    "Estructura de costes": [
        {"label": "Coste de ventas", "fuente": ("pyg", "coste_ventas"), "tipo": "moneda"},
        {"label": "SG&A", "fuente": ("pyg", "sga"), "pct": "pct_sga", "tipo": "moneda"},
        {"label": "I+D", "fuente": ("pyg", "rd"), "pct": "pct_rd", "tipo": "moneda"},
        {"label": "Amortizaciones (D&A)", "fuente": ("pyg", "da"), "pct": "pct_da", "tipo": "moneda"},
    ],
    "Fiscalidad y resultado financiero": [
        {"label": "Resultado financiero neto", "fuente": ("pyg", "resultado_financiero_neto"), "tipo": "moneda"},
        {"label": "Resultado antes de impuestos", "fuente": ("pyg", "resultado_antes_impuestos"), "pct": "pct_rai", "tipo": "moneda"},
        {"label": "Impuesto sobre beneficios", "fuente": ("pyg", "impuestos"), "tipo": "moneda"},
        {"label": "Tasa impositiva efectiva", "fuente": ("pyg", "tasa_impositiva"), "tipo": "porcentaje"},
    ],
    "Retorno al accionista (BPA, dividendos)": [
        {"label": "BPA diluido", "fuente": ("pyg", "bpa_diluido"), "tipo": "per_share"},
        {"label": "Dividendos pagados", "fuente": ("cashflow", "Dividendos pagados en efectivo"), "tipo": "moneda"},
    ],
    "Endeudamiento y solvencia": [
        {"label": "Deuda total", "fuente": ("balance_sheet", "Deuda total"), "tipo": "moneda"},
        {"label": "Deuda neta", "fuente": ("balance_sheet", "Deuda neta"), "tipo": "moneda"},
        {"label": "Patrimonio neto", "fuente": ("balance_sheet", "Patrimonio neto total"), "tipo": "moneda"},
        {"label": "Activos totales", "fuente": ("balance_sheet", "Activos totales"), "tipo": "moneda"},
    ],
    "Generación de caja": [
        {"label": "Flujo de caja operativo", "fuente": ("cashflow", "Flujo de caja operativo"), "tipo": "moneda"},
        {"label": "CAPEX", "fuente": ("cashflow", "Inversiones en inmovilizado (CAPEX)"), "tipo": "moneda"},
        {"label": "Flujo de caja libre", "fuente": ("cashflow", "Flujo de caja libre"), "tipo": "moneda"},
    ],
    "Capital circulante": [
        {"label": "Capital circulante", "fuente": ("balance_sheet", "Capital circulante"), "tipo": "moneda"},
        {"label": "Existencias", "fuente": ("balance_sheet", "Existencias"), "tipo": "moneda"},
        {"label": "Cuentas por cobrar", "fuente": ("balance_sheet", "Cuentas por cobrar comerciales"), "tipo": "moneda"},
        {"label": "Cuentas por pagar", "fuente": ("balance_sheet", "Cuentas por pagar comerciales"), "tipo": "moneda"},
    ],
    "Inversiones (CAPEX)": [
        {"label": "CAPEX", "fuente": ("cashflow", "Inversiones en inmovilizado (CAPEX)"), "tipo": "moneda"},
        {"label": "Inmovilizado material neto", "fuente": ("balance_sheet", "Inmovilizado material neto"), "tipo": "moneda"},
    ],
    # "Evolución de ingresos" y "Margen y rentabilidad" se cubren con los KPIs base.
}


def _buscar_fila_estado(datos, estado_key, concepto):
    """Busca una fila por concepto exacto en datos['estados'][estado_key]."""
    estado = (datos.get("estados") or {}).get(estado_key) or {}
    for fila in estado.get("filas", []):
        if fila.get("concepto") == concepto:
            return fila.get("valores", [])
    return []


def _valores_kpi(datos, fuente):
    """Devuelve [valor_ultimo, valor_anterior] para una fuente declarada en un KPI."""
    origen, clave = fuente
    if origen == "pyg":
        periodos = datos.get("periodos", []) or []
        ultimo = periodos[0].get(clave) if len(periodos) > 0 else None
        anterior = periodos[1].get(clave) if len(periodos) > 1 else None
        return ultimo, anterior
    valores = _buscar_fila_estado(datos, origen, clave)
    ultimo = valores[0] if len(valores) > 0 else None
    anterior = valores[1] if len(valores) > 1 else None
    return ultimo, anterior


def construir_metricas_relevantes(datos, enfoques_pred):
    """Selecciona las métricas a destacar en función del enfoque del informe.

    `enfoques_pred` es la lista de chips seleccionados por el usuario en el
    configurador. Si está vacía, devolvemos sólo los KPIs base.
    """
    seleccion = list(_KPI_BASE)
    vistos = {(k["fuente"][0], k["fuente"][1]) for k in seleccion}
    for enfoque in enfoques_pred or []:
        for kpi in _KPI_POR_ENFOQUE.get(enfoque, []):
            firma = (kpi["fuente"][0], kpi["fuente"][1])
            if firma in vistos:
                continue
            seleccion.append(kpi)
            vistos.add(firma)

    periodos = datos.get("periodos", []) or []
    label_ult = periodos[0]["periodo"] if periodos else ""
    label_ant = periodos[1]["periodo"] if len(periodos) > 1 else ""

    # Escala común para todas las cifras monetarias del cuadro, basada en el
    # mayor importe presente (igual criterio que las pestañas de Yahoo).
    vals_mon = []
    for kpi in seleccion:
        if kpi["tipo"] not in ("moneda",):
            continue
        u, a = _valores_kpi(datos, kpi["fuente"])
        for v in (u, a):
            if v is not None:
                try:
                    vals_mon.append(abs(float(v)))
                except (TypeError, ValueError):
                    pass
    max_v = max(vals_mon) if vals_mon else 0
    if max_v >= 1e9:
        escala = {"divisor": 1e6, "sufijo": "millones", "decimales": 0}
    elif max_v >= 1e6:
        escala = {"divisor": 1e6, "sufijo": "millones", "decimales": 1}
    elif max_v >= 1e3:
        escala = {"divisor": 1e3, "sufijo": "miles", "decimales": 0}
    else:
        escala = {"divisor": 1, "sufijo": "", "decimales": 0}

    def _fmt_valor(v, tipo):
        if v is None:
            return "—"
        try:
            vf = float(v)
        except (TypeError, ValueError):
            return "—"
        if tipo == "moneda":
            return _fmt_es(vf / escala["divisor"], escala["decimales"])
        if tipo == "per_share":
            return _fmt_es(vf, 2)
        if tipo == "porcentaje":
            return _fmt_es(vf, 1) + "%"
        return _fmt_es(vf, 0)

    def _fmt_variacion(ult, ant):
        if ult is None or ant is None:
            return "—", "neutro"
        try:
            u = float(ult); a = float(ant)
        except (TypeError, ValueError):
            return "—", "neutro"
        if a == 0:
            return "—", "neutro"
        var = ((u - a) / abs(a)) * 100
        sentido = "positivo" if var >= 0 else "negativo"
        signo = "+" if var >= 0 else "−"
        return f"{signo}{_fmt_es(abs(var), 1)}%", sentido

    filas = []
    for kpi in seleccion:
        u, a = _valores_kpi(datos, kpi["fuente"])
        # No mostrar la fila si ambos valores faltan.
        if u is None and a is None:
            continue
        fila = {
            "label": kpi["label"],
            "valor_ultimo": _fmt_valor(u, kpi["tipo"]),
            "valor_anterior": _fmt_valor(a, kpi["tipo"]),
            "tipo": kpi["tipo"],
        }
        pct_clave = kpi.get("pct")
        if pct_clave and periodos:
            pct_ult = periodos[0].get(pct_clave)
            pct_ant = periodos[1].get(pct_clave) if len(periodos) > 1 else None
            fila["pct_ultimo"] = (_fmt_es(pct_ult, 1) + "%") if pct_ult is not None else ""
            fila["pct_anterior"] = (_fmt_es(pct_ant, 1) + "%") if pct_ant is not None else ""
        var_txt, var_sentido = _fmt_variacion(u, a)
        fila["variacion"] = var_txt
        fila["variacion_sentido"] = var_sentido
        filas.append(fila)

    return {
        "filas": filas,
        "periodo_ultimo": label_ult,
        "periodo_anterior": label_ant,
        "escala": escala,
        "moneda": datos.get("moneda", ""),
        "etiqueta_unidad": etiqueta_unidad(escala, datos.get("moneda", "")),
    }


def preparar_datos_graficos(datos):
    """Construye un dict JSON-serializable para renderizar gráficos Chart.js.

    Ordena los periodos cronológicamente (más antiguo → más reciente) para
    que la serie temporal se lea de izquierda a derecha en los gráficos.
    """
    # `periodos[0]` es el más reciente; se invierte para series temporales.
    periodos_ord = list(reversed(datos["periodos"]))

    def _etiqueta_anio(periodo_str):
        return periodo_str[:4] if len(periodo_str) >= 4 else periodo_str

    labels = [_etiqueta_anio(p["periodo"]) for p in periodos_ord]

    ultimo = datos["periodos"][0]
    # Residual = impuestos, intereses, D&A (para cuadrar estructura de costes).
    residual = (
        ultimo["ingresos"]
        - ultimo["coste_ventas"]
        - ultimo["gastos_operativos"]
        - ultimo["beneficio_neto"]
    )
    if residual < 0:
        residual = 0.0

    return {
        "labels": labels,
        "moneda": datos.get("moneda", ""),
        "ingresos": [p["ingresos"] for p in periodos_ord],
        "margenes": {
            "margen_bruto": [p["pct_margen_bruto"] for p in periodos_ord],
            "ebitda": [p["pct_ebitda"] for p in periodos_ord],
            "neto": [p["pct_beneficio_neto"] for p in periodos_ord],
        },
        "estructura_costes": {
            "labels": [
                "Coste de ventas",
                "Gastos operativos",
                "Otros gastos (impuestos, intereses, D&A)",
                "Beneficio neto",
            ],
            "values": [
                ultimo["coste_ventas"],
                ultimo["gastos_operativos"],
                residual,
                ultimo["beneficio_neto"],
            ],
            "periodo": ultimo["periodo"],
        },
        "metricas_clave": {
            "labels": ["Ingresos", "Margen bruto", "EBITDA", "Beneficio neto"],
            "values": [
                ultimo["ingresos"],
                ultimo["margen_bruto"],
                ultimo["ebitda"],
                ultimo["beneficio_neto"],
            ],
            "periodo": ultimo["periodo"],
        },
    }


_IR_PENDIENTES_CACHE = {"ts": 0.0, "count": None}


def _contar_pendientes_ir(ttl_segundos=30):
    """Devuelve cuántos correos están pendientes ahora mismo.

    Usa la misma función `_leer_bandeja_gmail()` que la vista, para que el
    contador de la home y la bandeja real coincidan SIEMPRE (mismo filtro
    estricto: asunto que EMPIEZA por `{TICKER}-Inversores-` y recibido en
    las últimas 2 horas exactas). Caché TTL 30 s para no martillear IMAP.
    """
    import time
    if not GMAIL_USER or not GMAIL_APP_PASSWORD:
        return None
    ahora = time.time()
    if ahora - _IR_PENDIENTES_CACHE["ts"] < ttl_segundos and _IR_PENDIENTES_CACHE["count"] is not None:
        return _IR_PENDIENTES_CACHE["count"]
    try:
        correos = _leer_bandeja_gmail(_prefijo_asunto_actual(), horas_max=2)
        # Anotar cada correo con su estado de envío
        for c in correos:
            envio = _IR_ENVIADOS.get(_clave_envio(c))
            c["_enviado"] = envio  # dict o None
        # Pendientes = los no enviados
        pendientes = [c for c in correos if not c["_enviado"]]
        n = len(pendientes)
        _IR_PENDIENTES_CACHE.update({"ts": ahora, "count": n})
        return n
    except Exception:
        return None


def _empresas_disponibles():
    """Devuelve la lista de subcarpetas en Empresas/ ordenada alfabéticamente.

    Solo incluye las que parecen una carpeta de ticker válida (contiene
    alguno de: docs/, logo/, tema.txt, lineas_maestras.txt).
    """
    base = _ROOT_DIR / "Empresas"
    if not base.exists():
        return []
    out = []
    for d in sorted(base.iterdir()):
        if not d.is_dir():
            continue
        # Heurística suave: la carpeta cuenta como empresa si tiene algo
        # significativo dentro. Si está vacía o solo tiene `auditoria/`, la
        # filtramos para no ofrecer empresas vacías.
        marcadores = ["docs", "logo", "tema.txt", "lineas_maestras.txt"]
        if any((d / m).exists() for m in marcadores):
            out.append(d.name)
    return out


def _set_ticker_activo(nuevo_ticker):
    """Cambia el TICKER_FIJO activo y reinicializa branding / RAG / memoria."""
    global TICKER_FIJO
    nuevo = (nuevo_ticker or "").strip().upper()
    if not nuevo:
        return False
    TICKER_FIJO = nuevo
    _BRANDING["empresa"] = nuevo
    _BRANDING["tenant_sufijo"] = nuevo
    # Reset cachés dependientes del ticker.
    _TICKER_FIJO_DATOS["ticker"] = None
    _TICKER_FIJO_DATOS["datos"] = None
    _RAG_INDEX["ticker"] = None
    _RAG_INDEX["index"] = None
    _IR_PENDIENTES_CACHE.update({"ts": 0.0, "count": None})
    _IR_ENVIADOS.clear()
    _IR_CONVERSACIONES.clear()
    _GMAIL_CACHE.clear()
    # Reconstruir auditoría apuntando a la carpeta del nuevo ticker.
    nuevo_dir = _ROOT_DIR / "Empresas" / TICKER_FIJO / "auditoria"
    nuevo_dir.mkdir(parents=True, exist_ok=True)
    nuevo_log = nuevo_dir / "sesion.txt"
    # Reemplaza el FileHandler en el logger raíz para escribir en el nuevo path.
    root = _logging.getLogger()
    for h in list(root.handlers):
        if isinstance(h, _logging.FileHandler):
            root.removeHandler(h)
            try:
                h.close()
            except Exception:
                pass
    new_handler = _logging.FileHandler(nuevo_log, mode="a", encoding="utf-8")
    new_handler.setFormatter(_log_formatter)
    root.addHandler(new_handler)
    # Recalcular acumulado de coste para el nuevo ticker.
    global _AUDITORIA_DIR, _LOG_FILE, _COSTE_FILE
    _AUDITORIA_DIR = nuevo_dir
    _LOG_FILE = nuevo_log
    _COSTE_FILE = _AUDITORIA_DIR / "coste_tokens.txt"
    _COSTE_ACUMULADO["usd"] = 0.0
    _COSTE_ACUMULADO["llamadas"] = 0
    _cargar_acumulado_coste()
    logger.info("=" * 70)
    logger.info("TICKER cambiado a %s · fichero=%s", TICKER_FIJO, _LOG_FILE)
    # Cargar memoria de IR y datos básicos.
    try:
        _cargar_memoria_ir()
        _cargar_datos_ticker_fijo()
    except Exception as e:
        logger.exception("Fallo al inicializar datos del nuevo ticker: %s", e)
    try:
        _construir_rag_si_procede(forzar=False)
    except Exception as e:
        logger.exception("Fallo construyendo RAG para %s: %s", TICKER_FIJO, e)
    return True


@app.route("/seleccionar-empresa", methods=["POST"])
def seleccionar_empresa():
    from flask import redirect
    ticker = request.form.get("ticker", "").strip().upper()
    if not ticker:
        return redirect("/")
    _set_ticker_activo(ticker)
    return redirect("/")


@app.route("/")
def index():
    # Si no hay ticker activo, mostramos el selector inicial.
    if not TICKER_FIJO:
        return render_template(
            "seleccionar_empresa.html",
            empresas=_empresas_disponibles(),
        )
    return render_template(
        "index.html",
        ir_pendientes=_contar_pendientes_ir(),
        ir_prefijo=_prefijo_asunto_actual(),
    )


def _render_secciones_para_resultado(nota):
    """Convierte una memoria Markdown en la estructura que espera la plantilla."""
    secciones_asignadas, pendientes = segmentar_y_asignar(nota)
    secciones_html = []
    for sec in secciones_asignadas:
        body_html = markdown.markdown(
            sec.get("body", ""),
            extensions=["tables", "fenced_code", "sane_lists"],
        )
        secciones_html.append({
            "title": sec.get("title"),
            "body_html": body_html,
            "charts": sec.get("charts", []),
        })
    return secciones_html, pendientes


# Cache (a nivel proceso) de los datos de Yahoo del TICKER_FIJO, para evitar
# pegar a yfinance en cada navegación dentro de la misma sesión.
_TICKER_FIJO_DATOS = {"ticker": None, "datos": None}

# Índice RAG global por ticker (en proceso). Se construye al arrancar la app
# y al refrescar el ticker fijado. Si el ticker no tiene corpus, el índice
# está vacío y _resolver_correo_ir simplemente no inyecta fragmentos.
_RAG_INDEX = {"ticker": None, "index": None}


def _construir_rag_si_procede(forzar: bool = False):
    """Construye el índice RAG para TICKER_FIJO si hay corpus disponible."""
    if not TICKER_FIJO:
        return None
    if (not forzar and _RAG_INDEX["ticker"] == TICKER_FIJO
            and _RAG_INDEX["index"] is not None):
        return _RAG_INDEX["index"]
    try:
        idx = _rag.indexar_ticker(TICKER_FIJO, forzar=forzar)
        _RAG_INDEX["ticker"] = TICKER_FIJO
        _RAG_INDEX["index"] = idx
        return idx
    except Exception as e:
        _logging.getLogger("dfin.rag").exception("Fallo construyendo RAG: %s", e)
        return None


def _cargar_datos_ticker_fijo():
    if not TICKER_FIJO:
        return None
    if _TICKER_FIJO_DATOS.get("ticker") == TICKER_FIJO and _TICKER_FIJO_DATOS.get("datos"):
        return _TICKER_FIJO_DATOS["datos"]
    datos = obtener_datos_financieros(TICKER_FIJO)
    _TICKER_FIJO_DATOS["ticker"] = TICKER_FIJO
    _TICKER_FIJO_DATOS["datos"] = datos
    if datos and datos.get("nombre"):
        _BRANDING["empresa"] = datos["nombre"]
        _BRANDING["tenant_sufijo"] = datos["nombre"]
    return datos


def _listar_informes_descargados():
    """Devuelve la lista de informes guardados en Empresas/{TICKER}/descargas/.

    Cada item: {nombre, tamano_kb, fecha, url, extension}. Lista vacía si no
    hay TICKER_FIJO o la carpeta aún no existe / está vacía.
    """
    from datetime import datetime as _dt
    import urllib.parse as _up
    d = _dir_descargas_ticker()
    items = []
    if d and d.exists():
        for ruta in sorted(d.iterdir(), reverse=True):
            if ruta.suffix.lower() not in (".docx", ".pdf"):
                continue
            try:
                stat = ruta.stat()
            except OSError:
                continue
            items.append({
                "nombre": ruta.name,
                "tamano_kb": round(stat.st_size / 1024, 1),
                "fecha": _dt.fromtimestamp(stat.st_mtime).strftime("%d/%m/%Y %H:%M"),
                "url": "/descargar-archivado/" + _up.quote(ruta.name),
                "extension": ruta.suffix.lstrip(".").upper(),
            })
    return items


def _render_datos_yahoo(datos, modelo, error=None):
    cache_id = str(uuid.uuid4())
    _download_cache[cache_id] = {
        "datos": datos,
        "nota": None,
        "modelo": modelo,
        "instrucciones": "",
        "chat_historial": [],
    }
    return render_template(
        "datos_yahoo.html",
        datos=datos,
        cache_id=cache_id,
        modelo_actual=modelo,
        formatear=formatear_numero,
        escala_para_estado=escala_para_estado,
        formatear_celda_estado=formatear_celda_estado,
        error=error,
        informes_previos=_listar_informes_descargados(),
    )


@app.route("/analisis", methods=["GET", "POST"])
def analisis():
    """Si hay TICKER_FIJO, salta directamente al paso 2 con esa compañía.

    En caso contrario mantiene el flujo en dos pasos: el GET muestra el form
    de selección de ticker y el POST procesa el ticker introducido.
    """
    modelo_default = MODELO_POR_DEFECTO

    if TICKER_FIJO:
        try:
            datos = _cargar_datos_ticker_fijo()
            if datos is None:
                return render_template(
                    "analisis.html",
                    error=(f"No se encontraron datos financieros para el ticker "
                           f"fijado '{TICKER_FIJO}'. Revisa el símbolo bursátil "
                           f"con el que arrancaste la app."),
                    modelo_seleccionado=modelo_default,
                )
            return _render_datos_yahoo(datos, modelo_default)
        except Exception as e:
            return render_template(
                "analisis.html",
                error=f"Error al cargar datos de {TICKER_FIJO}: {e}",
                modelo_seleccionado=modelo_default,
            )

    if request.method == "POST":
        ticker_symbol = request.form.get("ticker", "").strip().upper()
        modelo = _normalizar_modelo(request.form.get("modelo", MODELO_POR_DEFECTO))
        if not ticker_symbol:
            return render_template(
                "analisis.html",
                error="Introduce un ticker válido.",
                modelo_seleccionado=modelo,
            )

        try:
            datos = obtener_datos_financieros(ticker_symbol)
            if datos is None:
                return render_template(
                    "analisis.html",
                    error=f"No se encontraron datos financieros para '{ticker_symbol}'.",
                    modelo_seleccionado=modelo,
                )
            return _render_datos_yahoo(datos, modelo)
        except Exception as e:
            return render_template(
                "analisis.html",
                error=f"Error: {str(e)}",
                modelo_seleccionado=modelo,
            )

    return render_template(
        "analisis.html",
        modelo_seleccionado=MODELO_POR_DEFECTO,
    )


@app.route("/generar_informe", methods=["POST"])
def generar_informe():
    """Paso 2: recibe el configurador y genera el informe narrativo."""
    cache_id = request.form.get("cache_id", "").strip()
    cache_entry = _download_cache.get(cache_id) if cache_id else None
    if not cache_entry:
        return render_template(
            "analisis.html",
            error="La sesión ha expirado. Vuelve a introducir el ticker.",
            modelo_seleccionado=MODELO_POR_DEFECTO,
        )

    datos = cache_entry["datos"]
    modelo = _normalizar_modelo(
        request.form.get("modelo", cache_entry.get("modelo", MODELO_POR_DEFECTO))
    )
    tipo_informe = request.form.get("tipo_informe", "nota_memoria").strip()
    enfoques_pred = request.form.getlist("enfoque_predefinido")
    enfoque_libre = request.form.get("enfoque_libre", "").strip()
    extension = request.form.get("extension", "estandar").strip()
    buscar_noticias = request.form.get("buscar_noticias") in ("1", "on", "true")
    instrucciones = request.form.get("instrucciones", "").strip()

    enfoque_partes = []
    if enfoques_pred:
        enfoque_partes.append("Áreas seleccionadas: " + ", ".join(enfoques_pred) + ".")
    if enfoque_libre:
        enfoque_partes.append(enfoque_libre)
    enfoque = "\n".join(enfoque_partes)

    try:
        nota = generar_nota_memoria(
            datos, modelo,
            instrucciones_usuario=instrucciones,
            permitir_busqueda=buscar_noticias,
            tipo_informe=tipo_informe,
            enfoque=enfoque,
            extension=extension,
        )
    except Exception as e:
        return render_template(
            "datos_yahoo.html",
            datos=datos,
            cache_id=cache_id,
            modelo_actual=modelo,
            formatear=formatear_numero,
            escala_para_estado=escala_para_estado,
            formatear_celda_estado=formatear_celda_estado,
            error=f"Error generando el informe: {e}",
        )

    secciones_html, pendientes = _render_secciones_para_resultado(nota)

    metricas_relevantes = construir_metricas_relevantes(datos, enfoques_pred)

    cache_entry.update({
        "nota": nota,
        "modelo": modelo,
        "instrucciones": instrucciones,
        "chat_historial": [],
        "tipo_informe": tipo_informe,
        "enfoque": enfoque,
        "extension": extension,
    })

    return render_template(
        "resultado.html",
        datos=datos,
        cache_id=cache_id,
        formatear=formatear_numero,
        graficos=preparar_datos_graficos(datos),
        secciones=secciones_html,
        pendientes=pendientes,
        chart_titulos=_CHART_TITULOS,
        modelo_actual=modelo,
        instrucciones_iniciales=instrucciones,
        busqueda_activa_inicial=buscar_noticias,
        metricas_relevantes=metricas_relevantes,
        meta_informe={
            "tipo": tipo_informe,
            "enfoques": enfoques_pred,
            "enfoque_libre": enfoque_libre,
            "extension": extension,
        },
    )


@app.route("/chat_memoria", methods=["POST"])
def chat_memoria():
    """Recibe un turno de chat y devuelve la memoria actualizada."""
    from flask import jsonify, render_template

    data = request.get_json(silent=True) or {}
    cache_id = (data.get("cache_id") or "").strip()
    mensaje = (data.get("mensaje") or "").strip()

    cached = _download_cache.get(cache_id)
    if not cached or "datos" not in cached:
        return jsonify({"error": "La sesión ha caducado. Genera la memoria de nuevo."}), 400
    if not mensaje:
        return jsonify({"error": "El mensaje no puede estar vacío."}), 400
    if len(mensaje) > 4000:
        return jsonify({"error": "El mensaje es demasiado largo (máx. 4000 caracteres)."}), 400

    modelo = _normalizar_modelo(data.get("modelo") or cached.get("modelo") or MODELO_POR_DEFECTO)
    buscar_noticias = bool(data.get("buscar_noticias"))
    datos = cached["datos"]
    memoria_actual = cached["nota"]
    historial = cached.setdefault("chat_historial", [])
    instrucciones_iniciales = cached.get("instrucciones", "")

    # Limitar historial a los últimos 10 turnos para contener tokens
    historial_acotado = historial[-10:]

    try:
        respuesta_chat, memoria_nueva, _citas = refinar_memoria_con_chat(
            datos, memoria_actual, historial_acotado, mensaje,
            instrucciones_iniciales, modelo,
            permitir_busqueda=buscar_noticias,
        )
    except Exception as e:
        return jsonify({"error": f"Error al consultar la IA: {str(e)}"}), 500

    # Actualiza el estado en cache
    historial.append({"role": "user", "content": mensaje})
    historial.append({"role": "assistant", "content": respuesta_chat})
    if memoria_nueva:
        cached["nota"] = memoria_nueva

    memoria_para_render = cached["nota"]
    secciones_html, pendientes = _render_secciones_para_resultado(memoria_para_render)

    memoria_fragmento = render_template(
        "_memoria_fragmento.html",
        datos=datos,
        secciones=secciones_html,
        pendientes=pendientes,
        chart_titulos=_CHART_TITULOS,
    )

    return jsonify({
        "chat_reply": respuesta_chat,
        "memoria_html": memoria_fragmento,
        "memoria_updated": memoria_nueva is not None,
    })


@app.route("/detectar_incoherencias", methods=["POST"])
def detectar_incoherencias():
    """Simula la conexión con Workiva y devuelve incoherencias inventadas
    entre la nota IA y la supuesta nota oficial del equipo financiero.
    """
    from flask import jsonify

    data = request.get_json(silent=True) or {}
    cache_id = (data.get("cache_id") or "").strip()
    cached = _download_cache.get(cache_id)
    if not cached or "datos" not in cached or "nota" not in cached:
        return jsonify({
            "error": "La sesión ha caducado. Genera la memoria de nuevo.",
        }), 400

    modelo = _normalizar_modelo(
        data.get("modelo") or cached.get("modelo") or MODELO_POR_DEFECTO
    )

    try:
        incoherencias = detectar_incoherencias_memoria(
            cached["datos"], cached["nota"], modelo,
        )
    except Exception as e:
        return jsonify({"error": f"Error consultando {SISTEMA_NOTA_OFICIAL}: {str(e)}"}), 500

    if not incoherencias:
        return jsonify({
            "sistema": SISTEMA_NOTA_OFICIAL,
            "incoherencias": [],
            "mensaje": (
                f"No se han detectado incoherencias relevantes entre la nota IA y "
                f"la versión oficial recuperada de {SISTEMA_NOTA_OFICIAL}."
            ),
        })

    return jsonify({
        "sistema": SISTEMA_NOTA_OFICIAL,
        "incoherencias": incoherencias,
    })


@app.route("/comparador", methods=["GET", "POST"])
def comparador():
    if request.method == "POST":
        tickers = []
        for i in range(1, 4):
            t = request.form.get(f"ticker{i}", "").strip().upper()
            if t:
                tickers.append(t)
        modelo = _normalizar_modelo(request.form.get("modelo", MODELO_POR_DEFECTO))

        if len(tickers) < 2:
            return render_template(
                "comparador.html",
                error="Introduce al menos 2 tickers para comparar.",
                modelo_seleccionado=modelo,
            )

        try:
            datos_empresas = []
            errores = []
            for t in tickers:
                datos = obtener_datos_financieros(t)
                if datos is None:
                    errores.append(f"No se encontraron datos para '{t}'")
                else:
                    datos_empresas.append(datos)

            if len(datos_empresas) < 2:
                return render_template(
                    "comparador.html",
                    error="No se pudieron obtener datos suficientes. " + "; ".join(errores),
                    modelo_seleccionado=modelo,
                )

            analisis = generar_analisis_comparativo(datos_empresas, modelo)
            analisis_html = markdown.markdown(analisis, extensions=["tables", "fenced_code", "sane_lists"])

            cache_id = str(uuid.uuid4())
            _download_cache[cache_id] = {
                "datos_empresas": datos_empresas,
                "analisis": analisis,
            }

            return render_template(
                "comparador_resultado.html",
                datos_empresas=datos_empresas,
                analisis=analisis_html,
                errores=errores,
                cache_id=cache_id,
                formatear=formatear_numero,
            )
        except Exception as e:
            return render_template(
                "comparador.html",
                error=f"Error: {str(e)}",
                modelo_seleccionado=modelo,
            )

    return render_template("comparador.html", modelo_seleccionado=MODELO_POR_DEFECTO)



def _decodificar_cabecera(h):
    """Decodifica cabeceras MIME tipo `=?utf-8?B?...?=` a string Unicode."""
    if not h:
        return ""
    from email.header import decode_header as _dh
    out = []
    for trozo, enc in _dh(h):
        if isinstance(trozo, bytes):
            try:
                out.append(trozo.decode(enc or "utf-8", errors="replace"))
            except (LookupError, UnicodeDecodeError):
                out.append(trozo.decode("utf-8", errors="replace"))
        else:
            out.append(trozo)
    return "".join(out)


def _limpiar_disclaimer_correo(cuerpo):
    """Corta los disclaimers legales del final del correo.

    Detecta el primer "punto de corte" típico (separadores largos de
    guiones, líneas de subrayados, o frases canónicas de cláusulas de
    confidencialidad) y devuelve sólo el texto anterior. Si no encuentra
    nada, devuelve el cuerpo intacto.
    """
    if not cuerpo:
        return cuerpo
    # Separadores: 8 o más guiones, em-dashes o subrayados consecutivos en
    # una línea. Cubre los típicos `-----`, `———`, `_____`.
    patrones = [
        r"\n[ \t]*[-–—]{8,}[ \t]*\n",
        r"\n[ \t]*_{8,}[ \t]*\n",
        r"\n[ \t]*\*{8,}[ \t]*\n",
        r"\n[ \t]*={8,}[ \t]*\n",
        # Frases inequívocas de disclaimer; cortamos justo antes.
        r"\n(?=[ \t]*(?:Este (?:mensaje|correo|e-?mail|email)"
        r"|AVISO LEGAL|AVISO DE CONFIDENCIALIDAD|CONFIDENCIALIDAD"
        r"|CL[ÁA]USULA DE CONFIDENCIALIDAD"
        r"|PROTECCI[ÓO]N DE DATOS"
        r"|This (?:message|e-?mail|email)"
        r"|CONFIDENTIAL(?:ITY)? NOTICE"
        r"|DISCLAIMER)[\s\S]{0,200}"
        r"(?:confidencial|destinatari|recipient|exclusiv|legal|privacy))",
    ]
    indices = []
    for p in patrones:
        m = re.search(p, cuerpo, flags=re.IGNORECASE)
        if m:
            indices.append(m.start())
    if not indices:
        return cuerpo
    corte = min(indices)
    return cuerpo[:corte].rstrip()


def _html_a_texto(html):
    """Conversión rudimentaria HTML → texto plano."""
    txt = re.sub(r"(?i)<br\s*/?>", "\n", html)
    txt = re.sub(r"(?i)</p>", "\n\n", txt)
    txt = re.sub(r"<[^>]+>", "", txt)
    txt = re.sub(r"\n{3,}", "\n\n", txt)
    return txt.strip()


_GMAIL_CACHE = {}

# Memoria en proceso de los correos ya respondidos durante esta sesión de la
# app. Clave estable: Message-ID del correo entrante (cambia entre apps, no
# entre re-aperturas del módulo). Si por algún motivo falta el Message-ID,
# se cae al id de IMAP. Se pierde al reiniciar la app, que es exactamente
# la semántica que pide el usuario ("memoria por sesión").
_IR_ENVIADOS = {}  # {clave: {timestamp, destinatario_nombre, destinatario_email, asunto, cuerpo}}
_IR_CONVERSACIONES = {}  # {clave: [{role, content, ts}, ...]}


def _clave_envio(correo):
    """Devuelve la clave estable para registrar un correo enviado."""
    return (correo.get("_message_id") or "").strip() or correo.get("id", "")


# --- Persistencia de memoria IR (capa 1 enviados + capa 2 conversaciones) -

def _dir_memoria_ir():
    if not TICKER_FIJO:
        return None
    d = _ROOT_DIR / "Empresas" / TICKER_FIJO / "memoria"
    d.mkdir(parents=True, exist_ok=True)
    return d


def _ruta_memoria_enviados():
    d = _dir_memoria_ir()
    return d / "enviados.json" if d else None


def _ruta_memoria_conversaciones():
    d = _dir_memoria_ir()
    return d / "conversaciones.json" if d else None


def _escribir_json_atomico(ruta, data):
    """Escribe JSON con un swap atómico (.tmp → rename) para evitar corrupciones."""
    import json as _json
    tmp = ruta.with_suffix(ruta.suffix + ".tmp")
    tmp.write_text(_json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    tmp.replace(ruta)


def _cargar_memoria_ir():
    """Carga enviados.json y conversaciones.json (si existen) a memoria."""
    import json as _json
    ruta_env = _ruta_memoria_enviados()
    if ruta_env and ruta_env.exists():
        try:
            data = _json.loads(ruta_env.read_text(encoding="utf-8"))
            if isinstance(data, dict):
                _IR_ENVIADOS.clear()
                _IR_ENVIADOS.update(data)
                _logging.getLogger("dfin.ir.memoria").info(
                    "Memoria de enviados cargada: %d correos", len(_IR_ENVIADOS),
                )
        except Exception as e:
            _logging.getLogger("dfin.ir.memoria").warning("enviados.json ilegible: %s", e)
    ruta_conv = _ruta_memoria_conversaciones()
    if ruta_conv and ruta_conv.exists():
        try:
            data = _json.loads(ruta_conv.read_text(encoding="utf-8"))
            if isinstance(data, dict):
                _IR_CONVERSACIONES.clear()
                _IR_CONVERSACIONES.update(data)
                _logging.getLogger("dfin.ir.memoria").info(
                    "Memoria de conversaciones cargada: %d hilos", len(_IR_CONVERSACIONES),
                )
        except Exception as e:
            _logging.getLogger("dfin.ir.memoria").warning("conversaciones.json ilegible: %s", e)


def _persistir_enviados():
    ruta = _ruta_memoria_enviados()
    if not ruta:
        return
    try:
        _escribir_json_atomico(ruta, _IR_ENVIADOS)
    except Exception as e:
        _logging.getLogger("dfin.ir.memoria").warning("No se pudo persistir enviados.json: %s", e)


def _persistir_conversaciones():
    ruta = _ruta_memoria_conversaciones()
    if not ruta:
        return
    try:
        _escribir_json_atomico(ruta, _IR_CONVERSACIONES)
    except Exception as e:
        _logging.getLogger("dfin.ir.memoria").warning("No se pudo persistir conversaciones.json: %s", e)


def _leer_bandeja_gmail(prefijo, horas_max=2, limite=50):
    """Lee la INBOX de Gmail por IMAP, filtra y devuelve correos.

    Filtros:
    - Asunto que EMPIEZA por `prefijo` (case-insensitive).
    - Recibidos en las últimas `horas_max` horas (medido en hora UTC con
      precisión real; IMAP solo permite filtrar por día, por lo que el
      filtro fino se aplica después).
    """
    log_bandeja = _logging.getLogger("dfin.ir.bandeja")
    if not GMAIL_USER or not GMAIL_APP_PASSWORD:
        log_bandeja.warning("Lectura Gmail abortada: faltan credenciales")
        raise RuntimeError(
            "Faltan GMAIL_USER y/o GMAIL_APP_PASSWORD en el archivo .env"
        )

    import imaplib
    import email as _emaillib
    from email.utils import parseaddr, parsedate_to_datetime
    from datetime import datetime as _dt, timezone as _tz, timedelta as _td
    import time as _time

    desde = _dt.now(_tz.utc) - _td(hours=horas_max)
    desde_imap_dia = desde.strftime("%d-%b-%Y")
    _t0 = _time.time()

    correos = []
    with imaplib.IMAP4_SSL("imap.gmail.com", 993) as M:
        M.login(GMAIL_USER, GMAIL_APP_PASSWORD)
        M.select("INBOX")
        criterio = f'(SUBJECT "{prefijo}" SINCE "{desde_imap_dia}")'
        status, data = M.search(None, criterio)
        if status != "OK" or not data or not data[0]:
            return []
        ids = data[0].split()
        for num in reversed(ids[-limite:]):
            status, msg_data = M.fetch(num, "(RFC822)")
            if status != "OK" or not msg_data or not msg_data[0]:
                continue
            msg = _emaillib.message_from_bytes(msg_data[0][1])

            # Fecha del correo
            try:
                fecha = parsedate_to_datetime(msg.get("Date", ""))
                if fecha.tzinfo is None:
                    fecha = fecha.replace(tzinfo=_tz.utc)
            except (TypeError, ValueError):
                fecha = None
            if fecha and fecha < desde:
                continue  # más viejo de lo permitido

            subject = _decodificar_cabecera(msg.get("Subject", ""))
            if not subject.upper().startswith(prefijo.upper()):
                continue

            from_nombre_raw, from_email = parseaddr(msg.get("From", ""))
            from_nombre = _decodificar_cabecera(from_nombre_raw) or from_email

            # Cuerpo: preferimos text/plain; fallback a HTML stripped.
            cuerpo = ""
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain" and "attachment" not in str(part.get("Content-Disposition", "")):
                        payload = part.get_payload(decode=True) or b""
                        cuerpo = payload.decode(part.get_content_charset() or "utf-8", errors="replace")
                        break
                if not cuerpo:
                    for part in msg.walk():
                        if part.get_content_type() == "text/html":
                            payload = part.get_payload(decode=True) or b""
                            html = payload.decode(part.get_content_charset() or "utf-8", errors="replace")
                            cuerpo = _html_a_texto(html)
                            break
            else:
                payload = msg.get_payload(decode=True)
                if payload:
                    cuerpo = payload.decode(msg.get_content_charset() or "utf-8", errors="replace")

            # Limpiamos el prefijo del asunto para no repetirlo en el "Re:".
            asunto_limpio = subject
            if asunto_limpio.upper().startswith(prefijo.upper()):
                asunto_limpio = asunto_limpio[len(prefijo):].lstrip(" :-—|")
            if not asunto_limpio:
                asunto_limpio = subject

            # El perfil del remitente lo clasificará la IA al pulsar
            # "Resolver con IA"; aquí dejamos un placeholder neutro.
            categoria = "Bandeja Gmail · Por clasificar"
            perfil = "Perfil del inversor pendiente de clasificación por IA"

            correos.append({
                "id": f"gmail-{num.decode()}",
                "remitente_nombre": (from_nombre or "").strip(),
                "remitente_email": (from_email or "").strip(),
                "remitente_perfil": perfil,
                "asunto": asunto_limpio.strip(),
                "categoria": categoria,
                "cuerpo": _limpiar_disclaimer_correo((cuerpo or "").strip()),
                "_message_id": msg.get("Message-ID", ""),
                "_fecha": fecha.strftime("%d/%m/%Y %H:%M") if fecha else "",
            })
    log_bandeja.info(
        "Gmail leído: %d correos coincidentes (prefijo='%s', %dh, %.2fs)",
        len(correos), prefijo, horas_max, _time.time() - _t0,
    )
    return correos


def _enviar_correo_smtp(destinatario_email, destinatario_nombre, asunto,
                        cuerpo, in_reply_to=None):
    """Envía un correo real vía SMTP de Gmail (SSL en :465)."""
    if not GMAIL_USER or not GMAIL_APP_PASSWORD:
        raise RuntimeError(
            "Faltan GMAIL_USER y/o GMAIL_APP_PASSWORD para enviar el correo."
        )
    import smtplib
    from email.mime.text import MIMEText
    from email.utils import formataddr, formatdate, make_msgid

    msg = MIMEText(cuerpo or "", "plain", "utf-8")
    remite = f"IR — {_BRANDING.get('empresa') or 'Equipo de IR'}"
    msg["From"] = formataddr((remite, GMAIL_USER))
    msg["To"] = formataddr((destinatario_nombre or "", destinatario_email))
    msg["Subject"] = asunto or ""
    msg["Date"] = formatdate(localtime=True)
    msg["Message-ID"] = make_msgid()
    if in_reply_to:
        msg["In-Reply-To"] = in_reply_to
        msg["References"] = in_reply_to

    log_smtp = _logging.getLogger("dfin.ir.enviar")
    import time as _t
    _t0 = _t.time()
    try:
        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
            server.login(GMAIL_USER, GMAIL_APP_PASSWORD)
            server.send_message(msg)
    except Exception as e:
        log_smtp.exception(
            "Fallo enviando a %s (asunto=%r): %s",
            destinatario_email, asunto, e,
        )
        raise
    log_smtp.info(
        "Correo enviado a %s (asunto=%r, %.2fs)",
        destinatario_email, (asunto or "")[:80], _t.time() - _t0,
    )



def _resolver_correo_ir(correo, datos, modelo=None):
    """Pide al modelo un borrador de respuesta para un correo de IR.

    Devuelve un dict con `puede_resolver`, `asunto_respuesta`, `cuerpo_respuesta`
    y `nota_interna`.
    """
    modelo = _normalizar_modelo(modelo or MODELO_POR_DEFECTO)
    empresa = _BRANDING.get("empresa") or TICKER_FIJO or "la compañía"
    ticker = TICKER_FIJO or ""
    tabla = _texto_datos_para_prompt(datos) if datos else "(no hay datos financieros disponibles en este momento)"

    prompt = f"""Eres el equipo de Relación con Inversores de {empresa} ({ticker}).
Acabas de recibir el siguiente correo electrónico de un inversor.

[CORREO RECIBIDO]
De: {correo['remitente_nombre']} <{correo['remitente_email']}>
Perfil del remitente: {correo['remitente_perfil']}
Asunto: {correo['asunto']}

Cuerpo:
{correo['cuerpo']}
[FIN DEL CORREO]

REGLAS DE OPERACIÓN (críticas, NO te desvíes):

PRIORIDAD ESTRICTA DE FUENTES (de mayor a menor confianza):
  A. LÍNEAS MAESTRAS DEL EQUIPO DE IR (si aparecen al inicio del prompt):
     son directrices internas no negociables sobre cómo responder. Tienen
     PRIORIDAD ABSOLUTA: si entran en conflicto con cualquier otra fuente,
     gana la línea maestra. Toda respuesta debe estar alineada con ellas.
  B. FRAGMENTOS DE DOCUMENTACIÓN CORPORATIVA (PDFs locales + crawl de la
     web corporativa): extractos literales de informes anuales, política
     de dividendos, presentaciones a inversores. Si la respuesta está
     aquí, úsala con preferencia sobre Yahoo y la web abierta.
  C. Datos de Yahoo Finance (cuenta de resultados, balance, cashflow).
  D. Búsqueda web abierta (solo para datos prospectivos/coyunturales que
     no estén en A, B ni C, o para confirmar lo encontrado).
La cita de fuentes va EN `nota_interna`, NUNCA en `cuerpo_respuesta`.

0. Si el cuerpo del correo termina con un separador de guiones, asteriscos o
   subrayados largos seguido de un aviso legal (frases tipo "Este mensaje va
   dirigido…", "AVISO LEGAL", "CONFIDENCIAL", "DISCLAIMER", "This message is
   intended only…"), IGNÓRALO POR COMPLETO: no es parte de la consulta, es
   la firma legal de la empresa del remitente. Trabaja únicamente con lo que
   viene ANTES de esa frontera.

1. Tu ÚNICA salida válida es el objeto JSON descrito al final. NUNCA escribas
   razonamiento, cálculos intermedios, "I'll search", "let me think", "Tengo
   todos los datos necesarios", "Procedo ahora", "CÁLCULOS INTERNOS", notas,
   comentarios o cualquier otro texto fuera del JSON. El primer carácter de
   tu salida debe ser '{{' y el último '}}'.

2. CLASIFICA al remitente como "minorista" (particular sin tecnicismos) o
   "mayorista" (profesional institucional con tono técnico).

3. Decide si puedes responder la consulta. ANTES de declararte incapaz,
   tienes que examinar a fondo los datos PÚBLICOS de Yahoo Finance que se
   adjuntan al final del prompt (los 3 estados completos: cuenta de
   resultados, balance y cashflow). Muchas consultas no piden un dato
   literal sino un RATIO o una DERIVADA que tienes que CALCULAR tú mismo
   sumando, restando o dividiendo líneas individuales del estado. Algunos
   ejemplos típicos (no exhaustivos):
     · Deuda Neta = "Deuda total" − "Efectivo y equivalentes" (o usar
       directamente "Deuda neta" si está disponible).
     · Deuda Neta / EBITDA = Deuda Neta / EBITDA del mismo ejercicio.
     · Flujo de Caja Libre (FCF) = "Flujo de caja operativo" + CAPEX
       (CAPEX viene con signo negativo, por eso se suma). Si "Flujo de
       caja libre" aparece como línea, úsala directamente.
     · Conversión EBITDA → FCF = FCF / EBITDA (en %).
     · CAPEX % de ingresos = |CAPEX| / Ingresos totales (en %).
     · ROE = Beneficio neto / Patrimonio neto.
     · ROA = Beneficio neto / Activos totales.
     · Crecimiento interanual = (valor actual − anterior) / anterior.
     · Margen bruto / operativo / neto = la línea correspondiente / Ingresos.
     · Pay-out = Dividendos pagados / Beneficio neto.
   SOLO si después de revisar a fondo los estados no puedes derivar la
   métrica pedida (porque genuinamente no está ni es derivable), entonces
   declararte incapaz.

   IMPORTANTE — DATOS FUTUROS / PROSPECTIVOS: si la consulta se refiere a
   guidance, previsiones, próximos resultados, próximas adquisiciones,
   estimaciones de analistas, cambios futuros en dividendo, próximo
   earnings call, calendario corporativo, o cualquier información sobre
   PERIODOS AÚN NO REPORTADOS, NO la busques en Yahoo Finance (esos datos
   históricos no la contienen). USA OBLIGATORIAMENTE la búsqueda web para
   localizar comunicados oficiales, notas de prensa o consensos de
   analistas recientes sobre {empresa}. Si pese a la búsqueda no
   encuentras información pública verificable, entonces declárate incapaz.

4. SI PUEDES responder con datos suficientes y verificables (literales o
   calculados a partir de los estados):
   - Redacta un BORRADOR PROFESIONAL adecuado al perfil clasificado.
   - Minorista: lenguaje claro y didáctico, evita tecnicismos.
   - Mayorista: tono técnico, conciso, cifras exactas con unidades.
   - Cifras concretas y exactas de los datos provistos, formato europeo
     (separador de miles ".", decimal ",").
   - Si la cifra es derivada, indica brevemente la fórmula entre paréntesis
     la primera vez que la usas, p. ej. "1,8x (Deuda Neta / EBITDA, con
     Deuda Neta = 56.950 M USD)".
   - El cuerpo es **íntegramente texto redactado en prosa**. PROHIBIDO usar
     tablas markdown, tablas ASCII, símbolos `|`, `---`, columnas alineadas
     con espacios, esquemas, bullets en formato de columnas, ni cualquier
     formato tabular: en un correo electrónico quedan mal y se rompen.
     Todo lo que quieras transmitir va en frases y, si necesitas enumerar,
     usa frases tipo "el ratio X fue 2,3 en 2024 frente a 1,9 en 2023" o
     listas con guion simple "- texto:" pero NUNCA columnas alineadas.
   - No incluyas listas de URLs, citas, "Fuentes consultadas" ni nada
     similar al final.

5. SI REALMENTE NO PUEDES responder con calidad (datos genuinamente
   internos / no públicos / no derivables tras revisar los estados):
   - Marca `puede_resolver` = false.
   - DEJA EL CUERPO COMPLETAMENTE VACÍO: `cuerpo_respuesta: ""`. No
     redactes un correo de excusa, no añadas saludo ni firma, no escribas
     [PENDIENTE: …]. El usuario humano redactará la respuesta desde cero.
   - Mantén el asunto como `"Re: <asunto original>"` por conveniencia.
   - En `nota_interna` explica brevemente por qué no se ha podido (1-2
     frases técnicas para el usuario interno, NO van al correo final).

6. CALIDAD: relee tu borrador antes de cerrar el JSON. Si detectas que has
   metido razonamiento, marcadores de cálculo, símbolos `**`, encabezados
   markdown `#`, código `` ` `` o tablas, vuelve a redactarlo en prosa limpia.

FORMATO DE SALIDA OBLIGATORIO — devuelve EXCLUSIVAMENTE un objeto JSON
válido, sin comentarios, sin texto antes ni después, sin acentos graves:

{{
  "tipo_inversor": "minorista",
  "perfil_estimado": "Frase breve describiendo al remitente y por qué.",
  "puede_resolver": true,
  "asunto_respuesta": "Re: {correo['asunto']}",
  "cuerpo_respuesta": "Estimado/a ...\\n\\n[texto en prosa, sin tablas ni listados tabulares]\\n\\nUn saludo,\\nEquipo de Relación con Inversores — {empresa}",
  "nota_interna": "1-2 frases para el usuario."
}}
"""

    # Líneas maestras del equipo de IR: prioridad absoluta.
    lineas_maestras = _cargar_lineas_maestras()

    # Recuperación de fragmentos del corpus RAG (PDFs locales + web crawl).
    # Si no hay índice o no hay matches, esta sección queda vacía.
    rag_chunks_texto = ""
    rag_resumen_fuentes = ""
    rag_idx = _RAG_INDEX.get("index")
    if rag_idx is not None and rag_idx.chunks:
        consulta_rag = f"{correo.get('asunto', '')} {correo.get('cuerpo', '')}"
        hits = rag_idx.buscar(consulta_rag, top_k=_rag.TOP_K_DEFECTO)
        if hits:
            rag_chunks_texto = _rag.formatear_chunks_para_prompt(hits)
            rag_resumen_fuentes = _rag.resumen_fuentes_para_nota_interna(hits)
            _logging.getLogger("dfin.rag").info(
                "RAG correo id=%s: %d chunks recuperados (top score=%.2f)",
                correo.get("id", "?"), len(hits), hits[0][0],
            )

    if lineas_maestras:
        prompt += (
            "\nLÍNEAS MAESTRAS DEL EQUIPO DE IR (PRIORIDAD ABSOLUTA — debes "
            "seguirlas siempre; si chocan con otra fuente, gana esto):\n"
            f"---\n{lineas_maestras}\n---\n"
        )
    if rag_chunks_texto:
        prompt += f"\n{rag_chunks_texto}\n"
    if datos:
        prompt += f"\nDatos financieros de Yahoo Finance disponibles:\n{tabla}\n"

    log_res = _logging.getLogger("dfin.ir.resolver")
    import time as _t
    _t0 = _t.time()
    log_res.info(
        "Resolviendo correo id=%s remitente=%s modelo=%s",
        correo.get("id", "?"), correo.get("remitente_email", "?"), modelo,
    )

    try:
        texto, _citas = llamar_ia_con_reintentos(
            prompt, max_tokens=2500, modelo=modelo, permitir_busqueda=True,
            modulo="ir.resolver",
        )
    except Exception as e:
        log_res.exception("Fallo en LLM al resolver id=%s: %s", correo.get("id", "?"), e)
        raise

    _latencia = _t.time() - _t0
    import json as _json
    bloque = re.search(r"\{.*\}", texto or "", re.DOTALL)
    if not bloque:
        log_res.warning(
            "Sin JSON válido para id=%s (modelo=%s, %.2fs)",
            correo.get("id", "?"), modelo, _latencia,
        )
        return {
            "puede_resolver": False,
            "asunto_respuesta": f"Re: {correo['asunto']}",
            "cuerpo_respuesta": (texto or "").strip() or "[La IA no devolvió respuesta]",
            "nota_interna": "La IA no devolvió JSON válido; se muestra el texto bruto.",
        }
    try:
        parsed = _json.loads(bloque.group(0), strict=False)
    except _json.JSONDecodeError:
        return {
            "puede_resolver": False,
            "asunto_respuesta": f"Re: {correo['asunto']}",
            "cuerpo_respuesta": bloque.group(0),
            "nota_interna": "JSON no parseable; se muestra el bloque tal cual.",
        }
    # Saneo defensivo
    parsed.setdefault("puede_resolver", False)
    parsed.setdefault("asunto_respuesta", f"Re: {correo['asunto']}")
    parsed.setdefault("cuerpo_respuesta", "")
    parsed.setdefault("nota_interna", "")
    tipo = (parsed.get("tipo_inversor") or "").lower().strip()
    if tipo not in ("minorista", "mayorista"):
        tipo = "mayorista"
    parsed["tipo_inversor"] = tipo
    parsed.setdefault("perfil_estimado", "")
    parsed = _validar_borrador_ir(parsed, correo)

    # Anexamos las fuentes RAG a la nota interna (solo se muestran al
    # usuario en la UI, NUNCA viajan al correo de respuesta).
    if rag_resumen_fuentes:
        nota_previa = (parsed.get("nota_interna") or "").strip()
        bloque_fuentes = "Fuentes consultadas (documentación corporativa):\n" + rag_resumen_fuentes
        parsed["nota_interna"] = (nota_previa + "\n\n" + bloque_fuentes).strip() if nota_previa else bloque_fuentes

    log_res.info(
        "Correo id=%s resuelto: puede_resolver=%s tipo=%s modelo=%s lat=%.2fs rag_chunks=%d",
        correo.get("id", "?"),
        parsed.get("puede_resolver"),
        parsed.get("tipo_inversor"),
        modelo, _latencia,
        len(rag_resumen_fuentes.splitlines()) if rag_resumen_fuentes else 0,
    )
    return parsed


# Patrones que delatan "fugas" de razonamiento del modelo en el cuerpo del
# correo. Si aparecen, el borrador no es publicable y se degrada a
# `puede_resolver=false` con un mensaje claro.
_PATRONES_RAZONAMIENTO = [
    r"(?:^|\s)(I'll|I will|Let me|I need to|I'm going to|I should)\s",
    r"\b(search for|let me search|search the web|let me calculate|I'll calculate)\b",
    r"\b(C[ÁA]LCULOS INTERNOS|REASONING|CHAIN OF THOUGHT)\b",
    r"\bTengo todos los datos\b",
    r"\bProcedo ahora a\b",
    r"\bvoy a (calcular|buscar|verificar|consultar)\b",
    r"\baqu[íi] est[áa] el JSON\b",
    r"\bd[ée]jame (calcular|buscar|verificar|pensar)\b",
]
_PATRONES_TABLA = [
    r"\n\s*\|[^\n]+\|",         # filas markdown tipo `| a | b |`
    r"\n\s*-{3,}\s*\|",         # separadores `---|`
    r"\n\s*\|?\s*:?-{3,}:?",    # separadores tipo `:---:`
]


def _validar_borrador_ir(parsed, correo):
    """Sanea el borrador y deja el cuerpo vacío cuando no es publicable.

    Reglas:
    - Si el cuerpo contiene fugas de razonamiento (`I'll search`,
      `CÁLCULOS INTERNOS`, etc.) o tablas markdown/ASCII, se descarta:
      `puede_resolver=false` y `cuerpo_respuesta=""`.
    - Si la IA marcó `puede_resolver=false` por su cuenta, el cuerpo
      también se fuerza a vacío (no queremos correos de "no puedo
      responderle"; el usuario humano redacta desde cero).
    En ambos casos se mantiene el asunto `Re: <original>` y se anota la
    razón en `nota_interna` para que el frontend pueda mostrar el badge
    correcto.
    """
    cuerpo = parsed.get("cuerpo_respuesta") or ""
    razones = []
    for patron in _PATRONES_RAZONAMIENTO:
        if re.search(patron, cuerpo, flags=re.IGNORECASE):
            razones.append("razonamiento expuesto en la respuesta")
            break
    for patron in _PATRONES_TABLA:
        if re.search(patron, cuerpo):
            razones.append("tabla/columna detectada en la respuesta")
            break

    if razones or not parsed.get("puede_resolver"):
        if razones:
            _logging.getLogger("dfin.ir.validador").warning(
                "Borrador descartado para id=%s: %s",
                correo.get("id", "?"), "; ".join(razones),
            )
        parsed["puede_resolver"] = False
        parsed["cuerpo_respuesta"] = ""
        if not parsed.get("asunto_respuesta"):
            parsed["asunto_respuesta"] = f"Re: {correo.get('asunto', '')}"
        if razones:
            parsed["nota_interna"] = (
                "Borrador descartado tras validación: " + ", ".join(razones)
                + ". Redáctalo manualmente."
            )
        else:
            # Mantenemos la nota_interna que la IA ya hubiera puesto, o una
            # genérica si no había.
            if not parsed.get("nota_interna"):
                parsed["nota_interna"] = (
                    "La IA no encontró información suficiente en los datos "
                    "disponibles para redactar una respuesta con calidad. "
                    "Redacta manualmente."
                )
    return parsed


@app.route("/relacion-inversores")
def relacion_inversores():
    """Bandeja en vivo: lee la INBOX de Gmail y muestra los correos del ticker."""
    if not GMAIL_USER or not GMAIL_APP_PASSWORD:
        return render_template(
            "relacion_inversores_no_config.html",
            prefijo=_prefijo_asunto_actual(),
        )
    try:
        correos = _leer_bandeja_gmail(_prefijo_asunto_actual(), horas_max=2)
    except Exception as e:
        return render_template(
            "relacion_inversores_no_config.html",
            prefijo=_prefijo_asunto_actual(),
            gmail_user=GMAIL_USER,
            error=f"No se pudo leer la bandeja de Gmail: {e}",
        )

    # Anotar cada correo con su estado de envío en esta sesión.
    for c in correos:
        c["_enviado"] = _IR_ENVIADOS.get(_clave_envio(c))  # dict o None

    _GMAIL_CACHE.clear()
    _GMAIL_CACHE.update({c["id"]: c for c in correos})
    # Refrescamos el contador (pendientes = no enviados aún en esta sesión).
    import time
    pendientes = sum(1 for c in correos if not c["_enviado"])
    _IR_PENDIENTES_CACHE.update({"ts": time.time(), "count": pendientes})

    return render_template(
        "relacion_inversores.html",
        correos=correos,
        modelo_actual=MODELO_POR_DEFECTO,
        fuente="gmail",
        prefijo=_prefijo_asunto_actual(),
        gmail_user=GMAIL_USER,
    )


@app.route("/relacion-inversores/resolver", methods=["POST"])
def relacion_inversores_resolver():
    from flask import jsonify
    data = request.get_json(silent=True) or {}
    correo_id = (data.get("correo_id") or "").strip()
    modelo = _normalizar_modelo(data.get("modelo") or MODELO_POR_DEFECTO)

    correo = _GMAIL_CACHE.get(correo_id)
    if not correo:
        return jsonify({"ok": False, "error": "Correo no encontrado. Refresca la bandeja."}), 404

    datos = _cargar_datos_ticker_fijo() if TICKER_FIJO else None
    try:
        resultado = _resolver_correo_ir(correo, datos, modelo=modelo)
    except Exception as e:
        return jsonify({"ok": False, "error": f"Error al resolver: {e}"}), 500

    return jsonify({
        "ok": True,
        "correo_id": correo_id,
        "destinatario_nombre": correo["remitente_nombre"],
        "destinatario_email": correo["remitente_email"],
        **resultado,
    })


def _refinar_correo_ir(correo, datos, borrador_actual, historial, mensaje_usuario, modelo=None):
    """Refina un borrador de respuesta a un correo de IR mediante chat.

    El usuario puede pedir cambios (tono, longitud, cifras, datos extra…) o
    hacer preguntas sin tocar el borrador. La IA devuelve una respuesta
    conversacional + el borrador (modificado o no).
    """
    modelo = _normalizar_modelo(modelo or MODELO_POR_DEFECTO)
    empresa = _BRANDING.get("empresa") or TICKER_FIJO or "la compañía"
    ticker = TICKER_FIJO or ""
    tabla = _texto_datos_para_prompt(datos) if datos else "(no hay datos financieros disponibles)"

    historial_txt = ""
    for turno in (historial or [])[-12:]:
        rol = "Usuario" if (turno.get("role") == "user") else "Asistente"
        historial_txt += f"\n{rol}: {(turno.get('content') or '').strip()}"

    prompt = f"""Eres el equipo de Relación con Inversores de {empresa} ({ticker}).
Estás revisando con el usuario el borrador de respuesta a un correo de un
inversor. El usuario te pedirá cambios o te preguntará cosas; aplica las
modificaciones necesarias al borrador.

[CORREO ORIGINAL DEL INVERSOR]
De: {correo['remitente_nombre']} <{correo['remitente_email']}>
Asunto: {correo['asunto']}
Cuerpo:
{correo['cuerpo']}
[FIN CORREO ORIGINAL]

[BORRADOR ACTUAL EN EDICIÓN]
Asunto: {borrador_actual.get('asunto', '')}
Cuerpo:
{borrador_actual.get('cuerpo', '')}
[FIN BORRADOR]
{('[HISTORIAL DEL CHAT]' + historial_txt + chr(10) + '[FIN HISTORIAL]') if historial_txt else ''}

NUEVO MENSAJE DEL USUARIO:
{mensaje_usuario}

Reglas:
- Si el usuario te pide CAMBIOS, devuelve el borrador con los cambios
  aplicados y `modificado` = true. Mantén el resto del borrador intacto.
- Si el usuario te HACE UNA PREGUNTA sin pedir cambios, no toques el
  borrador (devuélvelo tal cual) y pon `modificado` = false.
- Usa cifras exactas de los datos de Yahoo (formato europeo: miles ".",
  decimal ",") cuando el usuario te pida añadir datos.
- No incluyas nunca listas de URLs ni "Fuentes consultadas" en el cuerpo.
- `respuesta_chat`: 1-2 frases describiendo qué hiciste o respondiendo a
  la pregunta. NO repitas el cuerpo completo del correo aquí.

FORMATO DE SALIDA OBLIGATORIO — devuelve EXCLUSIVAMENTE un objeto JSON
válido, sin comentarios ni acentos graves:

{{
  "respuesta_chat": "1-2 frases para el usuario.",
  "asunto_respuesta": "{borrador_actual.get('asunto', '')}",
  "cuerpo_respuesta": "(borrador completo actualizado, o el actual si no hubo cambios)",
  "modificado": true
}}
"""

    lineas_maestras = _cargar_lineas_maestras()
    if lineas_maestras:
        prompt += (
            "\nLÍNEAS MAESTRAS DEL EQUIPO DE IR (PRIORIDAD ABSOLUTA):\n"
            f"---\n{lineas_maestras}\n---\n"
        )
    if datos:
        prompt += f"\nDatos financieros de Yahoo Finance disponibles:\n{tabla}\n"

    log_ref = _logging.getLogger("dfin.ir.refinar")
    import time as _t
    _t0 = _t.time()
    log_ref.info(
        "Refinando correo id=%s modelo=%s msg=%r",
        correo.get("id", "?"), modelo, (mensaje_usuario or "")[:120],
    )

    try:
        texto, _citas = llamar_ia_con_reintentos(
            prompt, max_tokens=2500, modelo=modelo, permitir_busqueda=True,
            modulo="ir.refinar",
        )
    except Exception as e:
        log_ref.exception("Fallo en LLM al refinar id=%s: %s", correo.get("id", "?"), e)
        raise

    _lat = _t.time() - _t0
    import json as _json
    bloque = re.search(r"\{.*\}", texto or "", re.DOTALL)
    if not bloque:
        log_ref.warning("Sin JSON válido refinando id=%s (lat=%.2fs)", correo.get("id", "?"), _lat)
        return {
            "respuesta_chat": (texto or "").strip()[:600] or "(sin respuesta)",
            "asunto_respuesta": borrador_actual.get("asunto", ""),
            "cuerpo_respuesta": borrador_actual.get("cuerpo", ""),
            "modificado": False,
        }
    try:
        parsed = _json.loads(bloque.group(0), strict=False)
    except _json.JSONDecodeError:
        return {
            "respuesta_chat": "No pude procesar la respuesta del modelo. Vuelve a intentarlo.",
            "asunto_respuesta": borrador_actual.get("asunto", ""),
            "cuerpo_respuesta": borrador_actual.get("cuerpo", ""),
            "modificado": False,
        }
    parsed.setdefault("respuesta_chat", "")
    parsed.setdefault("asunto_respuesta", borrador_actual.get("asunto", ""))
    parsed.setdefault("cuerpo_respuesta", borrador_actual.get("cuerpo", ""))
    parsed.setdefault("modificado", False)
    # Reusa la misma validación que la primera redacción: si la IA cuela
    # razonamiento o tablas al refinar, también se descarta.
    pseudo = {
        "puede_resolver": True,
        "asunto_respuesta": parsed["asunto_respuesta"],
        "cuerpo_respuesta": parsed["cuerpo_respuesta"],
        "nota_interna": "",
    }
    pseudo = _validar_borrador_ir(pseudo, correo)
    if not pseudo["puede_resolver"]:
        parsed["cuerpo_respuesta"] = pseudo["cuerpo_respuesta"]
        parsed["respuesta_chat"] = (
            "He intentado aplicar tus cambios pero la respuesta contenía contenido "
            "no apto (razonamiento expuesto o tablas). He marcado el borrador como "
            "pendiente para revisión manual."
        )
        parsed["modificado"] = True
    log_ref.info(
        "Refinado id=%s modificado=%s modelo=%s lat=%.2fs",
        correo.get("id", "?"), parsed.get("modificado"), modelo, _lat,
    )
    return parsed


@app.route("/relacion-inversores/lineas-maestras", methods=["GET", "POST"])
def relacion_inversores_lineas_maestras():
    """GET → devuelve el contenido actual. POST → guarda el contenido."""
    from flask import jsonify
    if not TICKER_FIJO:
        return jsonify({"ok": False, "error": "Hace falta arrancar la app con un ticker."}), 400
    if request.method == "POST":
        data = request.get_json(silent=True) or {}
        contenido = data.get("contenido", "")
        if not isinstance(contenido, str):
            return jsonify({"ok": False, "error": "Contenido inválido."}), 400
        try:
            _guardar_lineas_maestras(contenido)
        except ValueError as e:
            return jsonify({"ok": False, "error": str(e)}), 400
        except Exception as e:
            return jsonify({"ok": False, "error": f"Error al guardar: {e}"}), 500
        return jsonify({
            "ok": True,
            "contenido": _cargar_lineas_maestras(),
            "guardado_en": str(_ruta_lineas_maestras()),
        })
    # GET
    return jsonify({
        "ok": True,
        "contenido": _cargar_lineas_maestras(),
        "ruta": str(_ruta_lineas_maestras() or ""),
        "existe": bool(_ruta_lineas_maestras() and _ruta_lineas_maestras().exists()),
    })


@app.route("/relacion-inversores/refinar", methods=["POST"])
def relacion_inversores_refinar():
    from flask import jsonify
    data = request.get_json(silent=True) or {}
    correo_id = (data.get("correo_id") or "").strip()
    modelo = _normalizar_modelo(data.get("modelo") or MODELO_POR_DEFECTO)
    mensaje = (data.get("mensaje") or "").strip()
    historial = data.get("historial") or []
    borrador_actual = {
        "asunto": data.get("asunto_actual") or "",
        "cuerpo": data.get("cuerpo_actual") or "",
    }
    if not mensaje:
        return jsonify({"ok": False, "error": "Mensaje vacío"}), 400

    correo = _GMAIL_CACHE.get(correo_id)
    if not correo:
        return jsonify({"ok": False, "error": "Correo no encontrado. Refresca la bandeja."}), 404

    datos = _cargar_datos_ticker_fijo() if TICKER_FIJO else None
    try:
        res = _refinar_correo_ir(correo, datos, borrador_actual, historial, mensaje, modelo=modelo)
    except Exception as e:
        return jsonify({"ok": False, "error": f"Error al refinar: {e}"}), 500

    # Persistir conversación (capa 2 de memoria).
    from datetime import datetime as _dt
    clave_conv = _clave_envio(correo)
    if clave_conv:
        ts = _dt.now().strftime("%Y-%m-%d %H:%M:%S")
        hilo = list(historial) + [
            {"role": "user", "content": mensaje, "ts": ts},
            {"role": "assistant", "content": res.get("respuesta_chat", ""), "ts": ts,
             "modificado": bool(res.get("modificado"))},
        ]
        _IR_CONVERSACIONES[clave_conv] = hilo
        _persistir_conversaciones()

    return jsonify({"ok": True, **res})


@app.route("/relacion-inversores/conversaciones", methods=["GET"])
def relacion_inversores_conversaciones():
    """Devuelve el historial de refinamiento persistido para un correo."""
    from flask import jsonify
    correo_id = (request.args.get("correo_id") or "").strip()
    if not correo_id:
        return jsonify({"ok": True, "historial": []})
    correo = _GMAIL_CACHE.get(correo_id) or {}
    clave = _clave_envio(correo) or correo_id
    hilo = _IR_CONVERSACIONES.get(clave) or []
    return jsonify({"ok": True, "historial": hilo})


@app.route("/relacion-inversores/enviar", methods=["POST"])
def relacion_inversores_enviar():
    from flask import jsonify
    from datetime import datetime as _dt
    data = request.get_json(silent=True) or {}
    destinatario_nombre = data.get("destinatario_nombre", "")
    destinatario_email = data.get("destinatario_email", "")
    asunto = data.get("asunto", "")
    cuerpo = data.get("cuerpo", "")
    correo_orig = _GMAIL_CACHE.get(data.get("correo_id") or "") or {}
    try:
        _enviar_correo_smtp(
            destinatario_email=destinatario_email,
            destinatario_nombre=destinatario_nombre,
            asunto=asunto,
            cuerpo=cuerpo,
            in_reply_to=correo_orig.get("_message_id"),
        )
    except Exception as e:
        return jsonify({"ok": False, "error": f"Error al enviar: {e}"}), 500

    timestamp = _dt.now().strftime("%d/%m/%Y %H:%M")
    # Registramos el envío en memoria de sesión.
    clave = _clave_envio(correo_orig) if correo_orig else (data.get("correo_id") or "")
    if clave:
        _IR_ENVIADOS[clave] = {
            "timestamp": timestamp,
            "destinatario_nombre": destinatario_nombre,
            "destinatario_email": destinatario_email,
            "asunto": asunto,
            "cuerpo": cuerpo,
        }
        _persistir_enviados()
        # Invalidamos el cache del contador para que la home se entere ya.
        _IR_PENDIENTES_CACHE.update({"ts": 0.0, "count": None})

    return jsonify({
        "ok": True,
        "timestamp": timestamp,
        "destinatario_nombre": destinatario_nombre,
        "destinatario_email": destinatario_email,
        "asunto": asunto,
    })


def _dir_descargas_ticker():
    if not TICKER_FIJO:
        return None
    d = _ROOT_DIR / "Empresas" / TICKER_FIJO / "descargas"
    d.mkdir(parents=True, exist_ok=True)
    return d


def _guardar_descarga_en_empresa(buffer, datos, extension):
    """Guarda una copia del documento generado en Empresas/{TICKER}/descargas/.

    Devuelve la ruta absoluta del fichero guardado (o None si no hay
    TICKER_FIJO o falla la escritura). El buffer se rebobina al final.
    """
    dst = _dir_descargas_ticker()
    if not dst:
        return None
    from datetime import datetime as _dt
    ticker = (datos or {}).get("ticker") or TICKER_FIJO or "informe"
    ts = _dt.now().strftime("%Y-%m-%d_%H-%M-%S")
    nombre = f"{ts}_{ticker}_informe.{extension}"
    ruta = dst / nombre
    try:
        buffer.seek(0)
        ruta.write_bytes(buffer.read())
        buffer.seek(0)
        _logging.getLogger("dfin.informe").info("Copia guardada en %s", ruta)
        return ruta
    except Exception as e:
        _logging.getLogger("dfin.informe").warning("No se pudo guardar copia local: %s", e)
        return None


@app.route("/descargar_word", methods=["POST"])
def descargar_word():
    cache_id = request.form.get("cache_id", "")
    cached = _download_cache.get(cache_id)
    if not cached:
        return "Sesión expirada. Por favor, realiza el análisis de nuevo.", 400

    datos = cached["datos"]
    nota = cached["nota"]
    try:
        buffer = crear_documento_word(datos, nota)
        _guardar_descarga_en_empresa(buffer, datos, "docx")
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"informe_{datos['ticker']}.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        return f"Error generando documento: {str(e)}", 500


@app.route("/descargar_comparativo", methods=["POST"])
def descargar_comparativo():
    cache_id = request.form.get("cache_id", "")
    cached = _download_cache.get(cache_id)
    if not cached:
        return "Sesión expirada. Por favor, realiza la comparación de nuevo.", 400

    datos_empresas = cached["datos_empresas"]
    analisis = cached["analisis"]
    try:
        buffer = crear_documento_comparativo(datos_empresas, analisis)
        nombre = "_vs_".join([d["ticker"] for d in datos_empresas])
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"comparativo_{nombre}.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        return f"Error generando documento: {str(e)}", 500


@app.route("/descargar_pdf", methods=["POST"])
def descargar_pdf():
    cache_id = request.form.get("cache_id", "")
    cached = _download_cache.get(cache_id)
    if not cached or "datos" not in cached:
        return "Sesión expirada. Por favor, realiza el análisis de nuevo.", 400

    datos = cached["datos"]
    nota = cached["nota"]
    try:
        buffer = crear_documento_pdf(datos, nota)
        _guardar_descarga_en_empresa(buffer, datos, "pdf")
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"informe_{datos['ticker']}.pdf",
            mimetype="application/pdf",
        )
    except Exception as e:
        return f"Error generando documento: {str(e)}", 500


@app.route("/informes-generados")
def informes_generados():
    """Lista los informes Word/PDF guardados en Empresas/{TICKER}/descargas/."""
    d = _dir_descargas_ticker()
    return render_template(
        "informes_generados.html",
        items=_listar_informes_descargados(),
        carpeta=str(d) if d else "",
    )


@app.route("/descargar-archivado/<nombre>")
def descargar_archivado(nombre):
    """Sirve un fichero ya generado y archivado en la carpeta del ticker."""
    from flask import send_from_directory, abort
    import re as _re
    d = _dir_descargas_ticker()
    if not d or not _re.match(r"^[\w.\-]+$", nombre):
        abort(404)
    ruta = d / nombre
    if not ruta.exists() or not ruta.is_file():
        abort(404)
    return send_from_directory(str(d), nombre, as_attachment=True)


def _parsear_lineas_coste(limite_lineas=200):
    """Parsea las últimas N líneas de coste_tokens.txt en dicts.

    Devuelve: lista de {timestamp, modulo, modelo, input_tokens, output_tokens,
    coste_usd, acumulado_usd, llamada_n}, ordenadas de más reciente a más
    antigua, y agregados por módulo y por modelo.
    """
    import re as _re
    if not _COSTE_FILE.exists():
        return {"llamadas": [], "por_modulo": {}, "por_modelo": {},
                "total_usd": 0.0, "total_llamadas": 0}
    lineas = _COSTE_FILE.read_text(encoding="utf-8").splitlines()
    parseadas = []
    for linea in lineas:
        m = _re.match(r"^\[(?P<ts>[^\]]+)\]\s+(?P<rest>.+)$", linea)
        if not m:
            continue
        partes = dict(
            p.split("=", 1)
            for p in m.group("rest").split()
            if "=" in p
        )
        try:
            parseadas.append({
                "timestamp": m.group("ts"),
                "modulo": partes.get("modulo", ""),
                "modelo": partes.get("modelo", ""),
                "input_tokens": int(partes.get("input_tokens", 0)),
                "output_tokens": int(partes.get("output_tokens", 0)),
                "coste_usd": float(partes.get("coste_usd", 0)),
                "acumulado_usd": float(partes.get("acumulado_usd", 0)),
                "llamada_n": int(partes.get("llamada_n", 0)),
            })
        except (ValueError, KeyError):
            continue
    por_modulo = {}
    por_modelo = {}
    for x in parseadas:
        m_ = por_modulo.setdefault(x["modulo"], {"usd": 0.0, "llamadas": 0, "in": 0, "out": 0})
        m_["usd"] += x["coste_usd"]
        m_["llamadas"] += 1
        m_["in"] += x["input_tokens"]
        m_["out"] += x["output_tokens"]
        d_ = por_modelo.setdefault(x["modelo"], {"usd": 0.0, "llamadas": 0, "in": 0, "out": 0})
        d_["usd"] += x["coste_usd"]
        d_["llamadas"] += 1
        d_["in"] += x["input_tokens"]
        d_["out"] += x["output_tokens"]
    total_usd = parseadas[-1]["acumulado_usd"] if parseadas else 0.0
    return {
        "llamadas": list(reversed(parseadas[-limite_lineas:])),
        "por_modulo": por_modulo,
        "por_modelo": por_modelo,
        "total_usd": total_usd,
        "total_llamadas": len(parseadas),
    }


@app.route("/coste")
def coste_tokens_panel():
    info = _parsear_lineas_coste()
    return render_template(
        "coste.html",
        info=info,
        fichero=str(_COSTE_FILE),
        fichero_existe=_COSTE_FILE.exists(),
    )


@app.route("/coste/descargar")
def coste_tokens_descargar():
    from flask import send_file, abort
    if not _COSTE_FILE.exists():
        abort(404)
    return send_file(
        str(_COSTE_FILE),
        as_attachment=True,
        download_name=f"coste_tokens_{TICKER_FIJO or 'global'}.txt",
        mimetype="text/plain",
    )


# --- Funcionalidades mock (navegación completa, sin lógica productiva) ---

@app.route("/presupuesto")
def presupuesto():
    return render_template("presupuesto.html")


@app.route("/conciliacion")
def conciliacion():
    return render_template("conciliacion.html")


@app.route("/radar-normativo")
def radar_normativo():
    alertas = [
        {
            "categoria": "Fiscal",
            "tag": "tag-red",
            "color": "#DA291C",
            "jurisdiccion": "España",
            "impacto": "Alto",
            "fecha": "18/04/2026",
            "fuente": "BOE",
            "titulo": "Nueva obligación de reporte trimestral para pagos internacionales",
            "resumen": (
                "La AEAT publica la Orden HAC/392/2026 que amplía el Modelo 349 e introduce "
                "un reporte trimestral obligatorio para operaciones de servicios intragrupo "
                "por encima de 50.000 €. Entrada en vigor: 1 de julio de 2026."
            ),
        },
        {
            "categoria": "Contable",
            "tag": "tag-blue",
            "color": "#00A3E0",
            "jurisdiccion": "Unión Europea",
            "impacto": "Medio",
            "fecha": "15/04/2026",
            "fuente": "EFRAG / DOUE",
            "titulo": "Revisión de NIIF 18 — presentación y desglose de estados financieros",
            "resumen": (
                "La EFRAG emite dictamen favorable a la adopción de NIIF 18, que redefine "
                "subtotales obligatorios (operativo, financiación, inversión) y amplía las "
                "medidas definidas por la dirección. Primera aplicación: ejercicios 2027."
            ),
        },
        {
            "categoria": "Laboral",
            "tag": "tag-orange",
            "color": "#ED8B00",
            "jurisdiccion": "España",
            "impacto": "Medio",
            "fecha": "11/04/2026",
            "fuente": "BOE",
            "titulo": "Actualización del SMI y cotizaciones sociales para 2026",
            "resumen": (
                "Real Decreto 298/2026 que fija el SMI en 1.230 €/mes y actualiza los "
                "topes máximos de cotización. Revisar masa salarial y provisiones laborales."
            ),
        },
        {
            "categoria": "ESG",
            "tag": "tag-green",
            "color": "#86BC25",
            "jurisdiccion": "Unión Europea",
            "impacto": "Alto",
            "fecha": "08/04/2026",
            "fuente": "DOUE",
            "titulo": "CSRD — publicación de las NEIS sector specific para el ejercicio 2026",
            "resumen": (
                "La Comisión Europea publica los estándares sectoriales de la CSRD, que "
                "serán de aplicación obligatoria a partir del ejercicio 2026 para empresas "
                "cotizadas del sector financiero y energético."
            ),
        },
        {
            "categoria": "Financiera",
            "tag": "tag-black",
            "color": "#000000",
            "jurisdiccion": "Unión Europea",
            "impacto": "Bajo",
            "fecha": "03/04/2026",
            "fuente": "BCE",
            "titulo": "Guía técnica sobre el uso de LLMs en procesos KYC y AML",
            "resumen": (
                "El BCE publica principios supervisores para la utilización de modelos de "
                "lenguaje en procesos antiblanqueo, con foco en explicabilidad, trazabilidad "
                "y supervisión humana sobre las decisiones automatizadas."
            ),
        },
    ]
    return render_template("radar_normativo.html", alertas=alertas)


@app.route("/noticias")
def noticias():
    noticias_lista = [
        {
            "categoria": "M&A", "tag": "tag-green", "banner": "#43B02A",
            "fecha": "20/04/2026", "fuente": "Reuters",
            "titular": "BlackRock lidera el consorcio que adquiere Klarna por 14,2B USD",
            "resumen": "La operación convierte a la fintech sueca en la mayor OPA del sector BNPL en la última década.",
        },
        {
            "categoria": "Regulación", "tag": "tag-blue", "banner": "#00A3E0",
            "fecha": "19/04/2026", "fuente": "Bloomberg",
            "titular": "La EBA publica la guía definitiva sobre el uso de IA en procesos KYC",
            "resumen": "El marco europeo establece requisitos de auditabilidad, explicabilidad y supervisión humana sobre modelos generativos.",
        },
        {
            "categoria": "Mercados", "tag": "tag-black", "banner": "#000000",
            "fecha": "19/04/2026", "fuente": "Financial Times",
            "titular": "Las fintechs europeas cierran la semana con un rally del 4,2%",
            "resumen": "Adyen, Nu Holdings y Wise lideran las subidas tras los resultados del primer trimestre por encima del consenso.",
        },
        {
            "categoria": "Tecnología", "tag": "tag-orange", "banner": "#ED8B00",
            "fecha": "18/04/2026", "fuente": "Expansión",
            "titular": "CaixaBank firma un acuerdo con Anthropic para generalizar Claude en su banca privada",
            "resumen": "El despliegue cubrirá 3.200 gestores y se integrará con la plataforma propia de asesoramiento del banco.",
        },
        {
            "categoria": "Corporate", "tag": "tag-green", "banner": "#86BC25",
            "fecha": "17/04/2026", "fuente": "Cinco Días",
            "titular": "Endesa presenta un plan ESG de 4,5B € hasta 2030 centrado en hidrógeno verde",
            "resumen": "La compañía prevé triplicar su capacidad instalada renovable y emitir bonos verdes por 1,2B € en 2026.",
        },
        {
            "categoria": "M&A", "tag": "tag-green", "banner": "#43B02A",
            "fecha": "16/04/2026", "fuente": "Reuters",
            "titular": "Telefónica y Orange exploran una fusión de sus redes móviles en España",
            "resumen": "Fuentes próximas a la negociación indican que un eventual acuerdo se someterá a revisión por la CNMC.",
        },
    ]
    return render_template("noticias.html", noticias=noticias_lista)


@app.route("/admin")
def admin():
    return render_template("admin.html")


if __name__ == "__main__":
    _arranque_logger = _logging.getLogger("dfin.app")
    _arranque_logger.info(
        "ARRANQUE python app.py TICKER_FIJO=%s GMAIL=%s LOG_LEVEL=%s",
        TICKER_FIJO or "(ninguno)",
        "configurado" if (GMAIL_USER and GMAIL_APP_PASSWORD) else "no_configurado",
        _LOG_LEVEL,
    )
    if TICKER_FIJO:
        # Modo compatibilidad: se pasó ticker por argv. Precargamos.
        print(f"\n[DFin AI] Precargando datos de Yahoo Finance para {TICKER_FIJO}…")
        try:
            _cargar_datos_ticker_fijo()
            nombre = _BRANDING["empresa"]
            print(f"[DFin AI] Aplicación dedicada a: {nombre} ({TICKER_FIJO})\n")
            _arranque_logger.info("Datos de Yahoo precargados: %s (%s)", nombre, TICKER_FIJO)
        except Exception as _e:
            print(f"[DFin AI] AVISO: no se pudieron precargar datos de {TICKER_FIJO}: {_e}\n")
            _arranque_logger.exception("Fallo al precargar Yahoo para %s", TICKER_FIJO)

        print(f"[DFin AI] Indexando corpus RAG de {TICKER_FIJO}"
              f"{' (--recrawl forzado)' if RECRAWL_AL_ARRANCAR else ''}…")
        try:
            idx = _construir_rag_si_procede(forzar=RECRAWL_AL_ARRANCAR)
            n = len(idx.chunks) if idx else 0
            print(f"[DFin AI] RAG: {n} fragmentos indexados.\n")
        except Exception as _e:
            print(f"[DFin AI] AVISO: fallo indexando corpus RAG: {_e}\n")
            _arranque_logger.exception("Fallo en RAG para %s", TICKER_FIJO)

        _cargar_memoria_ir()
    else:
        # Modo normal: sin ticker, el usuario lo elegirá en la web.
        print("\n[DFin AI] Aplicación lista. Abre http://localhost:5000 y "
              "selecciona la empresa.\n")
        _arranque_logger.info("Arranque sin TICKER_FIJO: selección desde la web.")

    app.run(debug=True, host="0.0.0.0", port=5000)
