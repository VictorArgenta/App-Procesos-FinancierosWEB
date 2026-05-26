"""RAG ligero para el módulo Relación con Inversores.

Lee PDFs locales de `Empresas/{TICKER}/docs/`, crawlea las URLs corporativas
listadas en `Empresas/{TICKER}/docs/url_inversores.txt`, indexa todo con BM25
y permite consultas top-K para inyectar contexto relevante al prompt del LLM.

Todo en local, sin LLM ni embeddings: cero coste de tokens en la indexación.
Sólo se consumen tokens cuando el LLM consume los chunks recuperados (igual
que hoy).
"""
from __future__ import annotations

import hashlib
import json
import logging
import math
import re
import time
import urllib.error
import urllib.parse
import urllib.request
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable

log = logging.getLogger("dfin.rag")

# --- Configuración global ---------------------------------------------------

CRAWL_PROFUNDIDAD = 2
CRAWL_MAX_PAGINAS = 50
CRAWL_MAX_PDFS = 30
CRAWL_MAX_PDF_BYTES = 25 * 1024 * 1024  # 25 MB
CHUNK_PALABRAS = 600
CHUNK_SOLAPE = 100
TOP_K_DEFECTO = 8
HTTP_TIMEOUT = 15
USER_AGENT = "DFin-AI-RAG/1.0 (+contact: ir@dfin.local)"
TTL_CACHE_SEG = 24 * 3600  # 24 h


# Stopwords es+en mínimas para que BM25 no se distorsione con palabras vacías.
_STOPWORDS = {
    # Español
    "a", "al", "ante", "bajo", "cabe", "con", "contra", "de", "del", "desde",
    "en", "entre", "hacia", "hasta", "para", "por", "según", "sin", "sobre",
    "tras", "el", "la", "los", "las", "un", "una", "unos", "unas", "y", "o",
    "u", "e", "ni", "que", "qué", "como", "cómo", "cuando", "cuándo", "donde",
    "dónde", "porque", "pero", "ese", "esa", "eso", "este", "esta", "esto",
    "aquel", "aquella", "aquello", "es", "son", "fue", "ser", "ha", "han",
    "hay", "su", "sus", "lo", "le", "les", "me", "te", "se", "nos", "vos",
    "yo", "tú", "él", "ella", "nosotros", "ustedes", "ellos", "ellas",
    "muy", "mas", "más", "menos", "sí", "no", "ya", "aún", "también",
    # English (un mínimo, por si los PDFs vienen en inglés)
    "the", "a", "an", "and", "or", "but", "if", "of", "in", "on", "at",
    "to", "for", "from", "by", "with", "is", "are", "was", "were", "be",
    "been", "being", "have", "has", "had", "do", "does", "did", "this",
    "that", "these", "those", "it", "its", "as", "than", "then", "so",
    "not", "no",
}

# --- Utilidades de extracción ----------------------------------------------


def _tokenizar(texto: str) -> list[str]:
    """Tokenización simple para BM25: minúsculas, alfanumérico, sin stopwords."""
    if not texto:
        return []
    texto = texto.lower()
    tokens = re.findall(r"[a-záéíóúñü0-9]{2,}", texto, flags=re.IGNORECASE)
    return [t for t in tokens if t not in _STOPWORDS]


def _chunk_texto(texto: str, palabras_por_chunk: int = CHUNK_PALABRAS,
                  solape: int = CHUNK_SOLAPE) -> list[str]:
    """Divide un texto en chunks de N palabras con solape entre chunks."""
    if not texto or not texto.strip():
        return []
    palabras = texto.split()
    if len(palabras) <= palabras_por_chunk:
        return [texto.strip()]
    chunks = []
    paso = max(palabras_por_chunk - solape, 1)
    for i in range(0, len(palabras), paso):
        trozo = palabras[i : i + palabras_por_chunk]
        if not trozo:
            break
        chunks.append(" ".join(trozo).strip())
        if i + palabras_por_chunk >= len(palabras):
            break
    return chunks


def _extraer_texto_pdf(ruta: Path) -> str:
    """Extrae el texto de un PDF con pypdf. Devuelve string vacío si falla."""
    try:
        from pypdf import PdfReader
    except ImportError:
        log.error("pypdf no instalado; no se pueden leer PDFs. Ejecuta: pip install pypdf")
        return ""
    try:
        reader = PdfReader(str(ruta))
        partes = []
        for pagina in reader.pages:
            try:
                partes.append(pagina.extract_text() or "")
            except Exception:
                continue
        return "\n\n".join(partes).strip()
    except Exception as e:
        log.warning("No se pudo leer PDF %s: %s", ruta, e)
        return ""


def _extraer_paginas_pdf(ruta: Path) -> list[tuple[int, str]]:
    """Extrae el texto del PDF página a página.

    Devuelve [(n_pagina_1based, texto), ...]. Si falla, lista vacía.
    Necesario para que cada chunk lleve el número de página en su meta.
    """
    try:
        from pypdf import PdfReader
    except ImportError:
        return []
    try:
        reader = PdfReader(str(ruta))
        paginas = []
        for i, pagina in enumerate(reader.pages, start=1):
            try:
                txt = (pagina.extract_text() or "").strip()
            except Exception:
                txt = ""
            paginas.append((i, txt))
        return paginas
    except Exception as e:
        log.warning("No se pudo leer páginas del PDF %s: %s", ruta, e)
        return []


def _extraer_texto_docx(ruta: Path) -> str:
    """Extrae texto de un .docx (Word). Incluye párrafos y celdas de tablas."""
    try:
        from docx import Document
    except ImportError:
        log.error("python-docx no instalado; no se pueden leer DOCX.")
        return ""
    try:
        doc = Document(str(ruta))
        partes = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                partes.append(t)
        for tabla in doc.tables:
            for fila in tabla.rows:
                celdas = [c.text.strip() for c in fila.cells if c.text and c.text.strip()]
                if celdas:
                    partes.append(" | ".join(celdas))
        return "\n\n".join(partes).strip()
    except Exception as e:
        log.warning("No se pudo leer DOCX %s: %s", ruta, e)
        return ""


def _extraer_texto_html_local(ruta: Path) -> str:
    """Extrae texto de un .html/.htm local reutilizando el limpiador del crawler."""
    try:
        html = ruta.read_text(encoding="utf-8", errors="replace")
    except OSError as e:
        log.warning("No se pudo leer HTML %s: %s", ruta, e)
        return ""
    return _html_a_texto(html)


def _extraer_texto_plano(ruta: Path) -> str:
    """Lee un fichero de texto plano (.txt, .md)."""
    try:
        return ruta.read_text(encoding="utf-8", errors="replace").strip()
    except OSError as e:
        log.warning("No se pudo leer fichero plano %s: %s", ruta, e)
        return ""


# Mapping extensión → extractor.
_EXTRACTORES = {
    ".pdf": _extraer_texto_pdf,
    ".docx": _extraer_texto_docx,
    ".html": _extraer_texto_html_local,
    ".htm": _extraer_texto_html_local,
    ".txt": _extraer_texto_plano,
    ".md": _extraer_texto_plano,
}

EXTENSIONES_SOPORTADAS = tuple(_EXTRACTORES.keys())


def extraer_texto_documento(ruta: Path) -> str:
    """Despachador general: detecta extensión y llama al extractor adecuado.

    Devuelve string vacío si el formato no está soportado o si la extracción
    falla. No lanza excepciones para no romper la indexación cuando un único
    documento es defectuoso.
    """
    ext = ruta.suffix.lower()
    extractor = _EXTRACTORES.get(ext)
    if not extractor:
        log.warning("Formato no soportado, ignoro %s (extensión %s)", ruta.name, ext)
        return ""
    return extractor(ruta)


def _html_a_texto(html: str) -> str:
    """Convierte HTML a texto limpio (sin scripts/estilos/navs/footers)."""
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        log.error("beautifulsoup4 no instalado; falla extracción HTML.")
        return ""
    sopa = BeautifulSoup(html, "html.parser")
    for tag in sopa(["script", "style", "noscript", "nav", "footer", "header", "aside"]):
        tag.decompose()
    texto = sopa.get_text(separator="\n")
    # Compacta múltiples saltos y espacios
    texto = re.sub(r"[ \t]+", " ", texto)
    texto = re.sub(r"\n{3,}", "\n\n", texto)
    return texto.strip()


def _titulo_html(html: str) -> str:
    try:
        from bs4 import BeautifulSoup
        sopa = BeautifulSoup(html, "html.parser")
        if sopa.title and sopa.title.string:
            return sopa.title.string.strip()[:140]
        h1 = sopa.find("h1")
        if h1 and h1.get_text():
            return h1.get_text().strip()[:140]
    except Exception:
        pass
    return ""


# --- Crawler ---------------------------------------------------------------


def _descargar(url: str, max_bytes: int | None = None) -> tuple[bytes, str] | None:
    """Descarga una URL devolviendo (bytes, content_type) o None si falla."""
    try:
        req = urllib.request.Request(url, headers={"User-Agent": USER_AGENT})
        with urllib.request.urlopen(req, timeout=HTTP_TIMEOUT) as resp:
            ctype = resp.headers.get("Content-Type", "").lower()
            data = resp.read(max_bytes + 1) if max_bytes else resp.read()
            if max_bytes and len(data) > max_bytes:
                log.warning("Descarga %s descartada: supera %d bytes", url, max_bytes)
                return None
            return data, ctype
    except (urllib.error.URLError, urllib.error.HTTPError, TimeoutError, ValueError) as e:
        log.warning("Fallo descargando %s: %s", url, e)
        return None
    except Exception as e:
        log.warning("Excepción inesperada descargando %s: %s", url, e)
        return None


def _es_misma_base(url: str, base: str) -> bool:
    """True si url está bajo el mismo dominio que base."""
    try:
        u, b = urllib.parse.urlparse(url), urllib.parse.urlparse(base)
        return (u.netloc or b.netloc) == b.netloc
    except Exception:
        return False


def _normalizar_url(url: str) -> str:
    """Quita fragmentos y normaliza."""
    try:
        u = urllib.parse.urlparse(url)
        return urllib.parse.urlunparse(u._replace(fragment=""))
    except Exception:
        return url


def _enlaces_de_html(html: str, base_url: str) -> list[str]:
    """Extrae todos los <a href> resueltos contra la URL base."""
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        return []
    sopa = BeautifulSoup(html, "html.parser")
    out = []
    for a in sopa.find_all("a", href=True):
        href = a["href"]
        if not href or href.startswith(("mailto:", "tel:", "javascript:")):
            continue
        absu = urllib.parse.urljoin(base_url, href)
        out.append(_normalizar_url(absu))
    return out


def _es_pdf(url: str, content_type: str = "") -> bool:
    return url.lower().endswith(".pdf") or "application/pdf" in (content_type or "")


def _es_descartable(url: str) -> bool:
    """Extensiones que se ignoran (imágenes, vídeos, scripts...)."""
    extens_no = (".png", ".jpg", ".jpeg", ".gif", ".webp", ".svg", ".ico",
                 ".mp4", ".webm", ".mov", ".mp3", ".wav", ".zip", ".rar",
                 ".css", ".js", ".woff", ".woff2", ".ttf", ".eot")
    return url.lower().endswith(extens_no)


def _crawl_url_base(url_base: str,
                     cache_dir: Path,
                     profundidad: int = CRAWL_PROFUNDIDAD,
                     max_paginas: int = CRAWL_MAX_PAGINAS,
                     max_pdfs: int = CRAWL_MAX_PDFS) -> list[dict]:
    """Crawl BFS limitado a un dominio. Devuelve lista de documentos en bruto.

    Cada documento es {tipo, url, titulo, texto, fuente_path} donde fuente_path
    es la ruta local a la copia descargada para PDFs (en `.cache/web_pdfs/`),
    None para páginas HTML.
    """
    docs = []
    visitadas: set[str] = set()
    pendientes: list[tuple[str, int]] = [(url_base, 0)]
    pdfs_descargados = 0

    pdfs_dir = cache_dir / "web_pdfs"
    pdfs_dir.mkdir(parents=True, exist_ok=True)

    while pendientes and len(visitadas) < max_paginas:
        url, prof = pendientes.pop(0)
        if url in visitadas or _es_descartable(url):
            continue
        if not _es_misma_base(url, url_base):
            continue
        visitadas.add(url)

        if _es_pdf(url):
            if pdfs_descargados >= max_pdfs:
                continue
            r = _descargar(url, max_bytes=CRAWL_MAX_PDF_BYTES)
            if not r:
                continue
            data, _ = r
            nombre = hashlib.sha1(url.encode("utf-8")).hexdigest()[:16] + ".pdf"
            ruta = pdfs_dir / nombre
            ruta.write_bytes(data)
            texto = _extraer_texto_pdf(ruta)
            if texto:
                docs.append({
                    "tipo": "pdf_web",
                    "url": url,
                    "titulo": Path(urllib.parse.urlparse(url).path).name or url,
                    "texto": texto,
                    "fuente_path": str(ruta),
                })
                pdfs_descargados += 1
                log.info("Crawl: PDF descargado %s (%d KB)", url, len(data) // 1024)
            continue

        r = _descargar(url)
        if not r:
            continue
        data, ctype = r
        if "text/html" not in (ctype or "").lower() and not url.lower().endswith((".html", ".htm")):
            continue
        try:
            html = data.decode("utf-8", errors="replace")
        except Exception:
            continue
        texto = _html_a_texto(html)
        if texto:
            docs.append({
                "tipo": "web",
                "url": url,
                "titulo": _titulo_html(html) or url,
                "texto": texto,
                "fuente_path": None,
            })
        if prof < profundidad:
            for enlace in _enlaces_de_html(html, url):
                if enlace not in visitadas and _es_misma_base(enlace, url_base):
                    pendientes.append((enlace, prof + 1))

    log.info(
        "Crawl %s: %d páginas HTML, %d PDFs descargados",
        url_base, len([d for d in docs if d["tipo"] == "web"]), pdfs_descargados,
    )
    return docs


# --- Índice BM25 -----------------------------------------------------------


class BM25Index:
    """Índice BM25 ligero con persistencia JSON."""

    K1 = 1.5
    B = 0.75

    def __init__(self):
        self.chunks: list[dict] = []  # cada uno: {texto, meta, tokens, length}
        self.df: Counter = Counter()
        self.avgdl: float = 0.0
        self.idf: dict[str, float] = {}

    def añadir(self, texto: str, meta: dict) -> None:
        tokens = _tokenizar(texto)
        if not tokens:
            return
        self.chunks.append({"texto": texto, "meta": meta, "tokens": tokens, "length": len(tokens)})

    def construir(self) -> None:
        if not self.chunks:
            return
        self.df = Counter()
        for ch in self.chunks:
            for term in set(ch["tokens"]):
                self.df[term] += 1
        n = len(self.chunks)
        self.avgdl = sum(c["length"] for c in self.chunks) / n
        self.idf = {
            term: math.log(1 + (n - df + 0.5) / (df + 0.5))
            for term, df in self.df.items()
        }
        # Pre-calculamos term frequency por chunk para acelerar consultas.
        for ch in self.chunks:
            ch["tf"] = Counter(ch["tokens"])

    def buscar(self, query: str, top_k: int = TOP_K_DEFECTO) -> list[tuple[float, dict]]:
        if not self.chunks:
            return []
        terms = _tokenizar(query)
        if not terms:
            return []
        scores = []
        for ch in self.chunks:
            score = 0.0
            for t in terms:
                if t not in self.idf:
                    continue
                tf = ch["tf"].get(t, 0)
                if tf == 0:
                    continue
                num = tf * (self.K1 + 1)
                den = tf + self.K1 * (1 - self.B + self.B * ch["length"] / self.avgdl)
                score += self.idf[t] * (num / den)
            if score > 0:
                scores.append((score, ch))
        scores.sort(key=lambda x: x[0], reverse=True)
        return scores[:top_k]

    def serializar(self) -> dict:
        return {
            "chunks": [
                {"texto": c["texto"], "meta": c["meta"]} for c in self.chunks
            ],
            "construido_en": datetime.now(timezone.utc).isoformat(),
        }

    @classmethod
    def cargar(cls, data: dict) -> "BM25Index":
        idx = cls()
        for c in data.get("chunks", []):
            idx.añadir(c["texto"], c.get("meta", {}))
        idx.construir()
        return idx


# --- Pipeline completo por ticker ------------------------------------------


_CLAVES_CONFIG_EMPRESA = {"cotiza", "ticker"}


def _leer_urls_inversores(ruta_txt: Path) -> list[str]:
    if not ruta_txt.exists():
        return []
    urls = []
    for linea in ruta_txt.read_text(encoding="utf-8").splitlines():
        s = linea.strip()
        if not s or s.startswith("#"):
            continue
        # El fichero comparte config de empresa (cotiza=, ticker=) con las URLs;
        # esas líneas no son crawleables.
        if "://" not in s and "=" in s:
            clave = s.split("=", 1)[0].strip().lower()
            if clave in _CLAVES_CONFIG_EMPRESA:
                continue
        urls.append(s)
    return urls


def _leer_cache_incremental(cache_idx: Path) -> dict | None:
    """Carga el caché incremental (formato v2) si existe y es legible.

    Estructura esperada:
        {
          "version": 2,
          "documentos_locales": {nombre: {"hash_md5", "chunks": [...]}},
          "web": {"completado_en", "url_inversores_hash", "chunks": [...]}
        }
    Devuelve None si el caché no existe, es de versión vieja o está corrupto.
    """
    if not cache_idx.exists():
        return None
    try:
        data = json.loads(cache_idx.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as e:
        log.warning("Caché RAG ilegible (%s); reconstruyo desde cero.", e)
        return None
    # v3 añade pagina_inicio/pagina_fin (PDFs) y seccion (resto) en chunks.
    # Cachés v2 carecen de esos campos y los chunks no pueden citar página.
    if data.get("version") != 3:
        log.info("Caché RAG con formato antiguo; reconstruyo en v3 con paginación.")
        return None
    if not isinstance(data.get("documentos_locales"), dict):
        return None
    return data


def _md5_fichero(ruta: Path) -> str | None:
    try:
        return hashlib.md5(ruta.read_bytes()).hexdigest()
    except OSError:
        return None


def _chunks_de_documento(doc: Path) -> list[dict]:
    """Extrae texto del documento y devuelve la lista de chunks con metadata.

    Para PDFs cada chunk lleva `pagina_inicio` y `pagina_fin` (1-based)
    para que la respuesta del módulo de Consulta Documental pueda citar
    "documento.pdf, página 12". Para el resto de formatos no hay
    concepto de página, así que llevan `seccion` (índice del chunk).
    """
    ext_lower = doc.suffix.lower().lstrip(".")
    meta_base = {
        "fuente_tipo": f"{ext_lower}_local",
        "fuente_nombre": doc.name,
        "fuente_path": str(doc),
    }

    if doc.suffix.lower() == ".pdf":
        paginas = _extraer_paginas_pdf(doc)
        if not paginas:
            return []
        # Construimos un texto continuo concatenando páginas, pero
        # arrastramos los rangos de palabras [inicio, fin) para saber
        # luego de qué páginas viene cada chunk.
        palabras = []
        rangos_pag = []  # [(pagina, idx_palabra_inicio, idx_palabra_fin), ...]
        for n_pag, texto_pag in paginas:
            if not texto_pag:
                continue
            words_pag = texto_pag.split()
            if not words_pag:
                continue
            rangos_pag.append((n_pag, len(palabras), len(palabras) + len(words_pag)))
            palabras.extend(words_pag)
        if not palabras:
            return []

        chunks_meta = []
        paso = max(CHUNK_PALABRAS - CHUNK_SOLAPE, 1)
        idx = 0
        chunk_i = 0
        while idx < len(palabras):
            ini = idx
            fin = min(idx + CHUNK_PALABRAS, len(palabras))
            texto_chunk = " ".join(palabras[ini:fin]).strip()
            pag_ini = _pagina_de_palabra(rangos_pag, ini)
            pag_fin = _pagina_de_palabra(rangos_pag, max(fin - 1, ini))
            meta = {
                **meta_base,
                "fragmento": chunk_i,
                "pagina_inicio": pag_ini,
                "pagina_fin": pag_fin,
            }
            chunks_meta.append({"texto": texto_chunk, "meta": meta})
            chunk_i += 1
            if fin >= len(palabras):
                break
            idx += paso
        return chunks_meta

    # Resto de formatos: sin paginación; usamos índice de sección.
    texto = extraer_texto_documento(doc)
    if not texto:
        return []
    return [
        {
            "texto": chunk,
            "meta": {**meta_base, "fragmento": i, "seccion": i + 1},
        }
        for i, chunk in enumerate(_chunk_texto(texto))
    ]


def _pagina_de_palabra(rangos_pag: list[tuple[int, int, int]], idx_pal: int) -> int:
    """Dado un índice de palabra global, devuelve la página a la que pertenece."""
    for n_pag, ini, fin in rangos_pag:
        if ini <= idx_pal < fin:
            return n_pag
    # Si no cae en ningún rango (margen), devolvemos el último visto.
    return rangos_pag[-1][0] if rangos_pag else 1


def indexar_ticker(ticker: str, root: Path | None = None, forzar: bool = False) -> BM25Index:
    """Indexa el corpus de un ticker con caché INCREMENTAL.

    Reaprovecha los chunks de documentos cuyo hash MD5 no ha cambiado desde
    la última indexación; solo reprocesa los nuevos o modificados. El crawl
    web sigue siendo "todo o nada" (con TTL 24h) porque el contenido web es
    inherentemente cambiante.

    Si `forzar=True`, ignora completamente el caché y reconstruye desde cero.
    """
    root = root or Path(__file__).resolve().parent
    base = root / "Empresas" / ticker
    docs_dir = base / "docs"
    urls_txt = docs_dir / "url_inversores.txt"
    cache_dir = base / ".cache"
    cache_idx = cache_dir / "index.json"

    if not base.exists():
        log.info("RAG ticker=%s: carpeta Empresas/%s no existe, sin corpus.", ticker, ticker)
        idx = BM25Index()
        idx.construir()
        return idx

    cache_dir.mkdir(parents=True, exist_ok=True)
    cache = None if forzar else _leer_cache_incremental(cache_idx)
    docs_cache = (cache or {}).get("documentos_locales", {}) or {}
    web_cache = (cache or {}).get("web", {}) or {}

    log.info("RAG ticker=%s: indexando corpus (incremental)...", ticker)
    t0 = time.time()

    # -- 1) Documentos locales ------------------------------------------------
    nuevos_docs_cache = {}
    n_reusados, n_nuevos, n_modificados, n_eliminados = 0, 0, 0, 0
    hashes_vistos: set[str] = set()

    ficheros_locales = []
    if docs_dir.exists():
        for ext in EXTENSIONES_SOPORTADAS:
            ficheros_locales.extend(docs_dir.glob(f"*{ext}"))

    nombres_actuales = set()
    for doc in sorted(ficheros_locales):
        if doc.name == "url_inversores.txt":
            continue
        nombres_actuales.add(doc.name)
        md5 = _md5_fichero(doc)
        if md5 is None:
            continue
        hashes_vistos.add(md5)
        entrada_cache = docs_cache.get(doc.name)
        if entrada_cache and entrada_cache.get("hash_md5") == md5 and entrada_cache.get("chunks"):
            # Reutilizamos los chunks ya cacheados
            nuevos_docs_cache[doc.name] = entrada_cache
            n_reusados += 1
        else:
            # Nuevo o modificado: extraer + chunkear
            chunks = _chunks_de_documento(doc)
            nuevos_docs_cache[doc.name] = {"hash_md5": md5, "chunks": chunks}
            if entrada_cache:
                n_modificados += 1
            else:
                n_nuevos += 1

    # Documentos del caché que ya no están en disco → se eliminan
    eliminados_nombres = set(docs_cache.keys()) - nombres_actuales
    n_eliminados = len(eliminados_nombres)

    # -- 2) Web (crawl + descargas) ------------------------------------------
    urls = _leer_urls_inversores(urls_txt)
    url_hash = hashlib.sha1(("\n".join(urls)).encode("utf-8")).hexdigest()[:16] if urls else ""
    web_completado_en = web_cache.get("completado_en", 0) or 0
    web_url_hash = web_cache.get("url_inversores_hash", "")
    web_edad = time.time() - web_completado_en
    web_cache_valido = (
        not forzar
        and urls
        and web_url_hash == url_hash
        and web_completado_en > 0
        and web_edad < TTL_CACHE_SEG
        and isinstance(web_cache.get("chunks"), list)
    )

    if web_cache_valido:
        web_chunks = web_cache["chunks"]
        log.info(
            "RAG ticker=%s: web cacheada (edad %.1fh, %d chunks)",
            ticker, web_edad / 3600, len(web_chunks),
        )
    else:
        log.info("RAG ticker=%s: re-crawleando web (%d URLs)", ticker, len(urls))
        docs_web = []
        for u in urls:
            try:
                docs_web.extend(_crawl_url_base(u, cache_dir))
            except Exception as e:
                log.exception("Fallo crawleando %s: %s", u, e)
        web_chunks = []
        for d in docs_web:
            # Dedup PDFs web contra los locales por hash
            if d["tipo"] == "pdf_web" and d.get("fuente_path"):
                try:
                    h = hashlib.md5(Path(d["fuente_path"]).read_bytes()).hexdigest()
                    if h in hashes_vistos:
                        log.info("PDF web %s duplicado de uno local; ignorado", d["url"])
                        continue
                    hashes_vistos.add(h)
                except OSError:
                    pass
            for i, chunk in enumerate(_chunk_texto(d["texto"])):
                web_chunks.append({
                    "texto": chunk,
                    "meta": {
                        "fuente_tipo": d["tipo"],
                        "fuente_nombre": d["titulo"],
                        "fuente_url": d["url"],
                        "fragmento": i,
                    },
                })

    # -- 3) Construcción del índice BM25 -------------------------------------
    idx = BM25Index()
    n_chunks_locales = 0
    for entry in nuevos_docs_cache.values():
        for ch in entry.get("chunks", []):
            idx.añadir(ch["texto"], ch.get("meta", {}))
            n_chunks_locales += 1
    for ch in web_chunks:
        idx.añadir(ch["texto"], ch.get("meta", {}))
    idx.construir()

    # -- 4) Persistir caché incremental --------------------------------------
    try:
        payload = {
            "version": 3,
            "ticker": ticker,
            "construido_en": datetime.now(timezone.utc).isoformat(),
            "documentos_locales": nuevos_docs_cache,
            "web": {
                "completado_en": web_cache.get("completado_en") if web_cache_valido else time.time(),
                "url_inversores_hash": url_hash,
                "chunks": web_chunks,
            },
        }
        cache_idx.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception as e:
        log.warning("No se pudo persistir caché RAG: %s", e)

    log.info(
        "RAG ticker=%s indexado: docs locales %d (%d reusados, %d nuevos, %d modificados, %d eliminados) · "
        "%d chunks locales + %d chunks web · web %s · %.1fs",
        ticker, len(nuevos_docs_cache), n_reusados, n_nuevos, n_modificados, n_eliminados,
        n_chunks_locales, len(web_chunks),
        "cacheada" if web_cache_valido else "re-crawleada",
        time.time() - t0,
    )
    return idx


# --- Helpers de presentación ----------------------------------------------


_ETIQUETAS_LOCALES = {
    "pdf_local":  "PDF local",
    "docx_local": "Word local",
    "html_local": "HTML local",
    "txt_local":  "Texto local",
    "md_local":   "Markdown local",
}


def _etiqueta_fuente(meta: dict) -> str:
    tipo = meta.get("fuente_tipo", "")
    nombre = meta.get("fuente_nombre", "")
    if tipo in _ETIQUETAS_LOCALES:
        return f"{_ETIQUETAS_LOCALES[tipo]} — {nombre}"
    if tipo == "pdf_web":
        return f"PDF descargado de web — {nombre} ({meta.get('fuente_url', '')})"
    if tipo == "web":
        return f"Web corporativa — {nombre} ({meta.get('fuente_url', '')})"
    return f"{nombre} ({meta.get('fuente_url', '')})"


def formatear_chunks_para_prompt(resultados: list[tuple[float, dict]]) -> str:
    """Devuelve los chunks listos para inyectar en el prompt."""
    if not resultados:
        return ""
    partes = ["FRAGMENTOS RELEVANTES DE DOCUMENTACIÓN CORPORATIVA",
              "(consulta esta fuente PRIMERO; tiene prioridad sobre Yahoo y web abierta):", ""]
    for score, ch in resultados:
        partes.append(f"[Fuente: {_etiqueta_fuente(ch.get('meta', {}))} · score={score:.2f}]")
        partes.append(ch["texto"])
        partes.append("")
    return "\n".join(partes)


def resumen_fuentes_para_nota_interna(resultados: list[tuple[float, dict]]) -> str:
    """Devuelve una línea por fuente usada, para mostrar al usuario en la UI."""
    if not resultados:
        return ""
    iconos_locales = {
        "pdf_local": "📄", "docx_local": "📝", "html_local": "🌐",
        "txt_local": "📋", "md_local": "📋",
    }
    vistas = []
    seen = set()
    for _, ch in resultados:
        meta = ch.get("meta", {})
        tipo = meta.get("fuente_tipo", "")
        nombre = meta.get("fuente_nombre", "")
        if tipo in iconos_locales:
            etiqueta = f"{iconos_locales[tipo]} {nombre} ({_ETIQUETAS_LOCALES.get(tipo, 'local')})"
        elif tipo == "pdf_web":
            etiqueta = f"📄 {nombre} (PDF desde {meta.get('fuente_url', '')})"
        else:
            etiqueta = f"🌐 {nombre} ({meta.get('fuente_url', '')})"
        if etiqueta in seen:
            continue
        seen.add(etiqueta)
        vistas.append(etiqueta)
    return "\n".join(vistas)
